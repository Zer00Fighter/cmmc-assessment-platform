from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from src.ssp_export import SSPExportMetadata, export_ssp
from src.ssp_export.word_exporter import INPUT_REQUIRED


def _create_template(path: Path) -> None:
    document = Document()
    document.add_heading("ACME System Security Plan", level=0)
    document.add_paragraph("Prepared for ACME")
    document.sections[0].header.paragraphs[0].text = "ACME SSP"
    practice = document.add_table(rows=2, cols=7)
    practice.cell(0, 0).text = "Requirement Conformity:"
    practice.cell(0, 1).merge(
        practice.cell(0, 6)
    ).text = "MET    NOT MET    NOT APPLICABLE"
    practice.cell(1, 0).text = "CMMC Level:"
    practice.cell(1, 1).text = "L2"
    practice.cell(1, 2).text = "Practice ID:"
    practice.cell(1, 3).text = "AC.L2-3.1.1"
    practice.cell(1, 4).merge(practice.cell(1, 5)).text = "Practice Name:"
    practice.cell(1, 6).text = "Authorized Access Control"
    objective = practice.add_row()
    objective.cells[0].merge(
        objective.cells[4]
    ).text = "[a] authorized users are identified."
    objective.cells[5].merge(objective.cells[6]).text = "Met    Not Met    N/A"
    statement = practice.add_row()
    statement.cells[0].merge(
        statement.cells[6]
    ).text = "Assessment Objective Conformity Statement:"
    artifacts = document.add_table(rows=7, cols=1)
    artifacts.cell(0, 0).text = "Supporting Artifacts"
    artifacts.cell(1, 0).text = "System Design Documentation"
    artifacts.cell(2, 0).text = "Template guidance"
    artifacts.cell(3, 0).text = (
        "System Configuration Settings And Associated Documentation"
    )
    artifacts.cell(4, 0).text = "Template guidance"
    artifacts.cell(5, 0).text = "Supplemental Artifacts"
    artifacts.cell(6, 0).text = "Template guidance"
    document.save(path)


def _create_workbook(path: Path) -> None:
    workbook = Workbook()
    cover = workbook.active
    cover.title = "Cover"
    cover.cell(6, 2, "Organization Name")
    cover.cell(6, 3, "Workbook Organization")
    cover.cell(7, 2, "Assessment Name")
    cover.cell(7, 3, "Enterprise System")
    cover.cell(14, 2, "Workbook Version")
    cover.cell(14, 3, "1.0.0")

    assessment = workbook.create_sheet("Assessment")
    assessment.cell(5, 2, "Requirement ID")
    assessment.cell(6, 1, "AC")
    assessment.cell(6, 2, "AC.L2-3.1.1")
    assessment.cell(6, 3, "Authorized Access Control")
    assessment.cell(6, 4, "Limit access to authorized users and devices.")
    assessment.cell(6, 9, "MET")
    assessment.cell(6, 12, 5)
    assessment.cell(6, 16, "Security Officer")
    assessment.cell(6, 17, "SSP 3.1.1")
    assessment.cell(6, 18, "Access is limited through approved identities and devices.")

    crosswalk = workbook.create_sheet("SSP Crosswalk")
    crosswalk.cell(5, 2, "Requirement ID")
    crosswalk.cell(6, 1, "AC")
    crosswalk.cell(6, 2, "AC.L2-3.1.1")
    crosswalk.cell(6, 3, "SSP 3.1.1")
    crosswalk.cell(6, 4, "Access Control Policy")
    crosswalk.cell(6, 5, "Access review records")
    crosswalk.cell(6, 6, "Security Officer")
    crosswalk.cell(6, 7, "Mapped")
    workbook.save(path)


def _all_text(document: Document) -> str:
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text.extend(cell.text for cell in row.cells)
    for section in document.sections:
        text.extend(paragraph.text for paragraph in section.header.paragraphs)
    return "\n".join(text)


def test_export_preserves_template_and_adds_omni_crosswalk(tmp_path: Path) -> None:
    template = tmp_path / "source.docx"
    workbook = tmp_path / "omni.xlsx"
    output = tmp_path / "export.docx"
    _create_template(template)
    _create_workbook(workbook)
    original_hash = hashlib.sha256(template.read_bytes()).hexdigest()

    result = export_ssp(
        template,
        workbook,
        output,
        SSPExportMetadata(organization_name="R!SC", prepared_by="Omni Team"),
    )

    assert result == output.resolve()
    assert hashlib.sha256(template.read_bytes()).hexdigest() == original_hash
    document = Document(output)
    text = _all_text(document)
    assert "ACME" not in text
    assert "R!SC System Security Plan" in text
    assert "Omni Export Information" not in text
    assert "Omni Security Plan Crosswalk" not in text
    assert "AC.L2-3.1.1" in text
    assert "Access Control Policy" in text
    assert "Access review records" in text
    assert "Test Procedure" not in text
    assert INPUT_REQUIRED in text
    control_table = next(
        table for table in document.tables if "AC.L2-3.1.1" in table.cell(1, 6).text
    )
    assert control_table.cell(1, 0).text == "CMMC Level:"
    assert control_table.cell(1, 1).text == "L2"
    assert control_table.cell(1, 2).text == "SPRS Score:"
    assert control_table.cell(1, 3).text == "0"
    assert control_table.cell(1, 4).text == "Practice ID:"
    assert control_table.cell(1, 6).text == "AC.L2-3.1.1"
    assert control_table.cell(2, 0).text == "Practice Name:"
    assert control_table.cell(2, 2).text == "Authorized Access Control"
    assert "[X] MET" in control_table.cell(0, 1).text
    assert "[X] Met" in control_table.cell(3, 6).text
    assert control_table.cell(4, 0).text == (
        "Access is limited through approved identities and devices."
    )


def test_export_rejects_overwriting_source_template(tmp_path: Path) -> None:
    template = tmp_path / "source.docx"
    workbook = tmp_path / "omni.xlsx"
    _create_template(template)
    _create_workbook(workbook)

    try:
        export_ssp(template, workbook, template)
    except ValueError as error:
        assert "must differ" in str(error)
    else:
        raise AssertionError("Expected source-template overwrite to be rejected")


def test_not_met_uses_assessor_notes_as_finding(tmp_path: Path) -> None:
    template = tmp_path / "source.docx"
    workbook_path = tmp_path / "omni.xlsx"
    output = tmp_path / "export.docx"
    _create_template(template)
    _create_workbook(workbook_path)
    workbook = load_workbook(workbook_path)
    workbook["Assessment"]["I6"] = "NOT MET"
    workbook["Assessment"]["R6"] = "Authorized-user evidence was incomplete."
    workbook.save(workbook_path)

    export_ssp(template, workbook_path, output)

    document = Document(output)
    control_table = next(
        table for table in document.tables if "AC.L2-3.1.1" in table.cell(1, 6).text
    )
    assert control_table.cell(1, 3).text == "5"
    assert "[X] NOT MET" in control_table.cell(0, 1).text
    assert "[X] Not Met" in control_table.cell(3, 6).text
    assert control_table.cell(4, 0).text == "Authorized-user evidence was incomplete."
