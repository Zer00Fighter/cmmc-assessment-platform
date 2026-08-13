"""Export Omni's Security Plan crosswalk into a Word SSP template."""

from __future__ import annotations

import os
import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

INPUT_REQUIRED = "[REQUIRES ORGANIZATION INPUT]"


@dataclass(frozen=True)
class SSPExportMetadata:
    """Organization-specific values used to brand an SSP export."""

    organization_name: str = ""
    system_name: str = ""
    system_owner: str = ""
    prepared_by: str = ""
    version: str = "1.0"
    export_date: str = ""


@dataclass(frozen=True)
class SSPControlRow:
    domain: str
    requirement_id: str
    title: str
    statement: str
    ssp_reference: str
    governance_references: str
    evidence_references: str
    owner: str
    mapping_status: str
    notes: str


def _display(value: object, *, required: bool = True) -> str:
    if value is None or not str(value).strip():
        return INPUT_REQUIRED if required else ""
    text = str(value).strip()
    if text.startswith("="):
        return INPUT_REQUIRED if required else ""
    return text


def _workbook_value(formula_sheet, value_sheet, row: int, column: int) -> object:
    cached = value_sheet.cell(row, column).value
    if cached is not None:
        return cached
    return formula_sheet.cell(row, column).value


def _read_workbook(workbook_path: Path) -> tuple[dict[str, str], list[SSPControlRow]]:
    formulas = load_workbook(workbook_path, data_only=False, read_only=True)
    values = load_workbook(workbook_path, data_only=True, read_only=True)
    try:
        cover_f = formulas["Cover"]
        cover_v = values["Cover"]
        cover = {
            str(cover_f.cell(row, 2).value).strip(): _display(
                _workbook_value(cover_f, cover_v, row, 3), required=False
            )
            for row in range(1, cover_f.max_row + 1)
            if cover_f.cell(row, 2).value
        }

        assessment_f = formulas["Assessment"]
        assessment_v = values["Assessment"]
        assessment: dict[str, list[object]] = {}
        for row in range(6, assessment_f.max_row + 1):
            requirement_id = _display(assessment_f.cell(row, 2).value, required=False)
            if requirement_id:
                assessment[requirement_id] = [
                    _workbook_value(assessment_f, assessment_v, row, col)
                    for col in range(1, assessment_f.max_column + 1)
                ]

        crosswalk_f = formulas["SSP Crosswalk"]
        crosswalk_v = values["SSP Crosswalk"]
        rows: list[SSPControlRow] = []
        for row in range(6, crosswalk_f.max_row + 1):
            requirement_id = _display(crosswalk_f.cell(row, 2).value, required=False)
            if not requirement_id:
                continue
            source = assessment.get(requirement_id, [])
            rows.append(
                SSPControlRow(
                    domain=_display(crosswalk_f.cell(row, 1).value, required=False),
                    requirement_id=requirement_id,
                    title=_display(source[2] if len(source) > 2 else None),
                    statement=_display(source[3] if len(source) > 3 else None),
                    ssp_reference=_display(
                        _workbook_value(crosswalk_f, crosswalk_v, row, 3)
                    ),
                    governance_references=_display(
                        _workbook_value(crosswalk_f, crosswalk_v, row, 4)
                    ),
                    evidence_references=_display(
                        _workbook_value(crosswalk_f, crosswalk_v, row, 5)
                    ),
                    owner=_display(_workbook_value(crosswalk_f, crosswalk_v, row, 6)),
                    mapping_status=_display(
                        _workbook_value(crosswalk_f, crosswalk_v, row, 7)
                    ),
                    notes=_display(_workbook_value(crosswalk_f, crosswalk_v, row, 8)),
                )
            )
        return cover, rows
    finally:
        formulas.close()
        values.close()


def _iter_paragraphs(document: DocumentType) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _replace_literal(document: DocumentType, old: str, new: str) -> int:
    replacements = 0
    for paragraph in _iter_paragraphs(document):
        for run in paragraph.runs:
            if old in run.text:
                replacements += run.text.count(old)
                run.text = run.text.replace(old, new)
    return replacements


def _set_grid_span(cell_xml, span: int) -> None:
    properties = cell_xml.get_or_add_tcPr()
    grid_span = properties.find(qn("w:gridSpan"))
    if grid_span is None:
        grid_span = OxmlElement("w:gridSpan")
        properties.append(grid_span)
    grid_span.set(qn("w:val"), str(span))


def _add_blank_sprs_score_fields(document: DocumentType) -> set[str]:
    """Insert a blank SPRS Score between CMMC Level and Practice ID."""

    requirement_pattern = re.compile(r"[A-Z]{2}\.L2-3\.\d+\.\d+")
    updated: set[str] = set()
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        header = table.rows[1]
        header_text = "\n".join(cell.text for cell in header.cells)
        match = requirement_pattern.search(header_text)
        if (
            not match
            or "CMMC Level:" not in header_text
            or "Practice ID:" not in header_text
        ):
            continue

        practice_id = match.group(0)
        practice_name = header.cells[6].text
        physical_cells = list(header._tr.tc_lst)
        if len(physical_cells) != 6:
            raise ValueError(f"Unexpected control-header geometry for {practice_id}.")

        practice_row = deepcopy(header._tr)
        for cell in list(practice_row.tc_lst):
            practice_row.remove(cell)
        label_cell = deepcopy(physical_cells[4])
        value_cell = deepcopy(physical_cells[5])
        _set_grid_span(label_cell, 2)
        _set_grid_span(value_cell, 5)
        practice_row.append(label_cell)
        practice_row.append(value_cell)
        header._tr.addnext(practice_row)

        _write_artifact_value(header.cells[2], "SPRS Score:")
        _write_artifact_value(header.cells[3], "")
        _write_artifact_value(header.cells[4], "Practice ID:")
        _write_artifact_value(header.cells[6], practice_id)
        inserted_row = table.rows[2]
        _write_artifact_value(inserted_row.cells[0], "Practice Name:")
        _write_artifact_value(inserted_row.cells[2], practice_name)
        updated.add(practice_id)
    return updated


def _write_artifact_value(cell, value: str) -> None:
    """Replace template guidance in an artifact cell while preserving its style."""

    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)
    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def _populate_supporting_artifacts(
    document: DocumentType, rows: list[SSPControlRow]
) -> set[str]:
    """Populate each practice's existing Supporting Artifacts table."""

    controls = {row.requirement_id: row for row in rows}
    populated: set[str] = set()
    tables = document.tables
    requirement_pattern = re.compile(r"[A-Z]{2}\.L2-3\.\d+\.\d+")

    for index, table in enumerate(tables):
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        match = requirement_pattern.search(table_text)
        if not match or match.group(0) not in controls:
            continue

        artifact_table = None
        for candidate in tables[index + 1 :]:
            candidate_text = "\n".join(
                cell.text for row in candidate.rows for cell in row.cells
            )
            if requirement_pattern.search(candidate_text):
                break
            if "Supporting Artifacts" in candidate_text:
                artifact_table = candidate
                break
        if artifact_table is None:
            continue

        control = controls[match.group(0)]
        artifact_values = {
            "System Design Documentation": control.governance_references,
            "System Configuration Settings And Associated Documentation": INPUT_REQUIRED,
            "Supplemental Artifacts": control.evidence_references,
        }
        for row_index, artifact_row in enumerate(artifact_table.rows[:-1]):
            label = artifact_row.cells[0].text.strip()
            if label in artifact_values:
                _write_artifact_value(
                    artifact_table.rows[row_index + 1].cells[0], artifact_values[label]
                )
        populated.add(control.requirement_id)

    return populated


def export_ssp(
    template_path: str | Path,
    workbook_path: str | Path,
    output_path: str | Path,
    metadata: SSPExportMetadata | None = None,
) -> Path:
    """Create a Word SSP from an immutable template and an Omni workbook."""

    template = Path(template_path).resolve()
    workbook = Path(workbook_path).resolve()
    output = Path(output_path).resolve()
    if template == output:
        raise ValueError("Output path must differ from the source template path.")
    if not template.is_file():
        raise FileNotFoundError(f"SSP template not found: {template}")
    if not workbook.is_file():
        raise FileNotFoundError(f"Omni workbook not found: {workbook}")

    cover, rows = _read_workbook(workbook)
    if not rows:
        raise ValueError("The Omni workbook contains no SSP Crosswalk records.")

    supplied = metadata or SSPExportMetadata()
    organization = supplied.organization_name or cover.get("Organization Name", "")
    resolved = SSPExportMetadata(
        organization_name=_display(organization),
        system_name=supplied.system_name,
        system_owner=supplied.system_owner,
        prepared_by=supplied.prepared_by,
        version=supplied.version,
        export_date=supplied.export_date,
    )

    document = Document(template)
    _replace_literal(document, "ACME", resolved.organization_name)
    updated_headers = _add_blank_sprs_score_fields(document)
    expected_headers = {row.requirement_id for row in rows}
    if updated_headers != expected_headers:
        missing = sorted(expected_headers - updated_headers)
        raise ValueError(
            "The SSP template is missing control headers for: " + ", ".join(missing)
        )
    document.core_properties.title = (
        f"{resolved.organization_name} System Security Plan"
    )
    document.core_properties.subject = "Omni Security Plan (SSP) export"
    document.core_properties.comments = "Generated by Omni by R!SC"

    populated = _populate_supporting_artifacts(document, rows)
    if len(populated) != len(rows):
        missing = sorted(
            row.requirement_id for row in rows if row.requirement_id not in populated
        )
        raise ValueError(
            "The SSP template is missing Supporting Artifacts tables for: "
            + ", ".join(missing)
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    file_descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{output.stem}-", suffix=".docx", dir=output.parent
    )
    os.close(file_descriptor)
    temporary = Path(temporary_name)
    try:
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)
    return output
