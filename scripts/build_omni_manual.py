"""Build the public, framework-agnostic Omni user and administrator manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Omni-Comprehensive-User-and-Administrator-Manual.docx"
ASSET_DIR = ROOT / "output" / "manual_assets"
WORKFLOW_IMAGE = ASSET_DIR / "omni_workflow.png"
SCREENSHOT_DIR = ROOT / "docs" / "manual_screenshots"
ANNOTATED_SCREENSHOT_DIR = ASSET_DIR / "annotated_screenshots"

SCREENSHOTS = [
    ("01-sign-in.png", "Sign in", [(1, .50, .35), (2, .50, .53)], ["Enter your Omni account credentials.", "Select Sign in to open your authorized workspace."]),
    ("02-organizations.png", "Choose an organization", [(1, .20, .28), (2, .83, .20)], ["Open the organization that owns the engagement.", "Administrators can create a new organization from this action."]),
    ("03-systems.png", "Choose a system or environment", [(1, .22, .39), (2, .82, .20)], ["Open the system, enclave, or environment in assessment scope.", "Use Add system to establish a new assessment boundary."]),
    ("04-assessments.png", "Open or create an assessment", [(1, .28, .40), (2, .83, .20)], ["Open the applicable assessment period and purpose.", "Use New assessment to select one or more frameworks and load controls."]),
    ("05-dashboard.png", "Use the executive dashboard", [(1, .48, .18), (2, .28, .43), (3, .83, .20)], ["Use the action bar to move between planning, execution, evidence, remediation, review, and reports.", "Monitor completion, results, evidence, remediation, deadlines, and optional risk indicators.", "Use dashboard actions to maintain authorized assessment settings."]),
    ("06-assessment-plan.png", "Complete the assessment plan", [(1, .25, .31), (2, .73, .31), (3, .24, .83)], ["Define dates, boundaries, locations, and methodology.", "Assign the assessment team and team roles.", "Save the plan after confirming optional risk settings."]),
    ("07-assessment-execution.png", "Execute assessment objectives", [(1, .53, .24), (2, .26, .44), (3, .88, .67)], ["Record interviews, samples, and tests when the methodology requires them.", "Use the status metrics and filters to identify remaining objective work.", "Select Assess to document the objective status, notes, and supporting artifacts."]),
    ("08-control-result.png", "Record the control conclusion", [(1, .28, .31), (2, .27, .61), (3, .79, .84)], ["Select the supported control status and implementation state.", "Use Assessor Notes/Findings for either a conformity statement or a finding.", "Confirm ownership, dates, and evidence, then save and refresh the dashboard."]),
    ("09-evidence-workspace.png", "Manage evidence", [(1, .54, .21), (2, .25, .47), (3, .87, .20)], ["Filter and track evidence requests through the collection workflow.", "Open a request to manage ownership, controls, due dates, and artifacts.", "Register a file or governed external reference as a supporting artifact."]),
    ("10-remediation.png", "Track remediation", [(1, .26, .41), (2, .65, .41), (3, .88, .20)], ["Open a remediation plan to manage corrective action and linked controls.", "Monitor priority, status, due dates, and validation.", "Create a new plan for a confirmed assessment gap."]),
    ("11-quality-review.png", "Complete quality review", [(1, .25, .35), (2, .73, .35), (3, .79, .82)], ["Resolve assessment readiness blockers before approval.", "Review scoring, evidence, findings, demographics, remediation, and report warnings.", "Record the review decision and notes before sign-off."]),
    ("12-report-center.png", "Generate deliverables", [(1, .28, .35), (2, .73, .35), (3, .50, .78)], ["Review blockers and warnings before generating final deliverables.", "Select the workbook, Word plan, framework, consolidated, traceability, or package output.", "Use document history to confirm version, generator, date, and size."]),
    ("13-risk-register.png", "Use optional risk management", [(1, .30, .35), (2, .70, .35), (3, .86, .20)], ["Review the risk matrix and registered inherent and residual evaluations.", "Open risks to manage treatment, reassessment, acceptance, and closure.", "Add an independent risk only when risk management is enabled for the engagement."]),
    ("14-framework-catalog.png", "Administer frameworks", [(1, .28, .37), (2, .70, .37), (3, .84, .20)], ["Open a framework to review its governed requirements and versions.", "Use catalog tools for imports, mappings, evidence, sources, and risk relationships.", "Start a governed Excel, CSV, or PDF framework import."]),
    ("15-notification-policy.png", "Configure notification policy", [(1, .28, .36), (2, .70, .36), (3, .27, .80)], ["Use the master switch to enable or disable organization email automation.", "Choose reminder frequency and escalation recipients.", "Send a test only after protected SMTP configuration is complete."]),
    ("16-members-access.png", "Manage members and access", [(1, .27, .35), (2, .72, .35), (3, .85, .20)], ["Review active memberships, roles, and account activity.", "Invite users and assign the least-privileged suitable role.", "Export the access review for periodic governance."]),
]
WALKTHROUGH_SCREENSHOTS = [SCREENSHOTS[index] for index in (0, 3, 4, 5, 6, 7, 8, 11)]

NAVY = "15324B"
BLUE = "2E74B5"
TEAL = "187A78"
GREEN = "2E7D5B"
GOLD = "B78324"
LIGHT_BLUE = "E8F1F7"
LIGHT_TEAL = "E8F5F3"
LIGHT_GOLD = "FFF6DF"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "1F2933"
CONTENT_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        keep_table_row_together(row)
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, size=None, color=BLACK, bold=None, italic=None, font="Aptos") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run(run, size=9, color=MID_GRAY)


def add_inline_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run(run, size=8, color=MID_GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.17

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, TEAL, 0, 10),
        ("Heading 1", 17, NAVY, 14, 8),
        ("Heading 2", 13.5, BLUE, 11, 6),
        ("Heading 3", 11.5, TEAL, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name.startswith("Heading") or name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.17

    for style_name, fill, color in (
        ("Omni Note", LIGHT_BLUE, NAVY),
        ("Omni Tip", LIGHT_TEAL, GREEN),
        ("Omni Warning", LIGHT_GOLD, GOLD),
    ):
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.color.rgb = rgb(color)
        style.paragraph_format.left_indent = Inches(0.14)
        style.paragraph_format.right_indent = Inches(0.14)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(8)
        p_pr = style._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "16")
        left.set(qn("w:color"), color)
        borders.append(left)
        p_pr.append(borders)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("OMNI  |  USER AND ADMINISTRATOR MANUAL")
    set_run(r, size=8.5, color=MID_GRAY, bold=True)
    footer = section.footer
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Omni by R!SC  |  Version 1.1  |  August 2026  |  Page ")
    set_run(r1, size=8, color=MID_GRAY)
    add_inline_page_number(p1)


def add_callout(doc, label: str, text: str, kind="note") -> None:
    style = {"note": "Omni Note", "tip": "Omni Tip", "warning": "Omni Warning"}[kind]
    p = doc.add_paragraph(style=style)
    r = p.add_run(f"{label}: ")
    set_run(r, bold=True, color={"note": NAVY, "tip": GREEN, "warning": GOLD}[kind])
    r = p.add_run(text)
    set_run(r, color=BLACK)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def new_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    # OOXML requires every abstract numbering definition to precede concrete
    # w:num instances. Inserting after existing w:num nodes can corrupt Word's
    # built-in bullet and numbering styles.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_steps(doc, items: list[str]) -> None:
    num_id = new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_node)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9, color=WHITE, bold=True)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        if i % 2:
            for cell in cells:
                shade(cell, LIGHT_GRAY)
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run(r, size=9, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def create_workflow_image() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 46)
        box_font = ImageFont.truetype("arialbd.ttf", 27)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()
    draw.text((90, 45), "The Omni Assessment Lifecycle", fill="#15324B", font=title_font)
    stages = [
        ("1", "Establish context", "Organization, system, scope"),
        ("2", "Plan", "Frameworks, team, dates"),
        ("3", "Collect", "Requests and artifacts"),
        ("4", "Assess", "Objectives and controls"),
        ("5", "Respond", "Findings and remediation"),
        ("6", "Assure", "Quality review and sign-off"),
        ("7", "Report", "Framework and consolidated outputs"),
        ("8", "Monitor", "Optional risk and follow-up"),
    ]
    colors = ["#E8F1F7", "#E8F5F3", "#FFF6DF", "#E8F1F7", "#FCECEC", "#E8F5F3", "#E8F1F7", "#FFF6DF"]
    positions = [(90 + (i % 4) * 425, 170 + (i // 4) * 340) for i in range(8)]
    for i, ((num, title, subtitle), (x, y), color) in enumerate(zip(stages, positions, colors)):
        draw.rounded_rectangle((x, y, x + 335, y + 190), radius=22, fill=color, outline="#2E74B5", width=4)
        draw.ellipse((x + 20, y + 22, x + 78, y + 80), fill="#15324B")
        draw.text((x + 40, y + 29), num, fill="white", font=small_font, anchor="ma")
        draw.text((x + 20, y + 98), title, fill="#15324B", font=box_font)
        draw.multiline_text((x + 20, y + 140), subtitle, fill="#475467", font=small_font, spacing=4)
        if i % 4 != 3:
            draw.line((x + 342, y + 95, x + 405, y + 95), fill="#187A78", width=8)
            draw.polygon([(x + 405, y + 95), (x + 383, y + 82), (x + 383, y + 108)], fill="#187A78")
    draw.line((1595, 370, 1595, 500), fill="#187A78", width=8)
    draw.polygon([(1595, 500), (1582, 478), (1608, 478)], fill="#187A78")
    draw.text((90, 900), "Assessment conclusions remain framework-specific. Evidence and testing may be reused only through reviewed mappings and explicit applicability decisions.", fill="#475467", font=small_font)
    image.save(WORKFLOW_IMAGE, dpi=(180, 180))


def create_annotated_screenshots() -> None:
    """Create public manual figures from the committed synthetic screenshots."""
    ANNOTATED_SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    try:
        badge_font = ImageFont.truetype("arialbd.ttf", 34)
    except OSError:
        badge_font = ImageFont.load_default()
    for filename, _title, callouts, _legend in SCREENSHOTS:
        source = SCREENSHOT_DIR / filename
        if not source.exists():
            raise FileNotFoundError(f"Missing manual screenshot: {source}")
        with Image.open(source).convert("RGB") as screenshot:
            screenshot.thumbnail((1200, 760), Image.Resampling.LANCZOS)
            draw = ImageDraw.Draw(screenshot)
            radius = max(24, round(screenshot.width * .020))
            for number, x_ratio, y_ratio in callouts:
                x, y = round(screenshot.width * x_ratio), round(screenshot.height * y_ratio)
                draw.ellipse((x-radius, y-radius, x+radius, y+radius), fill="#15324B", outline="white", width=5)
                draw.text((x, y+1), str(number), fill="white", font=badge_font, anchor="mm")
            screenshot.save(ANNOTATED_SCREENSHOT_DIR / filename, dpi=(144, 144), optimize=True)


def add_screenshot_figure(doc: Document, filename: str, title: str, legend: list[str]) -> None:
    doc.add_page_break()
    doc.add_heading(title, 2)
    doc.add_picture(str(ANNOTATED_SCREENSHOT_DIR / filename), width=Inches(6.65))
    caption = doc.add_paragraph(f"Illustration: {title}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.runs[0], size=9, color=MID_GRAY, italic=True)
    for number, text in enumerate(legend, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(.18)
        p.paragraph_format.first_line_indent = Inches(-.18)
        badge = p.add_run(f"{number}  ")
        set_run(badge, size=10, color=NAVY, bold=True)
        set_run(p.add_run(text), size=10, color=BLACK)


def cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OMNI")
    set_run(r, size=38, color=NAVY, bold=True, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("by R!SC")
    set_run(r, size=13, color=TEAL, bold=True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Comprehensive User and\nAdministrator Manual")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Framework-agnostic guidance for planning, executing, governing, and reporting assessments")
    doc.add_paragraph()
    add_table(doc, ["Document", "Value"], [
        ["Version", "1.1"],
        ["Release date", "August 14, 2026"],
        ["Applies to", "Omni local web application through Sprint 18.3.1"],
        ["Audience", "Assessors, control owners, organization administrators, reviewers, and platform administrators"],
        ["Classification", "Public product documentation - contains no client or organization data"],
    ], [2500, 6860])
    add_callout(doc, "Purpose", "Use this manual to operate Omni end to end. It explains the application workflow; it does not replace the requirements, assessment methodology, or professional judgment associated with any selected framework.", "note")
    doc.add_page_break()


def front_matter(doc: Document) -> None:
    doc.add_heading("Revision history", 1)
    add_table(doc, ["Version", "Date", "Change", "Owner"], [
        ["1.1", "August 14, 2026", "Added synthetic, annotated application screenshots for Sprint 18.3.1", "R!SC"],
        ["1.0", "August 14, 2026", "Initial comprehensive user and administrator manual for Sprint 18.3", "R!SC"],
    ], [1100, 1900, 4860, 1500])
    doc.add_heading("How to use this manual", 1)
    add_bullets(doc, [
        "New assessors should read Sections 1 through 8 in sequence and keep Appendix A available as a quick reference.",
        "Control owners and client contributors should begin with Section 2.3, then use Sections 6 and 7 for evidence and remediation responsibilities.",
        "Organization administrators should use Sections 3, 4, 12, and 14 for access, configuration, notifications, and operational safeguards.",
        "Framework administrators should use Sections 9 and 13 for imports, mappings, catalogs, and governance.",
        "Risk management is optional. Skip Section 10 when the assessment plan disables it.",
    ])
    add_callout(doc, "Terminology", "Omni uses universal terms such as framework requirement, assessment objective, supporting artifact, finding, remediation plan, and security plan. Framework-specific aliases and scoring methods may appear in generated deliverables only when required by the selected framework.", "tip")
    doc.add_heading("Contents", 1)
    toc = [
        "1. Introduction and assessment lifecycle", "2. Roles and getting started", "3. Organizations, systems, and profiles",
        "4. Create and plan an assessment", "5. Execute and record assessment results", "6. Evidence requests and supporting artifacts",
        "7. Findings and remediation", "8. Quality review, sign-off, and reopening", "9. Multi-framework assessment and reuse",
        "10. Optional risk management", "11. Dashboards, analytics, and reports", "12. Notifications and workflow automation",
        "13. Framework and mapping administration", "14. Security and platform administration", "15. Troubleshooting",
        "Appendix A. Quick-reference workflows", "Appendix B. Status and terminology reference", "Appendix C. Administrator command reference",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()


def section_1(doc: Document) -> None:
    doc.add_heading("1. Introduction and assessment lifecycle", 1)
    doc.add_paragraph("Omni is a governance, risk, and compliance assessment platform. It organizes client environments, selected frameworks, assessment objectives, evidence, findings, remediation, reporting, and optional operational risk into one governed workflow.")
    doc.add_picture(str(WORKFLOW_IMAGE), width=Inches(6.65))
    p = doc.add_paragraph("Figure 1. The Omni assessment lifecycle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = rgb(MID_GRAY)
    doc.add_heading("1.1 Core operating principles", 2)
    add_bullets(doc, [
        "Framework conclusions remain independent. A result recorded for one framework is not silently copied into another.",
        "Evidence can support many requirements, but applicability, scope, period, and freshness must be reviewed.",
        "Mappings accelerate analysis; they do not replace assessor judgment.",
        "Control-level results summarize the conclusion. Objective-level records preserve the examination, interview, test, sampling, notes, and evidence behind it.",
        "Risk management is optional and does not block completion of the underlying assessment when disabled.",
        "Sign-off locks the assessment to protect the approved record. Reopening requires justification and is audited.",
    ])
    doc.add_heading("1.2 Framework-native behavior", 2)
    doc.add_paragraph("Omni preserves each framework's native identifiers, domains, assessment objectives, scoring method, and reporting context. A weighted score used by one framework is not presented as a universal GRC score. The dashboard identifies the active framework and labels its score accordingly.")
    add_callout(doc, "Example", "A government framework may use its own weighted scoring method, while another framework may use completion, maturity, or no numerical score. Omni displays the configured method without redefining it.", "note")


def section_2(doc: Document) -> None:
    doc.add_heading("2. Roles and getting started", 1)
    doc.add_heading("2.1 Sign in and navigate", 2)
    add_steps(doc, [
        "Start Omni with Run Omni Web.cmd and keep the command window open.",
        "Open http://127.0.0.1:8000/ in a browser and sign in.",
        "Choose an organization, then a system or environment, then an assessment.",
        "Use the assessment dashboard action bar for Plan, Execute, Harmonize, Quality review, Evidence, Remediation, Reports, and—when authorized—Access, Frameworks, and Assign owners.",
        "Use the top bar for Framework catalog, Notifications, profile settings, and sign out.",
    ])
    doc.add_heading("2.2 Role matrix", 2)
    add_table(doc, ["Role", "Typical responsibilities", "Key boundaries"], [
        ["Organization administrator", "Organization profile, systems, members, invitations, access review, notification policy, health, risk approvals", "Limited to assigned organizations"],
        ["Lead assessor / assessor", "Plan, execute, document conclusions, review evidence, create remediation and risk records, quality review and reports", "Assessment access and lock state apply"],
        ["Control owner / contributor", "Respond to requests, upload artifacts, support assigned controls and actions", "Cannot administer organization or approve final assessment"],
        ["Reviewer / viewer", "Read authorized assessment information and generated outputs", "Read-only unless explicitly granted another role"],
        ["Platform administrator", "Framework imports, mapping governance, source catalogs, backups, readiness and deployment configuration", "Must protect private source files and secrets"],
    ], [1900, 4710, 2750])
    doc.add_heading("2.3 Quick start by task", 2)
    add_table(doc, ["I need to...", "Go to...", "Action"], [
        ["Enter an overall control result", "Assessment dashboard > Control results", "Select Control result; enter status, implementation, notes/findings, owners, and save"],
        ["Record objective-level work", "Assessment dashboard > Execute", "Filter, select Assess, record status, notes, and evidence"],
        ["Upload evidence", "Assessment dashboard > Evidence", "Register artifact; upload a file or enter an external reference; link controls"],
        ["Create a finding response", "Control result or Remediation", "Set the control to Not Met, then create and track a remediation plan"],
        ["Generate deliverables", "Assessment dashboard > Reports", "Resolve readiness blockers, then select the required output"],
    ], [2300, 2900, 4160])
    doc.add_heading("2.4 Illustrated end-to-end walkthrough", 2)
    add_callout(doc, "Synthetic data", "Every screenshot in this manual was captured from an isolated demonstration database containing fictional names, accounts, systems, controls, evidence, remediation, and risk records. No client or organization data is shown.", "note")
    doc.add_paragraph("Follow the numbered markers in each illustration and the matching instructions below it. Screen contents vary by role, assessment state, selected frameworks, and whether optional risk management is enabled.")
    for filename, title, _callouts, legend in WALKTHROUGH_SCREENSHOTS:
        add_screenshot_figure(doc, filename, title, legend)


def section_3(doc: Document) -> None:
    doc.add_heading("3. Organizations, systems, and profiles", 1)
    doc.add_heading("3.1 Create or maintain an organization", 2)
    add_steps(doc, [
        "From Organizations, select New organization, or open an existing organization and select Profile.",
        "Record the legal or operating name, address, and available contact information.",
        "Save and confirm that the organization appears only to authorized members.",
    ])
    doc.add_heading("3.2 Add a system, enclave, or environment", 2)
    add_steps(doc, [
        "Open the organization and select Add system.",
        "Enter the system name, scope and boundaries, identifiers, owner, security officer, and other available demographics.",
        "Save. The system becomes the container for one or more assessments.",
    ])
    add_callout(doc, "Data quality", "System demographics feed report exports. Complete them before final reporting even when they are not required to create the assessment.", "tip")
    doc.add_heading("3.3 Maintain your profile", 2)
    doc.add_paragraph("Select your name in the top bar to update name, email, job title, phone, and time zone or to change your password. Notification delivery depends on a valid email address.")


def section_4(doc: Document) -> None:
    doc.add_heading("4. Create and plan an assessment", 1)
    doc.add_heading("4.1 Create the assessment", 2)
    add_steps(doc, [
        "Open the applicable system and select New assessment.",
        "Enter an engagement name that distinguishes the assessment period and purpose.",
        "Select one or more frameworks and identify the primary framework.",
        "Select Create and load controls. Omni creates separate result records for the selected frameworks.",
    ])
    doc.add_heading("4.2 Select and change frameworks", 2)
    doc.add_paragraph("Use Frameworks from the assessment dashboard. The primary framework controls the default scoring and reporting context. A secondary framework may be removed only while its results remain untouched; this protects completed work from accidental deletion.")
    doc.add_heading("4.3 Complete the assessment plan", 2)
    add_bullets(doc, [
        "Engagement dates and due date",
        "Scope boundaries, locations, sampling methodology, and assessment methodology",
        "Assessment team and team roles",
        "Optional risk management and risk reporting settings",
    ])
    add_callout(doc, "Optional risk", "Enable risk management only when the engagement includes organizational risk evaluation. Risk reporting can be independently disabled so risk records remain internal while assessment reports omit them.", "note")
    doc.add_heading("4.4 Assign access and control owners", 2)
    add_steps(doc, [
        "Administrators use Access to grant assessment-specific authorization when tighter access than organization membership is needed.",
        "Use Assign owners to allocate a primary owner and supporting owners by domain.",
        "Review the dashboard's control-owner workload to confirm coverage and balance.",
    ])


def section_5(doc: Document) -> None:
    doc.add_heading("5. Execute and record assessment results", 1)
    doc.add_heading("5.1 Understand the two result levels", 2)
    add_table(doc, ["Level", "Use", "Record"], [
        ["Assessment objective", "Detailed examination, interview, and test work", "Objective status, assessor notes, evidence, assessor identity; related interview, sample, and test records"],
        ["Framework requirement / control", "Final summarized conclusion and reporting record", "Status, implementation state, Assessor Notes/Findings, ownership, dates, score impact, and linked artifacts"],
    ], [2000, 3300, 4060])
    doc.add_heading("5.2 Record objective-level assessment work", 2)
    add_steps(doc, [
        "From the assessment dashboard, select Execute.",
        "Filter by framework or assessment method if needed.",
        "Select Assess beside an objective.",
        "Set the objective status and enter notes that state what was examined, interviewed, or tested and the resulting observation.",
        "Select the supporting artifacts and save.",
        "Use Schedule interview, Define sample, and Record test to preserve procedural details when applicable.",
    ])
    doc.add_heading("5.3 Record the control-level result", 2)
    add_steps(doc, [
        "Return to the dashboard and scroll to Control results.",
        "Locate the framework requirement and select Control result.",
        "Choose the assessment status and implementation state supported by the completed objectives.",
        "Use Assessor Notes/Findings for the conclusion: write a conformity statement when requirements are satisfied and a finding when they are not.",
        "Confirm the control owner, dates, and linked supporting artifacts, then select Save and update dashboard.",
    ])
    add_callout(doc, "Conformity statement", "State what was implemented, the scope and period evaluated, the evidence reviewed, and why the evidence supports the conclusion.", "tip")
    add_callout(doc, "Finding", "State the requirement, observed condition, objective evidence, scope, and impact. Avoid prescribing a solution unless the engagement calls for recommendations.", "warning")
    doc.add_heading("5.4 Status discipline", 2)
    add_bullets(doc, [
        "Met: evidence supports satisfaction of the requirement and its applicable objectives.",
        "Not Met: one or more applicable objectives are not satisfied or evidence demonstrates a gap.",
        "Not Applicable: justified exclusion allowed by the framework and assessment scope.",
        "Not Assessed: work is incomplete or no conclusion has been reached.",
        "Partially implemented is an implementation descriptor; do not substitute it for the framework's required assessment conclusion.",
    ])


def section_6(doc: Document) -> None:
    doc.add_heading("6. Evidence requests and supporting artifacts", 1)
    doc.add_heading("6.1 Generate and manage the request list", 2)
    add_steps(doc, [
        "Open Evidence and select Generate request list to create optimized requests from the selected frameworks.",
        "Review generated requests, consolidate duplicates where appropriate, assign owners and due dates, and add custom requests when necessary.",
        "Track each request through its workflow status and use notification automation only when enabled by policy.",
    ])
    doc.add_heading("6.2 Register an artifact", 2)
    add_steps(doc, [
        "Select Register artifact.",
        "Enter a clear title, source, collection and expiration information, and the related request.",
        "Upload the file or enter an external reference. A package export converts external links to text records so the reference is retained.",
        "Link every control the artifact supports and save.",
        "An authorized reviewer records acceptance, rejection, or further review information.",
    ])
    add_callout(doc, "Security", "Evidence files are organization-scoped and assessment-authorized. Never place client evidence, exported packages, local databases, credentials, or private source workbooks in the public Git repository.", "warning")
    doc.add_heading("6.3 Reuse evidence responsibly", 2)
    doc.add_paragraph("In the Shared evidence and testing workspace, record whether an artifact is fully applicable, partially applicable, not applicable, or requires additional evidence. Document period, scope, and system limitations. Freshness is a signal, not an automatic acceptance decision.")


def section_7(doc: Document) -> None:
    doc.add_heading("7. Findings and remediation", 1)
    doc.add_heading("7.1 Create a remediation plan", 2)
    add_steps(doc, [
        "Open a Not Met control and select Create remediation plan, or open Remediation and select New plan.",
        "Enter the remediation identifier, title, description, root cause, planned corrective action, owner, priority, dates, resources, status, and validation information.",
        "Link all related controls and supporting evidence.",
        "Add milestones with owners, due dates, status, and optional owner notifications.",
        "Update progress until corrective action and validation are complete.",
    ])
    doc.add_heading("7.2 Remediation terminology", 2)
    doc.add_paragraph("Omni's canonical object is a Remediation Action Plan. Framework-specific deliverables may use aliases such as corrective action plan or plan of action and milestones. The underlying object remains reusable across frameworks.")
    doc.add_heading("7.3 Export remediation", 2)
    doc.add_paragraph("Use Remediation > Export or the Report center to generate the workbook-compatible remediation deliverable. Review open, overdue, and high-priority items before release.")


def section_8(doc: Document) -> None:
    doc.add_heading("8. Quality review, sign-off, and reopening", 1)
    doc.add_heading("8.1 Perform quality review", 2)
    add_steps(doc, [
        "Open Quality review from the assessment dashboard.",
        "Resolve readiness blockers, including unassessed objectives and incomplete conclusions.",
        "Review framework scores, evidence acceptance, findings, remediation, demographics, and report warnings.",
        "Set the quality review status and record review notes.",
    ])
    doc.add_heading("8.2 Sign off and lock", 2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_paragraph("Final sign-off requires an approved quality review and assessment of every objective. Sign-off locks assessment edits and records the signer and timestamp.")
    doc.add_heading("8.3 Reopen with justification", 2)
    doc.add_paragraph("When correction is necessary after sign-off, use Reopen, enter a specific justification, make the controlled change, repeat quality review, and sign off again. The workflow history preserves the event.")


def section_9(doc: Document) -> None:
    doc.add_heading("9. Multi-framework assessment and reuse", 1)
    doc.add_heading("9.1 Harmonization workflow", 2)
    add_steps(doc, [
        "Select two or more frameworks for the assessment.",
        "Open Harmonize and select Analyze mappings.",
        "Review each suggested relationship and mapping path.",
        "Choose whether to reuse accepted evidence and/or recognize prior testing by reference.",
        "Enter reviewer rationale and approve or reject the reuse decision.",
        "Use Shared work to record applicability and identify framework-specific remaining work.",
    ])
    add_callout(doc, "Critical rule", "Omni never propagates compliance outcomes automatically. Each framework retains a separate conclusion and deliverable even when evidence or testing is reused.", "warning")
    doc.add_heading("9.2 Consolidated evidence requests", 2)
    doc.add_paragraph("The Shared work workspace can consolidate selected requests that ask for the same underlying artifact. Consolidation reduces burden without obscuring which framework requirements the request supports.")


def section_10(doc: Document) -> None:
    doc.add_heading("10. Optional risk management", 1)
    add_callout(doc, "Optional capability", "The assessor can disable risk management during planning. When disabled, the underlying assessment, findings, remediation, quality review, and framework reporting continue normally.", "note")
    doc.add_page_break()
    doc.add_heading("10.1 Interpret the two heatmaps", 2)
    add_table(doc, ["View", "Purpose", "Calculation"], [
        ["Weighted control exposure", "Prioritize gaps using the selected framework's configured control weights", "Not Met assessed weight divided by assessed applicable weight; unassessed weight remains unknown"],
        ["Organizational risk matrix", "Evaluate registered business or operational risks", "Likelihood (1-5) multiplied by impact (1-5), with separate inherent and residual evaluations"],
    ], [2200, 3400, 3760])
    doc.add_heading("10.2 Convert a finding to a risk", 2)
    add_steps(doc, [
        "Open the Organizational Risk Register and review approved catalog suggestions associated with current findings.",
        "Select Evaluate and register, or Add risk for an independently identified risk.",
        "Write the risk statement, identify source and affected controls, assign an owner, and evaluate inherent likelihood and impact.",
        "Select a treatment response, target date, review frequency, and treatment plan.",
    ])
    doc.add_heading("10.3 Treat, reassess, accept, and close", 2)
    add_steps(doc, [
        "Add treatment actions and track their owners, dates, priority, and status.",
        "Reassess residual likelihood and impact with rationale and evidence.",
        "When appropriate, request risk acceptance. An authorized administrator approves or rejects it and controls expiration.",
        "Close only after treatment actions, residual evaluation, and supporting evidence satisfy the closure gate.",
        "Reopen with justification when conditions materially change.",
    ])


def section_11(doc: Document) -> None:
    doc.add_heading("11. Dashboards, analytics, and reports", 1)
    doc.add_heading("11.1 Read the executive dashboard", 2)
    add_bullets(doc, [
        "Primary framework score or framework identity, control completion, objective completion, Met/Not Met, and N/A/unassessed counts",
        "Decision readiness and blockers",
        "Evidence acceptance, overdue requests, remediation workload, and deadlines",
        "Per-framework results and domain drill-down",
        "Control-owner workload",
        "Optional weighted exposure and organizational risk matrix",
    ])
    doc.add_heading("11.2 Generate reports", 2)
    add_table(doc, ["Output", "Purpose"], [
        ["Assessment workbook", "Detailed assessment results, evidence register, remediation, scoring, and dashboards"],
        ["Word security plan", "System demographics, assessment summary, control implementation and conclusions, owners, and supporting artifacts"],
        ["Remediation workbook", "Portable remediation action plan"],
        ["Framework report", "Framework-specific result set and native context"],
        ["Consolidated report", "Cross-framework executive view without merging conclusions"],
        ["Traceability CSV", "Mapping and result traceability for analysis"],
        ["Complete ZIP package", "Workbook, Word plan, remediation, evidence index, uploaded files, external-link records, and manifest"],
    ], [2600, 6760])
    doc.add_heading("11.3 Resolve report readiness", 2)
    doc.add_paragraph("The Report center distinguishes blockers from warnings. Completed security-plan and package exports are blocked until required assessment work is complete. Review the generated-document history to confirm who generated each version, when, and its size.")


def section_12(doc: Document) -> None:
    doc.add_heading("12. Notifications and workflow automation", 1)
    doc.add_heading("12.1 Personal preferences", 2)
    doc.add_paragraph("Use Notification preferences to control in-app and email categories available to your account. In-app notifications remain accessible from the top bar; mark individual items or all items as read.")
    doc.add_heading("12.2 Organization notification policy", 2)
    add_bullets(doc, [
        "Master enable/disable switch",
        "Reminder frequency and scheduled processing",
        "Escalation level: owner only; owner and lead assessor; or owner, lead assessor, and the system owner/security officer contact",
        "Email test function for administrators",
    ])
    add_callout(doc, "Email credentials", "Keep SMTP usernames and app passwords only in Omni.local.cmd or protected environment configuration. Omni.local.cmd is ignored by Git and must remain untracked.", "warning")


def section_13(doc: Document) -> None:
    doc.add_heading("13. Framework and mapping administration", 1)
    doc.add_heading("13.1 Framework catalog", 2)
    doc.add_paragraph("The Framework catalog is the entry point for imported frameworks, import jobs, mapping quality, mapping governance, the Omni Evidence Catalog, Authoritative Source Registry, and CCF Risk Catalog.")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_heading("13.2 Import a framework", 2)
    add_steps(doc, [
        "Open Framework catalog > Imports > New import.",
        "Upload Excel, CSV, or PDF. Identify the framework code, name, version, scoring method, and whether it is the Omni Control Framework.",
        "Review the parser preview, column interpretation, issues, and proposed requirements before activation.",
        "Correct source or mapping issues rather than forcing an ambiguous import.",
        "Approve and activate the governed version when validation is complete.",
    ])
    doc.add_heading("13.3 Mapping governance", 2)
    add_bullets(doc, [
        "Review relationship type, confidence, rationale, source provenance, and change impact.",
        "Approve mappings before they can drive reuse, risk suggestions, or reporting traceability.",
        "Use change requests and governance history rather than silently editing approved relationships.",
        "When importing a new framework, use existing Omni Control Framework mappings as a cross-reference accelerator, then validate every relationship.",
    ])
    doc.add_heading("13.4 Catalog boundaries", 2)
    add_table(doc, ["Catalog", "Use", "Not used as"], [
        ["Omni Evidence Catalog", "Canonical evidence objects, aliases, and framework request relationships", "Client evidence"],
        ["Authoritative Source Registry", "Official publications supporting framework definitions, mappings, citations, and versions", "Assessment artifact collection"],
        ["CCF Risk Catalog", "Canonical possible risks and approved control-to-risk relationships", "Automatic organizational risk conclusion"],
    ], [2300, 4360, 2700])


def section_14(doc: Document) -> None:
    doc.add_heading("14. Security and platform administration", 1)
    doc.add_heading("14.1 Membership and access lifecycle", 2)
    add_steps(doc, [
        "Open the organization > Team.",
        "Invite users by email and assign the least-privileged suitable role.",
        "Use assessment-specific Access when only selected members should see an engagement.",
        "Deactivate membership promptly when access is no longer required.",
        "Export the access review periodically and investigate unexpected assignments.",
    ])
    doc.add_heading("14.2 Local data protection", 2)
    add_bullets(doc, [
        "Do not commit omni.sqlite3, Omni.local.cmd, private_uploads, evidence files, client exports, or private source workbooks.",
        "Keep the local application bound to localhost unless a reviewed deployment architecture is implemented.",
        "Use unique user accounts; do not share administrator credentials.",
        "Protect backup archives and evidence with access controls appropriate to their sensitivity.",
    ])
    doc.add_heading("14.3 Health, backup, and readiness", 2)
    add_steps(doc, [
        "Open the organization's System health page and review database, email, invitations, audit activity, and deployment readiness.",
        "Run manage.py backup_omni to create a database and evidence archive with a SHA-256 sidecar.",
        "Run manage.py verify_omni_backup <archive> to verify integrity without restoring or overwriting data.",
        "Run manage.py omni_readiness before any deployment decision.",
    ])
    add_callout(doc, "Deployment boundary", "The local health page does not publish or expose Omni. Internet deployment requires a separate security architecture, production database, TLS, secrets management, monitoring, backup, and recovery implementation.", "warning")


def section_15(doc: Document) -> None:
    doc.add_heading("15. Troubleshooting", 1)
    add_table(doc, ["Symptom", "Likely cause", "Resolution"], [
        ["No organizations are assigned", "Account has no active membership", "Ask an organization administrator to invite or assign the account"],
        ["Cannot edit an assessment", "Viewer role, no assessment grant, or assessment locked", "Review role/access; if signed off, reopen with justification"],
        ["No controls loaded", "Framework selection/import issue", "Review Frameworks, import status, and activate the correct framework version"],
        ["Cannot find where to enter results", "Dashboard drill-down not yet opened", "Use Control results > Control result for the conclusion; Execute > Assess for objectives"],
        ["Report is blocked", "Readiness blockers remain", "Open Report center and Quality review; resolve listed blockers"],
        ["Email is not delivered", "Policy/preferences disabled or SMTP configuration failure", "Review organization policy, personal preferences, System health, and send a test email"],
        ["Risk pages are absent", "Risk management disabled", "Enable it in Assessment plan if the engagement requires risk evaluation"],
        ["Server command window closes", "Startup error", "Run .\\Run Omni Web.cmd from the VS Code terminal and review the visible error"],
    ], [2400, 3100, 3860])
    doc.add_heading("15.1 Safe diagnostic sequence", 2)
    add_steps(doc, [
        "Confirm the VS Code workspace is D:\\RISC\\Omni.",
        "Confirm Run Omni Web.cmd remains open and the browser uses http://127.0.0.1:8000/.",
        "Run .\\.venv\\Scripts\\python.exe manage.py check.",
        "Review the System health page and the terminal error without sharing credentials or client data.",
        "Back up before attempting any data repair or migration.",
    ])


def appendices(doc: Document) -> None:
    doc.add_heading("Appendix A. Quick-reference workflows", 1)
    doc.add_heading("A.1 Complete an assessment", 2)
    add_steps(doc, [
        "Create organization and system context.", "Create the assessment and select framework(s).", "Complete plan, team, access, and owner assignments.",
        "Generate and manage evidence requests.", "Register and review supporting artifacts.", "Execute every applicable objective.",
        "Record each control conclusion and Assessor Notes/Findings.", "Create and track remediation for findings.",
        "Complete quality review and resolve blockers.", "Sign off, generate reports, and preserve the approved package.",
    ])
    doc.add_heading("A.2 Record a result", 2)
    add_steps(doc, [
        "Dashboard > Execute > Assess each objective.",
        "Dashboard > Control results > Control result.",
        "Select status and implementation; write the conformity statement or finding.",
        "Confirm owners, dates, and evidence; save and verify dashboard totals.",
    ])
    doc.add_heading("A.3 Register evidence", 2)
    add_steps(doc, ["Dashboard > Evidence > Register artifact.", "Upload a file or add an external reference.", "Record source, dates, request, and linked controls.", "Review and record acceptance status."])

    doc.add_heading("Appendix B. Status and terminology reference", 1)
    add_table(doc, ["Omni term", "Meaning", "Common aliases/examples"], [
        ["Framework requirement / control", "An assessable requirement from a selected authority or framework", "Practice, safeguard, criterion"],
        ["Assessment objective", "A framework-native determination supporting the control conclusion", "Examine/interview/test objective"],
        ["Supporting artifact", "Evidence file or governed external reference", "Evidence, record, document"],
        ["Assessor Notes/Findings", "One field used for a conformity statement or a finding based on the result", "Observation, conclusion"],
        ["Remediation Action Plan", "Governed response to a gap", "Corrective action plan, POA&M"],
        ["Security plan", "System design, boundaries, implementation, and control narrative", "System security plan, system design document"],
        ["Not Assessed", "No final conclusion", "Open, incomplete"],
    ], [2350, 4100, 2910])

    doc.add_page_break()
    doc.add_heading("Appendix C. Administrator command reference", 1)
    add_table(doc, ["Command", "Purpose"], [
        [".\\Run Omni Web.cmd", "Start the local web application"],
        [".\\.venv\\Scripts\\python.exe manage.py check", "Run Django configuration and model checks"],
        [".\\.venv\\Scripts\\python.exe manage.py backup_omni", "Create a database and evidence backup with integrity sidecar"],
        [".\\.venv\\Scripts\\python.exe manage.py verify_omni_backup <archive>", "Verify backup integrity without restoring"],
        [".\\.venv\\Scripts\\python.exe manage.py omni_readiness", "Evaluate future deployment prerequisites"],
    ], [4300, 5060])
    add_callout(doc, "Stop the server", "Return to the command window and press Ctrl+C. Do not close the window during an active save, import, export, or backup operation.", "tip")


def set_alt_text_for_image(doc: Document) -> None:
    descriptions = ["Eight-stage Omni assessment lifecycle from context and planning through reporting and monitoring"]
    descriptions.extend(f"Annotated Omni application screenshot: {title}; contains synthetic demonstration data" for _filename, title, _callouts, _legend in WALKTHROUGH_SCREENSHOTS)
    for index, drawing in enumerate(doc.element.body.iter(qn("w:drawing"))):
        for doc_pr in drawing.iter(qn("wp:docPr")):
            doc_pr.set("descr", descriptions[index] if index < len(descriptions) else "Omni product documentation illustration")


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Omni Comprehensive User and Administrator Manual"
    props.subject = "Framework-agnostic operating guidance for the Omni GRC assessment platform"
    props.author = "R!SC"
    props.last_modified_by = "R!SC"
    props.keywords = "Omni, GRC, assessment, evidence, remediation, risk, administration"
    props.comments = "Public product documentation; no client or organization data."


def build() -> Path:
    create_workflow_image()
    create_annotated_screenshots()
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    cover(doc)
    front_matter(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    section_11(doc)
    section_12(doc)
    section_13(doc)
    section_14(doc)
    section_15(doc)
    appendices(doc)
    set_alt_text_for_image(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
