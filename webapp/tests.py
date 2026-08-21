import tempfile
import hashlib
import json
import zipfile
import sqlite3
from datetime import date, timedelta
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core import mail
from django.conf import settings
from django.core.management import call_command
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from openpyxl import Workbook, load_workbook
from docx import Document

from .models import (
    Assessment,
    AssessmentAccess,
    AssessmentBaseline,
    AssessmentFramework,
    AssessmentReuseDecision,
    AssessmentObjective,
    AssessmentProcedure,
    AssessmentProcedureCustomization,
    AssessmentSample,
    AssessmentTeamMember,
    AssessmentTemplate,
    AuditEvent,
    AuthoritativeDocument,
    ControlAssessment,
    ComplianceAutomationPolicy,
    ComplianceAutomationRun,
    ControlMonitoringEvent,
    ControlMonitoringProfile,
    ControlReassessmentTask,
    EvidenceArtifact,
    EvidenceApplicability,
    EvidenceRequest,
    ExternalAuthority,
    Framework,
    FrameworkImport,
    GeneratedDocument,
    Membership,
    LoginAttempt,
    MappingReference,
    MappingChangeRequest,
    MappingHistory,
    RevalidationTask,
    Notification,
    OmniEvidenceSourceRequest,
    NotificationPolicy,
    NotificationPreference,
    IntegrationPolicy,
    ImplementationActivity,
    ImplementationActivityMapping,
    OutboundWorkItem,
    ObjectiveAssessment,
    Organization,
    OrganizationInvitation,
    RemediationMilestone,
    RemediationPlan,
    Requirement,
    RequirementMapping,
    RequirementRiskMapping,
    RiskCatalogEntry,
    RiskRegisterEntry,
    RiskRegisterHistory,
    RiskAcceptanceRequest,
    RiskReassessment,
    RiskTolerancePolicy,
    RiskTreatmentAction,
    Soc2AssessmentProfile,
    Soc2PointOfFocus,
    System,
    TestExecution,
    TestReuseReference,
    WorkflowHistory,
    UserProfile,
)
from .framework_import import materialize_mapping_references, parse_upload, resolve_catalog_mappings
from .harmonization import refresh_harmonization, review_reuse
from .mapping_governance import review_change
from .omni_evidence_catalog import import_catalog, normalize_cmmc
from .authoritative_sources import import_authoritative_sources
from .risk_heatmap import build_weighted_risk_heatmap
from .risk_catalog import import_risk_catalog
from .soc2_tsc import EXPECTED_BY_DOMAIN, install_baseline, load_catalog, validate_catalog
from .soc2_activity_import import ACTIVITY_TARGETS, MIRROR_SHEETS, REQUIRED_COLUMNS, import_activities, normalize_workbook
from .soc2_points_of_focus import import_points_of_focus
from .soc2_procedures import ensure_soc2_execution_catalog
from .soc2_evidence import (
    approve_test_reuse, create_soc2_evidence_requests, soc2_evidence_expectations,
)
from .reporting import (
    ReportNotReady, build_soc2_drl, build_soc2_readiness_package,
    build_soc2_report, soc2_report_readiness,
)


class Soc2TscBaselineTests(TestCase):
    def test_catalog_has_exact_authoritative_identifier_inventory(self):
        catalog, _ = load_catalog()
        report = validate_catalog(catalog)
        self.assertTrue(report["valid"], report["errors"])
        self.assertEqual(report["criterion_count"], 61)
        self.assertEqual(report["domain_counts"], {
            "Security — Common Criteria": 33,
            "Availability": 3,
            "Processing Integrity": 5,
            "Confidentiality": 2,
            "Privacy": 18,
        })
        actual = {row[0] for row in catalog["criteria"]}
        expected = {identifier for identifiers in EXPECTED_BY_DOMAIN.values() for identifier in identifiers}
        self.assertEqual(actual, expected)
        self.assertNotIn("CC8.2", actual)
        self.assertNotIn("CC9.3", actual)

    def test_installer_creates_immutable_idempotent_framework(self):
        framework, created, report = install_baseline()
        self.assertTrue(created)
        self.assertEqual(report["criterion_count"], 61)
        self.assertEqual(framework.requirements.count(), 61)
        self.assertEqual(framework.scoring_method, Framework.ScoringMethod.NONE)
        self.assertTrue(framework.requirements.filter(requirement_id="CC3.4").exists())
        self.assertTrue(framework.requirements.filter(requirement_id="P8.1").exists())
        same, created_again, _ = install_baseline()
        self.assertFalse(created_again)
        self.assertEqual(same.pk, framework.pk)

    def test_validator_rejects_missing_and_invented_identifiers(self):
        catalog, _ = load_catalog()
        catalog["criteria"] = [row for row in catalog["criteria"] if row[0] != "CC7.5"]
        catalog["criteria"].append(["CC8.2", "Security — Common Criteria", "Invented"])
        report = validate_catalog(catalog)
        self.assertFalse(report["valid"])
        self.assertTrue(any("CC7.5" in error for error in report["errors"]))
        self.assertTrue(any("CC8.2" in error for error in report["errors"]))

    def test_existing_version_with_different_digest_is_rejected(self):
        framework, _, _ = install_baseline()
        framework.source_sha256 = "0" * 64
        framework.save(update_fields=("source_sha256",))
        with self.assertRaisesRegex(ValueError, "immutable"):
            install_baseline()

    def test_catalog_metadata_does_not_claim_to_reproduce_aicpa_text(self):
        catalog, _ = load_catalog()
        metadata = catalog["framework"]
        self.assertIn("not reproduced", metadata["content_scope"])
        self.assertIn("copyrighted", metadata["copyright_notice"])
        for _, _, label in catalog["criteria"]:
            self.assertTrue(label.strip())


class Soc2ActivityNormalizationTests(TestCase):
    def setUp(self):
        install_baseline()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)

    def workbook_path(self, mutate=None):
        workbook = Workbook()
        master = workbook.active
        master.title = "All Controls"
        master.append(list(REQUIRED_COLUMNS))
        rows = {}
        for identifier in ACTIVITY_TARGETS:
            if identifier.startswith("CC"):
                area = "Security (Common Criteria)"
            elif identifier.startswith("SEC-"):
                area = "Security Controls"
            elif identifier.startswith("PRIV-"):
                area = "Privacy Controls"
            elif identifier.startswith("CONF-"):
                area = "Confidentiality Controls"
            elif identifier.startswith("AVAIL-"):
                area = "Availability Controls"
            else:
                area = "Processing Integrity Controls"
            row = [identifier, area, "Synthetic category", f"Activity {identifier}",
                   "Preventive", "", "", "", "", "", "Not Started", "", "", ""]
            rows[identifier] = row
            master.append(row)
        for sheet_name, prefix in MIRROR_SHEETS.items():
            sheet = workbook.create_sheet(sheet_name)
            sheet.append(list(REQUIRED_COLUMNS))
            for identifier, row in rows.items():
                if identifier.startswith(prefix):
                    sheet.append(row)
        if mutate:
            mutate(workbook)
        path = Path(self.temp_dir.name) / "soc2-activities.xlsx"
        workbook.save(path)
        workbook.close()
        return path

    def test_normalizer_reads_all_tabs_and_preserves_all_activities(self):
        normalized, report = normalize_workbook(self.workbook_path())
        self.assertTrue(report["valid"], report["errors"])
        self.assertEqual(report["activity_count"], 60)
        self.assertEqual(len(normalized["activities"]), 60)
        self.assertEqual(report["mirror_counts"], {
            "CC-Series (Required)": 32, "Security Controls": 6,
            "Privacy Controls": 6, "Confidentiality Controls": 5,
            "Availability Controls": 5, "Processing Integrity": 6,
        })
        self.assertGreater(report["mapping_count"], report["activity_count"])

    def test_normalizer_rejects_category_tab_that_differs_from_master(self):
        def mutate(workbook):
            workbook["Security Controls"]["D2"] = "Changed activity"
        _, report = normalize_workbook(self.workbook_path(mutate))
        self.assertFalse(report["valid"])
        self.assertTrue(any("differs from master" in error for error in report["errors"]))

    def test_import_creates_proposed_resolved_support_mappings_and_is_idempotent(self):
        path = self.workbook_path()
        result, report = import_activities(path)
        self.assertEqual(result["created"], 60)
        self.assertEqual(ImplementationActivity.objects.count(), 60)
        self.assertEqual(ImplementationActivityMapping.objects.count(), report["mapping_count"])
        self.assertFalse(ImplementationActivityMapping.objects.filter(target_requirement__isnull=True).exists())
        self.assertFalse(ImplementationActivityMapping.objects.exclude(
            review_status=ImplementationActivityMapping.ReviewStatus.PROPOSED,
            relationship=RequirementMapping.Relationship.SUPPORTS,
        ).exists())
        again, _ = import_activities(path)
        self.assertEqual(again, {"created": 0, "existing": 60, "mappings_created": 0})


class Soc2AssessmentModelTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.assessor = users.objects.create_user("soc2-assessor", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic SOC 2", slug="soc2")
        Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.system = System.objects.create(organization=self.organization, name="SOC 2 System")
        self.framework, _, _ = install_baseline()
        self.client.login(username="soc2-assessor", password="test-password")

    def create_type_two(self, optional_categories=None):
        response = self.client.post(
            reverse("assessment-create", args=("soc2", self.system.id)),
            {
                "name": "SOC 2 Type II Assessment",
                "frameworks": [self.framework.id],
                "primary_framework": self.framework.id,
                "examination_type": Soc2AssessmentProfile.ExaminationType.TYPE_II,
                "optional_categories": optional_categories or [],
                "period_start": "2026-01-01", "period_end": "2026-06-30",
                "scope_notes": "Synthetic service scope.",
            },
        )
        return response, Assessment.objects.filter(name="SOC 2 Type II Assessment").first()

    def test_type_two_creation_requires_security_and_scopes_optional_categories(self):
        response, assessment = self.create_type_two([
            Soc2AssessmentProfile.Category.PRIVACY
        ])
        self.assertIsNotNone(assessment)
        self.assertRedirects(response, reverse("assessment-dashboard", args=("soc2", assessment.id)))
        profile = assessment.soc2_profile
        self.assertEqual(profile.included_categories, ["SECURITY", "PRIVACY"])
        self.assertEqual(assessment.control_results.filter(in_scope=True).count(), 51)
        excluded = assessment.control_results.filter(in_scope=False)
        self.assertEqual(excluded.count(), 10)
        self.assertFalse(excluded.exclude(
            status=ControlAssessment.Status.NOT_APPLICABLE,
            implementation_state=ControlAssessment.Implementation.NA,
        ).exists())

    def test_type_one_requires_as_of_date_and_rejects_period(self):
        response = self.client.post(
            reverse("assessment-create", args=("soc2", self.system.id)),
            {
                "name": "Invalid Type I", "frameworks": [self.framework.id],
                "primary_framework": self.framework.id,
                "examination_type": Soc2AssessmentProfile.ExaminationType.TYPE_I,
                "period_start": "2026-01-01", "period_end": "2026-01-31",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Type I requires an as-of date")
        self.assertFalse(Assessment.objects.filter(name="Invalid Type I").exists())

    def test_plan_can_expand_scope_without_retaining_auto_na_conclusions(self):
        _, assessment = self.create_type_two()
        self.assertEqual(assessment.control_results.filter(in_scope=True).count(), 33)
        response = self.client.post(
            reverse("assessment-plan", args=("soc2", assessment.id)),
            {
                "action": "soc2",
                "soc2-examination_type": Soc2AssessmentProfile.ExaminationType.TYPE_II,
                "soc2-optional_categories": [Soc2AssessmentProfile.Category.AVAILABILITY],
                "soc2-period_start": "2026-01-01", "soc2-period_end": "2026-06-30",
            },
        )
        self.assertRedirects(response, reverse("assessment-plan", args=("soc2", assessment.id)))
        availability = assessment.control_results.filter(requirement__domain="Availability")
        self.assertEqual(availability.filter(in_scope=True).count(), 3)
        self.assertFalse(availability.exclude(
            status=ControlAssessment.Status.NOT_ASSESSED,
            implementation_state=ControlAssessment.Implementation.UNASSESSED,
        ).exists())

    def test_dashboard_displays_soc2_profile(self):
        _, assessment = self.create_type_two([Soc2AssessmentProfile.Category.CONFIDENTIALITY])
        response = self.client.get(reverse("assessment-dashboard", args=("soc2", assessment.id)))
        self.assertContains(response, "Type II")
        self.assertContains(response, "Security (required), Confidentiality")


class Soc2ExecutionProcedureTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.assessor = users.objects.create_user("soc2-executor", password="test-password")
        self.organization = Organization.objects.create(name="SOC 2 Execution", slug="soc2-exec")
        self.member = Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.system = System.objects.create(
            organization=self.organization, name="Execution System",
            description="Production services and supporting operational processes.",
        )
        self.framework, _, _ = install_baseline()
        ensure_soc2_execution_catalog()
        self.client.login(username="soc2-executor", password="test-password")

    def create_assessment(self, examination_type="TYPE_II"):
        data = {
            "name": f"SOC 2 {examination_type}", "frameworks": [self.framework.id],
            "primary_framework": self.framework.id, "examination_type": examination_type,
            "service_commitments": "Protect customer information and maintain reliable service.",
        }
        if examination_type == "TYPE_I":
            data["as_of_date"] = "2026-06-30"
        else:
            data.update({"period_start": "2026-01-01", "period_end": "2026-06-30"})
        response = self.client.post(
            reverse("assessment-create", args=(self.organization.slug, self.system.id)), data
        )
        assessment = Assessment.objects.get(name=f"SOC 2 {examination_type}")
        self.assertEqual(response.status_code, 302)
        return assessment

    def test_catalog_creates_five_omni_procedures_per_criterion(self):
        result = ensure_soc2_execution_catalog()
        self.assertEqual(result["requirements"], 61)
        self.assertEqual(AssessmentObjective.objects.filter(
            requirement__framework=self.framework, objective_id="OMNI-SOC2-CRITERION"
        ).count(), 61)
        self.assertEqual(AssessmentProcedure.objects.filter(
            requirement__framework=self.framework
        ).count(), 305)
        self.assertEqual(set(AssessmentProcedure.objects.filter(
            requirement__framework=self.framework
        ).values_list("method", flat=True)), {"EXAMINE", "INTERVIEW", "OBSERVE", "TEST", "REPERFORM"})
        procedure = AssessmentProcedure.objects.filter(requirement__framework=self.framework).first()
        self.assertEqual(str(procedure), f"{procedure.get_method_display()}: {procedure.assessment_object}")

    def test_type_two_requires_all_three_conclusions(self):
        assessment = self.create_assessment()
        result = assessment.control_results.filter(in_scope=True).first().objective_results.get()
        url = reverse("objective-edit", args=(self.organization.slug, assessment.id, result.id))
        invalid = self.client.post(url, {
            "action": "objective", "status": "MET",
            "design_conclusion": "EFFECTIVE", "implementation_conclusion": "EFFECTIVE",
            "operating_effectiveness_conclusion": "NOT_TESTED",
            "assessor_notes": "Designed and implemented.",
        })
        self.assertContains(invalid, "Type II requires an operating-effectiveness conclusion")
        self.assertNotContains(invalid, "This field is required.")
        valid = self.client.post(url, {
            "action": "objective", "status": "MET",
            "design_conclusion": "EFFECTIVE", "implementation_conclusion": "EFFECTIVE",
            "operating_effectiveness_conclusion": "EFFECTIVE",
            "assessor_notes": "Designed, implemented, and operating effectively.",
        })
        self.assertRedirects(valid, reverse("assessment-execution", args=(self.organization.slug, assessment.id)))

    def test_type_one_forces_operating_effectiveness_not_applicable(self):
        assessment = self.create_assessment("TYPE_I")
        result = assessment.control_results.filter(in_scope=True).first().objective_results.get()
        response = self.client.post(
            reverse("objective-edit", args=(self.organization.slug, assessment.id, result.id)),
            {"action": "objective", "status": "MET", "design_conclusion": "EFFECTIVE",
             "implementation_conclusion": "EFFECTIVE",
             "operating_effectiveness_conclusion": "EFFECTIVE",
             "assessor_notes": "Suitably designed and implemented as of the date."},
        )
        self.assertEqual(response.status_code, 302)
        result.refresh_from_db()
        self.assertEqual(result.operating_effectiveness_conclusion, "NOT_APPLICABLE")

    def test_custom_procedure_is_assessment_specific(self):
        assessment = self.create_assessment()
        result = assessment.control_results.filter(in_scope=True).first().objective_results.get()
        base = result.objective.procedures.first()
        response = self.client.post(
            reverse("objective-edit", args=(self.organization.slug, assessment.id, result.id)),
            {"action": "procedure", "procedure-base_procedure": base.id,
             "procedure-method": "REPERFORM",
             "procedure-procedure_text": "Reperform the selected access review sample.",
             "procedure-enabled": "on"},
        )
        self.assertEqual(response.status_code, 302)
        customization = AssessmentProcedureCustomization.objects.get()
        self.assertEqual(customization.objective_result, result)
        self.assertEqual(customization.updated_by, self.assessor)

    def test_type_two_sample_tracks_population_and_period(self):
        assessment = self.create_assessment()
        objective = assessment.control_results.filter(in_scope=True).first().objective_results.get()
        url = reverse("sample-create", args=(self.organization.slug, assessment.id))
        invalid = self.client.post(url, {
            "name": "Outside period", "population_description": "Access changes",
            "population_size": 20, "sample_size": 5, "period_start": "2025-12-01",
            "period_end": "2026-06-30", "selection_method": "Random",
            "selected_items": "1, 2, 3, 4, 5", "objectives": [objective.id],
        })
        self.assertContains(invalid, "Cannot precede the Type II examination period")
        valid = self.client.post(url, {
            "name": "Valid population", "population_description": "Access changes",
            "population_size": 20, "sample_size": 5, "period_start": "2026-01-01",
            "period_end": "2026-06-30", "selection_method": "Random",
            "selected_items": "1, 2, 3, 4, 5", "objectives": [objective.id],
        })
        self.assertEqual(valid.status_code, 302)
        sample = AssessmentSample.objects.get(name="Valid population")
        self.assertEqual(sample.population_size, 20)
        self.assertEqual(sample.period_start.isoformat(), "2026-01-01")

    def test_authorized_points_of_focus_import_retains_provenance(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Licensed POF"
        sheet.append(["Criterion ID", "Point of Focus ID", "Point of Focus Text", "Source Reference", "Page"])
        sheet.append(["CC1.1", "POF-1", "Authorized synthetic point text.", "Licensed test source", 10])
        path = Path(tempfile.mkdtemp()) / "licensed-pof.xlsx"
        self.addCleanup(lambda: path.unlink(missing_ok=True))
        workbook.save(path)
        workbook.close()
        result, report = import_points_of_focus(path)
        self.assertTrue(report["valid"])
        self.assertEqual(result["created"], 1)
        point = Soc2PointOfFocus.objects.get()
        self.assertEqual(point.requirement.requirement_id, "CC1.1")
        self.assertEqual(point.source_row, 2)
        self.assertEqual(point.source_page, 10)

    def test_private_activity_guidance_generates_consolidated_evidence_requests(self):
        assessment = self.create_assessment()
        requirement = assessment.control_results.filter(in_scope=True).first().requirement
        activity = ImplementationActivity.objects.create(
            source_identifier="SYN-01", source_area="Synthetic", activity="Review access.",
            source_filename="private.xlsx", source_sha256="a" * 64,
            source_sheet="All Controls", source_row=2,
            source_metadata={"Evidence Artifact": "Access review records",
                             "Evidence Source": "Identity platform", "Frequency": "Quarterly"},
        )
        mapping = ImplementationActivityMapping.objects.create(
            activity=activity, target_framework_code=self.framework.code,
            target_requirement_id_text=requirement.requirement_id,
            target_requirement=requirement, review_status="APPROVED",
        )
        suggestions = soc2_evidence_expectations(assessment)
        self.assertEqual(len(suggestions), 1)
        result = create_soc2_evidence_requests(assessment, [mapping.id], self.assessor)
        self.assertEqual(result, {"selected": 1, "created": 1, "control_links": 1})
        evidence_request = EvidenceRequest.objects.get(title="Access review records")
        self.assertEqual(evidence_request.controls.get().requirement, requirement)

    def test_approved_test_reuse_references_work_without_copying_conclusion(self):
        assessment = self.create_assessment()
        source_result = assessment.control_results.filter(in_scope=True).first()
        target_framework = Framework.objects.create(code="TARGET", name="Target", version="1")
        target_requirement = Requirement.objects.create(
            framework=target_framework, requirement_id="T-1", title="Target",
            statement="Synthetic target requirement.",
        )
        target_result = ControlAssessment.objects.create(
            assessment=assessment, requirement=target_requirement,
        )
        target_objective = AssessmentObjective.objects.create(
            requirement=target_requirement, objective_id="T-1.a", text="Assess target.",
        )
        target_objective_result = ObjectiveAssessment.objects.create(
            control_result=target_result, objective=target_objective,
        )
        source_objective_result = source_result.objective_results.get()
        source_test = TestExecution.objects.create(
            assessment=assessment, objective_result=source_objective_result,
            performed_by=self.member, performed_at=timezone.now(),
            steps_performed="Inspected a synthetic sample.", actual_result="No exceptions.",
            outcome=TestExecution.Outcome.PASS,
        )
        decision = AssessmentReuseDecision.objects.create(
            assessment=assessment, source_result=source_result, target_result=target_result,
            basis=AssessmentReuseDecision.Basis.DIRECT, relationship="RELATED",
            status=AssessmentReuseDecision.Status.APPROVED, reuse_testing=True,
        )
        reference, created = approve_test_reuse(
            decision, source_test.id, target_objective_result.id, self.assessor,
            "Confirm target-specific language separately.",
        )
        self.assertTrue(created)
        self.assertEqual(TestReuseReference.objects.get(), reference)
        target_objective_result.refresh_from_db()
        self.assertEqual(target_objective_result.status, ObjectiveAssessment.Status.NOT_ASSESSED)

    def test_unapproved_mapping_does_not_generate_harmonization_candidate(self):
        assessment = self.create_assessment()
        source = assessment.control_results.filter(in_scope=True).first()
        other_framework = Framework.objects.create(code="DRAFT-MAP", name="Draft", version="1")
        other_requirement = Requirement.objects.create(
            framework=other_framework, requirement_id="D-1", title="Draft", statement="Draft.",
        )
        other_result = ControlAssessment.objects.create(
            assessment=assessment, requirement=other_requirement,
        )
        RequirementMapping.objects.create(
            source=source.requirement, target=other_requirement,
            relationship="RELATED", lifecycle="DRAFT",
        )
        self.assertEqual(refresh_harmonization(assessment)["candidates"], 0)
        self.assertFalse(AssessmentReuseDecision.objects.filter(
            source_result__in=(source, other_result), target_result__in=(source, other_result)
        ).exists())

    def _complete_soc2_execution(self, assessment):
        assessment.control_results.filter(in_scope=True).update(
            status=ControlAssessment.Status.MET,
            assessor_notes_findings="Criterion requirements were satisfied.",
        )
        values = {
            "status": ObjectiveAssessment.Status.MET,
            "design_conclusion": ObjectiveAssessment.Conclusion.EFFECTIVE,
            "implementation_conclusion": ObjectiveAssessment.Conclusion.EFFECTIVE,
            "assessor_notes": "Control design and implementation were evaluated without exception.",
            "assessed_by": self.assessor, "assessed_at": timezone.now(),
        }
        if assessment.soc2_profile.examination_type == "TYPE_I":
            values["operating_effectiveness_conclusion"] = "NOT_APPLICABLE"
        else:
            values["operating_effectiveness_conclusion"] = "EFFECTIVE"
        ObjectiveAssessment.objects.filter(
            control_result__assessment=assessment, control_result__in_scope=True
        ).update(**values)

    def test_soc2_type_two_readiness_requires_operating_conclusions(self):
        assessment = self.create_assessment()
        assessment.control_results.filter(in_scope=True).update(status="MET")
        ObjectiveAssessment.objects.filter(
            control_result__assessment=assessment, control_result__in_scope=True
        ).update(
            status="MET", design_conclusion="EFFECTIVE",
            implementation_conclusion="EFFECTIVE", assessor_notes="Implemented.",
        )
        readiness = soc2_report_readiness(assessment)
        self.assertFalse(readiness["ready"])
        self.assertTrue(any("operating-effectiveness" in item for item in readiness["blockers"]))
        with self.assertRaises(ReportNotReady):
            build_soc2_report(assessment)

    def test_type_one_and_type_two_reports_are_scope_aware_and_not_cpa_opinions(self):
        for examination_type, measurement_text in (
            ("TYPE_I", "Measurement date"), ("TYPE_II", "Examination period")
        ):
            assessment = self.create_assessment(examination_type)
            self._complete_soc2_execution(assessment)
            readiness = soc2_report_readiness(assessment)
            self.assertTrue(readiness["ready"], readiness["blockers"])
            self.assertEqual(readiness["total_criteria"], 33)
            content = build_soc2_report(assessment)
            document = Document(BytesIO(content))
            text = "\n".join(item.text for item in document.paragraphs)
            self.assertIn(measurement_text, text)
            self.assertIn("not an AICPA SOC 2 report", text)
            self.assertIn("Security (required)", text)
            self.assertIn("Protect customer information", text)
            self.assertIn("Production services and supporting operational processes", text)
            self.assertGreater(len(document.tables[0].rows), 33)
            assessment.delete()

    def test_soc2_report_download_records_audited_generated_document(self):
        assessment = self.create_assessment()
        self._complete_soc2_execution(assessment)
        url = reverse("report-download", args=(
            self.organization.slug, assessment.id, "SOC2_REPORT"
        ))
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        record = GeneratedDocument.objects.get(
            assessment=assessment, kind=GeneratedDocument.Kind.SOC2_REPORT
        )
        self.assertIn("SOC-2-TYPE-II-Readiness-Report", record.filename)
        self.assertTrue(record.content_sha256)
        self.assertTrue(AuditEvent.objects.filter(
            action="document.generated", object_id=str(record.id)
        ).exists())

    def test_soc2_readiness_package_contains_drl_evidence_and_verified_manifest(self):
        assessment = self.create_assessment()
        self._complete_soc2_execution(assessment)
        control = assessment.control_results.filter(in_scope=True).first()
        evidence_request = EvidenceRequest.objects.create(
            assessment=assessment, evidence_code="SOC2-DRL-001",
            title="Access review records", description="Provide quarterly access reviews.",
            owner=self.member, created_by=self.assessor,
        )
        evidence_request.controls.add(control)
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=assessment,
            title="Access review portal", external_reference="https://example.test/access-review",
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED,
            period_start=date(2026, 1, 1), period_end=date(2026, 6, 30),
            uploaded_by=self.assessor,
        )
        artifact.controls.add(control); artifact.requests.add(evidence_request)
        drl = load_workbook(BytesIO(build_soc2_drl(assessment)))
        self.assertEqual(drl["Document Request List"]["A2"].value, "SOC2-DRL-001")
        content, readiness = build_soc2_readiness_package(assessment, self.assessor)
        self.assertTrue(readiness["ready"])
        with zipfile.ZipFile(BytesIO(content)) as package:
            names = set(package.namelist())
            expected = {
                "Deliverables/Omni-SOC-2-Readiness-Report.docx",
                "Attachments/Omni-SOC-2-Document-Request-List.xlsx",
                "Attachments/Omni-SOC-2-Traceability-Matrix.csv",
                "Attachments/Omni-Remediation-Action-Plan.xlsx",
                "Attachments/SOC-2-Scope-Summary.json",
                "Evidence/Evidence-Index.csv", "Package-Manifest.json",
            }
            self.assertTrue(expected.issubset(names))
            self.assertTrue(any(name.endswith("EA-0001_External_Reference.txt") for name in names))
            manifest = json.loads(package.read("Package-Manifest.json"))
            self.assertEqual(manifest["package_type"], "SOC 2 readiness assessment package")
            for item in manifest["files"]:
                payload = package.read(item["path"])
                self.assertEqual(len(payload), item["size_bytes"])
                self.assertEqual(hashlib.sha256(payload).hexdigest(), item["sha256"])

    def test_soc2_package_download_is_audited(self):
        assessment = self.create_assessment()
        self._complete_soc2_execution(assessment)
        response = self.client.get(reverse("report-download", args=(
            self.organization.slug, assessment.id, "SOC2_PACKAGE"
        )))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/zip")
        record = GeneratedDocument.objects.get(kind=GeneratedDocument.Kind.SOC2_PACKAGE)
        self.assertTrue(record.filename.endswith("SOC-2-Readiness-Assessment-Package.zip"))
        self.assertTrue(record.content_sha256)
        drl_response = self.client.get(reverse("report-download", args=(
            self.organization.slug, assessment.id, "SOC2_DRL"
        )))
        self.assertEqual(drl_response.status_code, 200)
        self.assertEqual(
            drl_response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        self.assertTrue(GeneratedDocument.objects.filter(
            assessment=assessment, kind=GeneratedDocument.Kind.SOC2_DRL,
            filename__endswith="SOC-2-Document-Request-List.xlsx",
        ).exists())


class SprintSeventeenPointSevenRiskHeatmapTests(TestCase):
    def test_ccf_column_f_weight_is_imported_and_not_treated_as_a_mapping(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append([
            "CCF Domain", "CCF Control", "CCF #",
            "Comprehensive Controls Framework (CCF) Control Description",
            "CCF Control Question", "Control Weighting", "NIST CSF 2.0",
        ])
        sheet.append(["Governance", "Security program", "GOV-01", "Maintain a program.",
                      "Is a program maintained?", 9, "GV.OC-01"])
        sheet.append(["Technology", "Deprecated control", "TDA-11.2", "[deprecated]",
                      "[deprecated]", 0, ""])
        stream = BytesIO()
        workbook.save(stream)
        workbook.close()
        upload = SimpleUploadedFile(
            "ccf.xlsx", stream.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        normalized, report, _, _ = parse_upload(upload, {
            "code": "OMNI-CCF", "name": "Omni CCF", "version": "2026.2",
            "is_omni_control_framework": True,
        })

        requirement = normalized["requirements"][0]
        self.assertTrue(report["valid"])
        self.assertEqual(requirement["risk_weight"], 9)
        self.assertEqual(normalized["requirements"][1]["risk_weight"], 0)
        self.assertEqual(report["authority_count"], 1)
        self.assertEqual(requirement["mapping_refs"][0]["source_column"], 7)
        self.assertEqual(requirement["mapping_refs"][0]["target_framework"], "NIST CSF 2.0")

    def test_heatmap_uses_sprs_and_omni_weights_without_inventing_likelihood(self):
        sprs = SimpleNamespace(scoring_method=Framework.ScoringMethod.SPRS,
                               is_omni_control_framework=False)
        omni = SimpleNamespace(scoring_method=Framework.ScoringMethod.NONE,
                               is_omni_control_framework=True)
        results = [
            SimpleNamespace(requirement=SimpleNamespace(framework=sprs, domain="Access Control",
                                                        full_deduction=5, risk_weight=None),
                            status=ControlAssessment.Status.NOT_MET, calculated_deduction=5),
            SimpleNamespace(requirement=SimpleNamespace(framework=sprs, domain="Access Control",
                                                        full_deduction=1, risk_weight=None),
                            status=ControlAssessment.Status.MET, calculated_deduction=0),
            SimpleNamespace(requirement=SimpleNamespace(framework=omni, domain="Governance",
                                                        full_deduction=1, risk_weight=10),
                            status=ControlAssessment.Status.NOT_MET, calculated_deduction=1),
            SimpleNamespace(requirement=SimpleNamespace(framework=omni, domain="Governance",
                                                        full_deduction=1, risk_weight=5),
                            status=ControlAssessment.Status.NOT_ASSESSED, calculated_deduction=0),
        ]

        heatmap = build_weighted_risk_heatmap(results)

        self.assertEqual(heatmap["sources"], "Omni 0–10 + SPRS")
        access = next(item for item in heatmap["cells"] if item["domain"] == "Access Control")
        governance = next(item for item in heatmap["cells"] if item["domain"] == "Governance")
        self.assertEqual(access["exposure"], 83)
        self.assertEqual(access["severity"], "critical")
        self.assertEqual(governance["exposure"], 100)
        self.assertEqual(governance["unknown_weight"], 5)


class SprintSeventeenPointEightRiskCatalogTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.admin = users.objects.create_superuser(
            "risk-admin", "risk-admin@example.com", "test-password"
        )
        self.regular = users.objects.create_user("risk-user", password="test-password")

    def _catalog(self, directory):
        path = Path(directory) / "CCF Risk Catalog.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Risk Catalog"
        sheet.append(["RISK CATALOG use case", None, None, None])
        for _ in range(4):
            sheet.append([])
        sheet.append(["Risk Grouping", "Risk #", "Risk", "Description"])
        sheet.append([None, None, None, "IF THE CONTROL FAILS"])
        sheet.append(["Access Control", "R-AC-1", "Loss of accountability", "Accountability may be lost."])
        sheet.append([None, "R-AC-2", "Unauthorized access", "Unauthorized access may occur."])
        workbook.save(path)
        workbook.close()
        return path

    def test_private_import_preserves_canonical_risks_and_provenance(self):
        with tempfile.TemporaryDirectory() as directory:
            report = import_risk_catalog(self._catalog(directory), apply=True)
        self.assertTrue(report["valid"])
        self.assertEqual(report["records"], 2)
        self.assertEqual(report["groups"], 1)
        risk = RiskCatalogEntry.objects.get(risk_id="R-AC-2")
        self.assertEqual(risk.grouping, "Access Control")
        self.assertEqual(risk.source_row, 9)
        self.assertEqual(len(risk.source_sha256), 64)

    def test_only_approved_mapping_drives_assessment_risk_visibility(self):
        organization = Organization.objects.create(name="Risk Client", slug="risk-client")
        Membership.objects.create(user=self.regular, organization=organization,
                                  role=Membership.Role.ASSESSOR)
        system = System.objects.create(organization=organization, name="Risk System")
        framework = Framework.objects.create(code="RISK-FW", name="Risk Framework", version="1")
        requirement = Requirement.objects.create(framework=framework, requirement_id="AC-1",
                                                 domain="Access Control", title="Access",
                                                 statement="Restrict access.")
        assessment = Assessment.objects.create(system=system, framework=framework,
                                               name="Risk Assessment", created_by=self.regular,
                                               risk_management_enabled=True)
        AssessmentFramework.objects.create(assessment=assessment, framework=framework,
                                           is_primary=True, added_by=self.regular)
        ControlAssessment.objects.create(assessment=assessment, requirement=requirement,
                                         status=ControlAssessment.Status.NOT_MET)
        risk = RiskCatalogEntry.objects.create(
            risk_id="R-AC-4", grouping="Access Control", title="Unauthorized access",
            description="Unauthorized access may occur.", source_row=11,
            source_filename="private.xlsx", source_sha256="a" * 64,
        )
        mapping = RequirementRiskMapping.objects.create(
            requirement=requirement, risk=risk, rationale="Control directly restricts access.",
            proposed_by=self.admin,
        )
        self.client.login(username="risk-user", password="test-password")
        url = reverse("assessment-dashboard", args=(organization.slug, assessment.id))
        self.assertNotContains(self.client.get(url), "R-AC-4")
        mapping.review_status = RequirementRiskMapping.ReviewStatus.APPROVED
        mapping.reviewed_by = self.admin
        mapping.reviewed_at = timezone.now()
        mapping.save()
        self.assertContains(self.client.get(url), "R-AC-4")

    def test_registry_is_superuser_only(self):
        self.client.force_login(self.regular)
        self.assertEqual(self.client.get(reverse("risk-catalog-registry")).status_code, 404)
        self.client.force_login(self.admin)
        self.assertEqual(self.client.get(reverse("risk-catalog-registry")).status_code, 200)


class SprintEighteenRiskRegisterTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.admin = users.objects.create_user("risk-owner-admin", password="test-password")
        self.assessor = users.objects.create_user("risk-assessor", password="test-password")
        self.outsider = users.objects.create_user("risk-outsider", password="test-password")
        self.organization = Organization.objects.create(name="Risk Organization", slug="risk-org")
        other = Organization.objects.create(name="Other Risk Organization", slug="other-risk-org")
        self.admin_membership = Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.assessor_membership = Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        Membership.objects.create(user=self.outsider, organization=other, role=Membership.Role.VIEWER)
        system = System.objects.create(organization=self.organization, name="Risk System")
        framework = Framework.objects.create(code="RISK-18", name="Risk Framework", version="1")
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="AC-18", domain="Access Control",
            title="Access control", statement="Restrict system access.",
        )
        self.assessment = Assessment.objects.create(
            system=system, framework=framework, name="Sprint 18 Assessment", created_by=self.assessor,
            risk_management_enabled=True,
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=framework, is_primary=True, added_by=self.assessor
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=requirement,
            status=ControlAssessment.Status.NOT_MET,
            assessor_notes_findings="Access is not sufficiently restricted.",
        )

    def _risk_payload(self, **overrides):
        payload = {
            "catalog_risk": "", "title": "Unauthorized access risk",
            "description": "Threat actors may gain unauthorized access.",
            "category": "Access Control", "controls": [str(self.control.id)],
            "remediation_plans": [], "owner": str(self.assessor_membership.id),
            "status": RiskRegisterEntry.Status.IDENTIFIED, "likelihood": "4", "impact": "5",
            "treatment": RiskRegisterEntry.Treatment.MITIGATE,
            "treatment_plan": "Strengthen access restrictions and validate effectiveness.",
            "target_date": "2026-12-31", "residual_likelihood": "2", "residual_impact": "3",
            "acceptance_rationale": "", "acceptance_expires": "", "next_review_date": "2026-10-01",
        }
        payload.update(overrides)
        return payload

    def test_finding_creates_scored_risk_with_history_and_dashboard_matrix(self):
        self.client.login(username="risk-assessor", password="test-password")
        create_url = reverse("risk-register-create", args=(self.organization.slug, self.assessment.id))
        preview = self.client.get(create_url, {"control": self.control.id})
        self.assertContains(preview, "Access is not sufficiently restricted")
        response = self.client.post(create_url, self._risk_payload())
        risk = RiskRegisterEntry.objects.get()
        self.assertRedirects(response, reverse(
            "risk-register-detail", args=(self.organization.slug, self.assessment.id, risk.id)
        ))
        self.assertEqual(risk.risk_id, "RISK-0001")
        self.assertEqual(risk.inherent_score, 20)
        self.assertEqual(risk.residual_score, 6)
        self.assertEqual(risk.controls.get(), self.control)
        self.assertTrue(RiskRegisterHistory.objects.filter(risk=risk, action="CREATED").exists())
        dashboard = self.client.get(reverse(
            "assessment-dashboard", args=(self.organization.slug, self.assessment.id)
        ))
        self.assertContains(dashboard, "Likelihood × impact heatmap")
        self.assertContains(dashboard, "1 critical")

    def test_only_admin_can_accept_risk_and_acceptance_is_audited(self):
        risk = RiskRegisterEntry.objects.create(
            organization=self.organization, system=self.assessment.system, assessment=self.assessment,
            risk_id="RISK-0001", title="Acceptance candidate", description="Candidate risk.",
            category="Governance", created_by=self.assessor,
        )
        edit_url = reverse("risk-register-edit", args=(self.organization.slug, self.assessment.id, risk.id))
        acceptance = self._risk_payload(
            title=risk.title, description=risk.description, category=risk.category,
            treatment=RiskRegisterEntry.Treatment.ACCEPT,
            treatment_plan="Accept temporarily while replacement is funded.",
            acceptance_rationale="Exposure is within approved tolerance temporarily.",
            acceptance_expires="2027-01-31",
        )
        self.client.login(username="risk-assessor", password="test-password")
        self.assertEqual(self.client.post(edit_url, acceptance).status_code, 200)
        risk.refresh_from_db(); self.assertNotEqual(risk.status, RiskRegisterEntry.Status.ACCEPTED)
        self.client.login(username="risk-owner-admin", password="test-password")
        self.assertRedirects(self.client.post(edit_url, acceptance), reverse(
            "risk-register-detail", args=(self.organization.slug, self.assessment.id, risk.id)
        ))
        risk.refresh_from_db()
        self.assertEqual(risk.status, RiskRegisterEntry.Status.ACCEPTED)
        self.assertEqual(risk.accepted_by, self.admin)
        self.assertTrue(RiskRegisterHistory.objects.filter(risk=risk, action="UPDATED").exists())

    def test_register_and_export_are_tenant_scoped(self):
        list_url = reverse("risk-register-list", args=(self.organization.slug, self.assessment.id))
        export_url = reverse("risk-register-export", args=(self.organization.slug, self.assessment.id))
        self.client.login(username="risk-outsider", password="test-password")
        self.assertEqual(self.client.get(list_url).status_code, 404)
        self.assertEqual(self.client.get(export_url).status_code, 404)
        self.client.login(username="risk-assessor", password="test-password")
        response = self.client.get(export_url)
        self.assertEqual(response.status_code, 200)
        self.assertIn("Risk ID,Title,Category", response.content.decode("utf-8"))
        self.assertTrue(AuditEvent.objects.filter(
            organization=self.organization, action="risk.exported"
        ).exists())

    def test_approved_finding_risk_is_guided_and_cannot_be_registered_twice(self):
        catalog = RiskCatalogEntry.objects.create(
            risk_id="R-AC-4", grouping="Access Control", title="Unauthorized access",
            description="Unauthorized access may occur.", source_row=11,
            source_filename="private.xlsx", source_sha256="b" * 64,
        )
        RequirementRiskMapping.objects.create(
            requirement=self.control.requirement, risk=catalog,
            rationale="The failed control directly restricts unauthorized access.",
            review_status=RequirementRiskMapping.ReviewStatus.APPROVED,
            proposed_by=self.admin, reviewed_by=self.admin, reviewed_at=timezone.now(),
        )
        self.client.login(username="risk-assessor", password="test-password")
        list_url = reverse("risk-register-list", args=(self.organization.slug, self.assessment.id))
        listing = self.client.get(list_url)
        self.assertContains(listing, "Approved catalog risks associated with findings")
        self.assertContains(listing, "R-AC-4")
        self.assertContains(listing, "Evaluate and register")
        create_url = reverse("risk-register-create", args=(self.organization.slug, self.assessment.id))
        guided = self.client.get(create_url, {"control": self.control.id, "catalog": catalog.id})
        self.assertContains(guided, "Unauthorized access")
        payload = self._risk_payload(
            catalog_risk=str(catalog.id), title=catalog.title,
            description=catalog.description, category=catalog.grouping,
        )
        self.assertEqual(self.client.post(create_url, payload).status_code, 302)
        self.assertEqual(RiskRegisterEntry.objects.count(), 1)
        duplicate = self.client.post(create_url, payload)
        self.assertEqual(duplicate.status_code, 200)
        self.assertContains(duplicate, "already registered")
        self.assertEqual(RiskRegisterEntry.objects.count(), 1)
        self.assertContains(self.client.get(list_url), "Registered")


class SprintEighteenPointTwoTests(SprintEighteenRiskRegisterTests):
    def _create_risk(self, **overrides):
        values = {
            "organization": self.organization, "system": self.assessment.system,
            "assessment": self.assessment, "risk_id": "RISK-0001",
            "title": "Operational risk", "description": "A material event may occur.",
            "category": "Governance", "created_by": self.assessor,
            "owner": self.assessor_membership, "likelihood": 3, "impact": 3,
            "treatment": RiskRegisterEntry.Treatment.MITIGATE,
            "treatment_plan": "Implement and validate safeguards.",
        }
        values.update(overrides)
        return RiskRegisterEntry.objects.create(**values)

    def test_risk_workflow_is_optional_hidden_and_non_destructive(self):
        risk = self._create_risk()
        self.assessment.risk_management_enabled = False
        self.assessment.include_risk_in_reports = False
        self.assessment.save(update_fields=("risk_management_enabled", "include_risk_in_reports"))
        self.client.login(username="risk-assessor", password="test-password")
        dashboard = self.client.get(reverse(
            "assessment-dashboard", args=(self.organization.slug, self.assessment.id)
        ))
        self.assertNotContains(dashboard, "Organizational risk register")
        self.assertEqual(self.client.get(reverse(
            "risk-register-list", args=(self.organization.slug, self.assessment.id)
        )).status_code, 404)
        self.assertTrue(RiskRegisterEntry.objects.filter(pk=risk.pk).exists())

    def test_treatment_action_and_evidence_backed_residual_reassessment(self):
        risk = self._create_risk()
        evidence = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Treatment validation", external_reference="https://example.test/risk-evidence",
            uploaded_by=self.assessor,
        )
        self.client.login(username="risk-assessor", password="test-password")
        action_url = reverse("risk-treatment-action-create", args=(
            self.organization.slug, self.assessment.id, risk.id
        ))
        response = self.client.post(action_url, {
            "title": "Deploy safeguard", "description": "Deploy and test the safeguard.",
            "owner": self.assessor_membership.id, "status": "COMPLETE", "priority": "HIGH",
            "planned_start": "2026-09-01", "due_date": "2026-09-30",
            "completed_date": "2026-09-20", "completion_notes": "Safeguard tested successfully.",
            "remediation_plan": "", "evidence": [evidence.id], "dependencies": [],
        })
        self.assertEqual(response.status_code, 302)
        action = RiskTreatmentAction.objects.get(risk=risk)
        self.assertEqual(action.status, RiskTreatmentAction.Status.COMPLETE)
        reassess_url = reverse("risk-reassess", args=(self.organization.slug, self.assessment.id, risk.id))
        response = self.client.post(reassess_url, {
            "new_likelihood": 2, "new_impact": 2,
            "rationale": "Validated safeguards reduced probability and impact.",
            "evidence": [evidence.id],
        })
        self.assertEqual(response.status_code, 302)
        risk.refresh_from_db()
        self.assertEqual(risk.residual_score, 4)
        self.assertIsNotNone(risk.last_reviewed_at)
        self.assertTrue(RiskReassessment.objects.filter(risk=risk, evidence=evidence).exists())

    def test_acceptance_requires_request_policy_and_admin_review(self):
        risk = self._create_risk()
        policy = RiskTolerancePolicy.objects.create(
            organization=self.organization, updated_by=self.admin,
            maximum_residual_score=11, maximum_acceptance_days=90,
        )
        expiration = date.today() + timedelta(days=30)
        self.client.login(username="risk-assessor", password="test-password")
        request_url = reverse("risk-acceptance-request", args=(
            self.organization.slug, self.assessment.id, risk.id
        ))
        self.assertEqual(self.client.post(request_url, {
            "rationale": "Temporary acceptance while replacement is procured.",
            "requested_expiration": expiration.isoformat(),
        }).status_code, 302)
        acceptance = RiskAcceptanceRequest.objects.get(risk=risk)
        self.assertEqual(acceptance.status, RiskAcceptanceRequest.Status.PENDING)
        self.client.login(username="risk-owner-admin", password="test-password")
        review_url = reverse("risk-acceptance-review", args=(
            self.organization.slug, self.assessment.id, acceptance.id
        ))
        self.assertEqual(self.client.post(review_url, {"action": "approve"}).status_code, 302)
        risk.refresh_from_db(); acceptance.refresh_from_db()
        self.assertEqual(acceptance.status, RiskAcceptanceRequest.Status.APPROVED)
        self.assertEqual(risk.status, RiskRegisterEntry.Status.ACCEPTED)
        self.assertEqual(risk.acceptance_expires, expiration)
        self.assertEqual(policy.maximum_residual_score, 11)

    def test_closure_requires_completed_actions_residual_score_and_evidence_then_reopens(self):
        risk = self._create_risk(residual_likelihood=1, residual_impact=2)
        evidence = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Closure evidence", external_reference="https://example.test/closure",
            uploaded_by=self.assessor,
        )
        risk.supporting_evidence.add(evidence)
        RiskTreatmentAction.objects.create(
            risk=risk, title="Complete treatment", status=RiskTreatmentAction.Status.COMPLETE,
            completed_date=date.today(), completion_notes="Validated.", created_by=self.assessor,
        )
        self.client.login(username="risk-assessor", password="test-password")
        close_url = reverse("risk-close", args=(self.organization.slug, self.assessment.id, risk.id))
        self.assertEqual(self.client.post(close_url, {
            "closure_rationale": "Validated evidence demonstrates treatment completion."
        }).status_code, 302)
        risk.refresh_from_db(); self.assertEqual(risk.status, RiskRegisterEntry.Status.CLOSED)
        self.client.login(username="risk-owner-admin", password="test-password")
        reopen_url = reverse("risk-reopen", args=(self.organization.slug, self.assessment.id, risk.id))
        self.assertEqual(self.client.post(reopen_url, {
            "reason": "A material trigger event requires a new assessment."
        }).status_code, 302)
        risk.refresh_from_db(); self.assertEqual(risk.status, RiskRegisterEntry.Status.MONITORING)
        self.assertTrue(RiskRegisterHistory.objects.filter(risk=risk, action="REOPENED").exists())


class SprintSeventeenPointSixAuthoritativeSourcesTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_superuser(
            "source-admin", "admin@example.com", "test-password"
        )
        self.regular = user_model.objects.create_user(
            "source-user", password="test-password"
        )

    def _workbook(self, directory):
        path = Path(directory) / "authoritative-sources.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Focal Documents"
        sheet.append(["Locale", "CCF Column Header", "ADI", "Source", "Focal Document Name", "Authoritative Documents URL"])
        sheet.append(["Global", "ISO 27001", "iso-27001", "ISO", "ISO/IEC 27001", "https://www.iso.org/standard/27001.html"])
        sheet.append(["US", "NIST CSF", "duplicate-id", "NIST", "NIST CSF", "https://www.nist.gov/cyberframework"])
        sheet.append(["US", "NIST SP 800-53", "duplicate-id", "NIST", "NIST SP 800-53", "https://csrc.nist.gov/publications/detail/sp/800-53/rev-5/final"])
        sheet.append(["US", "Invalid Source", "invalid-source", "Example", "Invalid publication", "not a URL"])
        sheet.append(["Not Complete", "", "", "", "Footer", ""])
        workbook.save(path)
        workbook.close()
        return path

    def test_private_import_preserves_provenance_and_quality(self):
        with tempfile.TemporaryDirectory() as directory:
            report = import_authoritative_sources(self._workbook(directory), apply=True)

        self.assertEqual(report["records"], 4)
        self.assertEqual(report["valid"], 1)
        self.assertEqual(report["duplicate_adi"], 2)
        self.assertEqual(report["invalid_url"], 1)
        self.assertEqual(report["excluded_rows"], 1)
        self.assertEqual(AuthoritativeDocument.objects.count(), 4)
        invalid = AuthoritativeDocument.objects.get(authoritative_document_id="invalid-source")
        self.assertEqual(invalid.official_url, "")
        self.assertEqual(invalid.source_url_text, "not a URL")
        self.assertEqual(invalid.source_row, 5)
        self.assertEqual(invalid.authority.canonical_name, "Invalid Source")
        self.assertEqual(len(invalid.source_sha256), 64)

    def test_registry_is_restricted_to_superusers(self):
        self.client.force_login(self.regular)
        self.assertEqual(self.client.get(reverse("authoritative-source-registry")).status_code, 404)
        self.client.force_login(self.admin)
        self.assertEqual(self.client.get(reverse("authoritative-source-registry")).status_code, 200)


class SprintOneWorkflowTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.assessor = user_model.objects.create_user(
            "assessor", password="test-password"
        )
        self.outsider = user_model.objects.create_user(
            "outsider", password="test-password"
        )
        self.client_org = Organization.objects.create(name="Acme Defense", slug="acme")
        self.other_org = Organization.objects.create(name="Other Client", slug="other")
        Membership.objects.create(
            user=self.assessor,
            organization=self.client_org,
            role=Membership.Role.ASSESSOR,
        )
        Membership.objects.create(
            user=self.outsider,
            organization=self.other_org,
            role=Membership.Role.ASSESSOR,
        )
        self.system = System.objects.create(
            organization=self.client_org, name="CUI Enclave"
        )
        self.framework = Framework.objects.create(
            code="CMMC-L2", name="CMMC Level 2", version="2.13",
            scoring_method=Framework.ScoringMethod.SPRS, maximum_score=110,
        )
        self.requirement = Requirement.objects.create(
            framework=self.framework,
            requirement_id="AC.L2-3.1.1",
            domain="AC",
            title="Authorized Access Control",
            statement="Limit access to authorized users.",
            full_deduction=5,
        )
        self.assessment = Assessment.objects.create(
            system=self.system,
            framework=self.framework,
            name="2026 CMMC Assessment",
            created_by=self.assessor,
            status=Assessment.Status.IN_PROGRESS,
        )
        self.result = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement
        )

    def test_login_is_required(self):
        response = self.client.get(reverse("organization-list"))
        self.assertRedirects(response, f"{reverse('login')}?next=/")

    def test_tenant_membership_limits_organization_access(self):
        self.client.login(username="assessor", password="test-password")
        response = self.client.get(reverse("organization-list"))
        self.assertContains(response, "Acme Defense")
        self.assertNotContains(response, "Other Client")
        response = self.client.get(reverse("system-list", args=("acme",)))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.system.name)
        response = self.client.get(reverse("system-list", args=("other",)))
        self.assertEqual(response.status_code, 404)

    def test_create_assessment_loads_framework_controls(self):
        self.client.login(username="assessor", password="test-password")
        response = self.client.post(
            reverse("assessment-create", args=("acme", self.system.id)),
            {
                "frameworks": [self.framework.id],
                "primary_framework": self.framework.id,
                "name": "New Assessment",
            },
        )
        created = Assessment.objects.get(name="New Assessment")
        self.assertRedirects(
            response, reverse("assessment-dashboard", args=("acme", created.id))
        )
        self.assertEqual(created.control_results.count(), 1)

    def test_score_control_updates_dashboard_and_audit_log(self):
        self.client.login(username="assessor", password="test-password")
        response = self.client.post(
            reverse("control-edit", args=("acme", self.assessment.id, self.result.id)),
            {
                "status": ControlAssessment.Status.NOT_MET,
                "implementation_state": ControlAssessment.Implementation.NONE,
                "assessor_notes_findings": "Authorized-user evidence was incomplete.",
                "control_owner": "Security Officer",
                "ssp_reference": "SSP 3.1.1",
            },
            follow=True,
        )
        self.assertEqual(response.status_code, 200)
        self.result.refresh_from_db()
        self.assertEqual(self.result.calculated_deduction, 5)
        self.assertContains(response, "105")
        self.assertContains(response, "Authorized Access Control")
        self.assertTrue(
            AuditEvent.objects.filter(action="control_assessment.updated").exists()
        )

    def test_read_only_member_cannot_edit(self):
        viewer = get_user_model().objects.create_user(
            "viewer", password="test-password"
        )
        Membership.objects.create(
            user=viewer, organization=self.client_org, role=Membership.Role.VIEWER
        )
        self.client.login(username="viewer", password="test-password")
        response = self.client.get(
            reverse("control-edit", args=("acme", self.assessment.id, self.result.id))
        )
        self.assertEqual(response.status_code, 404)


class SprintTwoOnboardingTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin_user = user_model.objects.create_user(
            "orgadmin", email="admin@example.com", password="test-password"
        )
        self.owner_user = user_model.objects.create_user(
            "controlowner", email="owner@example.com", password="test-password"
        )
        self.client.login(username="orgadmin", password="test-password")

    def test_user_can_create_first_organization_and_becomes_admin(self):
        response = self.client.post(reverse("organization-create"), {
            "name": "Acme Defense", "legal_name": "Acme Defense LLC",
            "kind": Organization.Kind.CLIENT, "industry": "Defense",
            "primary_contact_email": "ciso@example.com",
        })
        organization = Organization.objects.get(name="Acme Defense")
        self.assertRedirects(response, reverse("system-list", args=(organization.slug,)))
        self.assertTrue(Membership.objects.filter(
            user=self.admin_user, organization=organization,
            role=Membership.Role.ADMIN,
        ).exists())

    def test_admin_must_confirm_assessment_name_before_permanent_delete(self):
        organization = Organization.objects.create(name="Delete Test", slug="delete-test")
        Membership.objects.create(
            user=self.admin_user, organization=organization, role=Membership.Role.ADMIN
        )
        system = System.objects.create(organization=organization, name="Synthetic System")
        framework = Framework.objects.create(code="DELETE-TEST", name="Delete Test", version="1")
        assessment = Assessment.objects.create(
            system=system, framework=framework, name="Synthetic Assessment",
            created_by=self.admin_user,
        )
        url = reverse("assessment-delete", args=(organization.slug, assessment.id))
        dashboard = self.client.get(reverse(
            "assessment-dashboard", args=(organization.slug, assessment.id)
        ))
        self.assertContains(dashboard, "Delete assessment")
        response = self.client.get(url)
        self.assertContains(response, "This action cannot be undone")
        response = self.client.post(url, {"confirmation": "wrong name"})
        self.assertContains(response, "Nothing was deleted")
        self.assertTrue(Assessment.objects.filter(id=assessment.id).exists())
        response = self.client.post(url, {"confirmation": assessment.name}, follow=True)
        self.assertContains(response, "permanently deleted")
        self.assertFalse(Assessment.objects.filter(id=assessment.id).exists())
        self.assertTrue(AuditEvent.objects.filter(
            organization=organization, action="assessment.deleted",
            object_id=str(assessment.id),
        ).exists())

    def test_non_admin_cannot_delete_assessment(self):
        organization = Organization.objects.create(name="Restricted", slug="restricted-delete")
        Membership.objects.create(
            user=self.admin_user, organization=organization, role=Membership.Role.ASSESSOR
        )
        system = System.objects.create(organization=organization, name="Synthetic System")
        framework = Framework.objects.create(code="NO-DELETE", name="No Delete", version="1")
        assessment = Assessment.objects.create(
            system=system, framework=framework, name="Protected Assessment",
            created_by=self.admin_user,
        )
        dashboard = self.client.get(reverse(
            "assessment-dashboard", args=(organization.slug, assessment.id)
        ))
        self.assertNotContains(dashboard, "Delete assessment")
        response = self.client.post(
            reverse("assessment-delete", args=(organization.slug, assessment.id)),
            {"confirmation": assessment.name},
        )
        self.assertEqual(response.status_code, 404)
        self.assertTrue(Assessment.objects.filter(id=assessment.id).exists())

    @override_settings(OMNI_EMAIL_ENABLED=True)
    def test_notification_test_is_an_explicit_submit_button(self):
        organization = Organization.objects.create(name="Notifications", slug="notifications-ui")
        Membership.objects.create(
            user=self.admin_user, organization=organization, role=Membership.Role.ADMIN
        )
        response = self.client.get(reverse("notification-policy", args=(organization.slug,)))
        self.assertContains(
            response,
            '<button class="secondary" type="submit">Send test email to me</button>',
            html=True,
        )

    def test_admin_can_onboard_system_and_existing_member(self):
        organization = Organization.objects.create(name="Acme", slug="acme")
        Membership.objects.create(
            user=self.admin_user, organization=organization, role=Membership.Role.ADMIN
        )
        response = self.client.post(reverse("system-create", args=("acme",)), {
            "name": "CUI Enclave", "description": "Protected environment",
            "system_owner_name": "Alex Owner", "system_owner_email": "alex@example.com",
            "environment": "Production", "data_types": "CUI", "scope": "In scope",
        })
        system = System.objects.get(name="CUI Enclave")
        self.assertRedirects(response, reverse("assessment-list", args=("acme", system.id)))
        response = self.client.post(reverse("membership-list", args=("acme",)), {
            "username_or_email": "owner@example.com", "role": Membership.Role.CLIENT,
        })
        self.assertRedirects(response, reverse("membership-list", args=("acme",)))
        self.assertTrue(Membership.objects.filter(
            user=self.owner_user, organization=organization,
            role=Membership.Role.CLIENT,
        ).exists())

    def test_bulk_owner_assignment_is_limited_to_organization_members(self):
        organization = Organization.objects.create(name="Acme", slug="acme")
        admin_membership = Membership.objects.create(
            user=self.admin_user, organization=organization, role=Membership.Role.ADMIN
        )
        owner_membership = Membership.objects.create(
            user=self.owner_user, organization=organization, role=Membership.Role.CLIENT
        )
        system = System.objects.create(organization=organization, name="Enclave")
        framework = Framework.objects.create(code="TEST", name="Test", version="1")
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="AC.1", domain="AC",
            title="Access", statement="Control access.", full_deduction=1,
        )
        assessment = Assessment.objects.create(
            system=system, framework=framework, name="Assessment",
            created_by=self.admin_user,
        )
        result = ControlAssessment.objects.create(
            assessment=assessment, requirement=requirement
        )
        response = self.client.post(
            reverse("bulk-control-owners", args=("acme", assessment.id)),
            {"domain": "AC", "primary_owner": owner_membership.id,
             "supporting_owners": [admin_membership.id]},
        )
        self.assertRedirects(
            response, reverse("assessment-dashboard", args=("acme", assessment.id))
        )
        result.refresh_from_db()
        self.assertEqual(result.primary_owner, owner_membership)
        self.assertEqual(list(result.supporting_owners.all()), [admin_membership])


class SprintThreeEvidenceTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.assessor = user_model.objects.create_user(
            "evidence-assessor", password="test-password"
        )
        self.client_user = user_model.objects.create_user(
            "evidence-client", password="test-password"
        )
        self.outsider = user_model.objects.create_user(
            "evidence-outsider", password="test-password"
        )
        self.organization = Organization.objects.create(name="Acme", slug="acme")
        self.other_org = Organization.objects.create(name="Other", slug="other-evidence")
        self.assessor_membership = Membership.objects.create(
            user=self.assessor, organization=self.organization,
            role=Membership.Role.ASSESSOR,
        )
        self.client_membership = Membership.objects.create(
            user=self.client_user, organization=self.organization,
            role=Membership.Role.CLIENT,
        )
        Membership.objects.create(
            user=self.outsider, organization=self.other_org,
            role=Membership.Role.ASSESSOR,
        )
        self.system = System.objects.create(
            organization=self.organization, name="CUI Enclave"
        )
        self.framework = Framework.objects.create(
            code="EVIDENCE-TEST", name="Evidence Test", version="1"
        )
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="AC.1", domain="AC",
            title="Access Control", statement="Control access.", full_deduction=1,
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Evidence Assessment",
            created_by=self.assessor, status=Assessment.Status.IN_PROGRESS,
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement
        )

    def test_curated_request_populates_canonical_object_and_control_mapping(self):
        self.client.login(username="evidence-assessor", password="test-password")
        response = self.client.post(
            reverse("evidence-request-create", args=("acme", self.assessment.id)),
            {
                "catalog_object": "EV-0001", "title": "", "description": "",
                "status": EvidenceRequest.Status.REQUESTED,
                "owner": self.client_membership.id, "controls": [self.control.id],
            },
        )
        request_item = EvidenceRequest.objects.get(assessment=self.assessment)
        self.assertRedirects(
            response, reverse("evidence-list", args=("acme", self.assessment.id))
        )
        self.assertEqual(request_item.evidence_code, "EV-0001")
        self.assertEqual(request_item.title, "Security Plan")
        self.assertEqual(list(request_item.controls.all()), [self.control])

    def test_generate_request_list_imports_optimized_drl_idempotently(self):
        self.framework.code = "CMMC-TEST"
        self.framework.save(update_fields=("code",))
        generated = SimpleNamespace(
            requests=[SimpleNamespace(
                requested_item="Access Control Policy",
                description="Provide the current policy.",
                controls=[SimpleNamespace(control_id="AC.1")],
            )]
        )
        self.client.login(username="evidence-assessor", password="test-password")
        url = reverse("evidence-request-generate", args=("acme", self.assessment.id))
        with patch("webapp.views._generate_cmmc_drl", return_value=generated):
            self.client.post(url)
            self.client.post(url)
        self.assertEqual(
            EvidenceRequest.objects.filter(title="Access Control Policy").count(), 1
        )
        item = EvidenceRequest.objects.get(title="Access Control Policy")
        self.assertEqual(list(item.controls.all()), [self.control])

    def test_client_can_register_artifact_and_request_becomes_received(self):
        request_item = EvidenceRequest.objects.create(
            assessment=self.assessment, title="Security Plan",
            created_by=self.assessor, owner=self.client_membership,
        )
        self.client.login(username="evidence-client", password="test-password")
        response = self.client.post(
            reverse("evidence-artifact-create", args=("acme", self.assessment.id)),
            {
                "title": "Current SSP", "external_reference": "https://example.com/ssp",
                "source": "GRC repository", "review_status": "ACCEPTED",
                "requests": [request_item.id], "controls": [self.control.id],
            },
        )
        self.assertRedirects(
            response, reverse("evidence-list", args=("acme", self.assessment.id))
        )
        artifact = EvidenceArtifact.objects.get(title="Current SSP")
        self.assertEqual(artifact.review_status, EvidenceArtifact.ReviewStatus.RECEIVED)
        request_item.refresh_from_db()
        self.assertEqual(request_item.status, EvidenceRequest.Status.RECEIVED)

    def test_assessor_acceptance_updates_request_and_dashboard_readiness(self):
        request_item = EvidenceRequest.objects.create(
            assessment=self.assessment, title="Security Plan", created_by=self.assessor
        )
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="SSP", external_reference="https://example.com/ssp",
            uploaded_by=self.client_user,
        )
        artifact.requests.add(request_item)
        self.client.login(username="evidence-assessor", password="test-password")
        response = self.client.post(
            reverse("evidence-artifact-edit", args=(
                "acme", self.assessment.id, artifact.id,
            )),
            {
                "title": "SSP", "external_reference": "https://example.com/ssp",
                "review_status": EvidenceArtifact.ReviewStatus.ACCEPTED,
                "assessor_notes": "Accepted and current.",
                "requests": [request_item.id], "controls": [self.control.id],
            },
        )
        self.assertRedirects(
            response, reverse("evidence-list", args=("acme", self.assessment.id))
        )
        request_item.refresh_from_db()
        self.assertEqual(request_item.status, EvidenceRequest.Status.ACCEPTED)
        dashboard = self.client.get(
            reverse("assessment-dashboard", args=("acme", self.assessment.id))
        )
        self.assertContains(dashboard, "100.0%")

    def test_cross_tenant_user_cannot_download_private_evidence(self):
        with tempfile.TemporaryDirectory() as media_root:
            with override_settings(MEDIA_ROOT=media_root):
                artifact = EvidenceArtifact.objects.create(
                    organization=self.organization, assessment=self.assessment,
                    title="Private Evidence",
                    file=SimpleUploadedFile("evidence.txt", b"private"),
                    uploaded_by=self.assessor,
                )
                self.client.login(username="evidence-outsider", password="test-password")
                response = self.client.get(
                    reverse("evidence-artifact-download", args=(
                        "acme", self.assessment.id, artifact.id,
                    ))
                )
                self.assertEqual(response.status_code, 404)


class SprintFourRemediationTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.admin = users.objects.create_user("rem-admin", password="test-password")
        self.assessor = users.objects.create_user("rem-assessor", password="test-password")
        self.owner = users.objects.create_user("rem-owner", password="test-password")
        self.outsider = users.objects.create_user("rem-outsider", password="test-password")
        self.organization = Organization.objects.create(name="Acme Remediation", slug="acme-rem")
        self.other_org = Organization.objects.create(name="Other Remediation", slug="other-rem")
        self.admin_member = Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.assessor_member = Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.owner_member = Membership.objects.create(
            user=self.owner, organization=self.organization, role=Membership.Role.CLIENT
        )
        Membership.objects.create(
            user=self.outsider, organization=self.other_org, role=Membership.Role.ADMIN
        )
        self.system = System.objects.create(organization=self.organization, name="Enclave")
        self.framework = Framework.objects.create(code="REM", name="Remediation", version="1")
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="AC.1", domain="AC",
            title="Access Control", statement="Control access.", full_deduction=5,
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Remediation Assessment",
            created_by=self.assessor,
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement,
            status=ControlAssessment.Status.NOT_MET,
            assessor_notes_findings="Access reviews were not performed.",
        )

    def _plan(self, **overrides):
        values = {
            "assessment": self.assessment, "remediation_id": "RAP-0001",
            "title": "Restore access reviews",
            "weakness_description": "Access reviews were not performed.",
            "owner": self.owner_member, "date_identified": date.today(),
            "planned_completion": date.today() + timedelta(days=30),
            "created_by": self.assessor,
        }
        values.update(overrides)
        plan = RemediationPlan.objects.create(**values)
        plan.controls.add(self.control)
        return plan

    def _post_values(self, plan, **overrides):
        values = {
            "title": plan.title, "controls": [self.control.id],
            "weakness_description": plan.weakness_description,
            "root_cause": plan.root_cause, "corrective_action": plan.corrective_action,
            "compensating_controls": plan.compensating_controls,
            "closure_criteria": plan.closure_criteria, "owner": self.owner_member.id,
            "status": plan.status, "priority": plan.priority, "severity": plan.severity,
            "likelihood": plan.likelihood, "residual_risk": plan.residual_risk,
            "date_identified": plan.date_identified.isoformat(),
            "planned_completion": plan.planned_completion.isoformat(),
            "validation_status": plan.validation_status,
        }
        values.update(overrides)
        return values

    def test_create_from_finding_prefills_control_and_finding(self):
        self.client.login(username="rem-assessor", password="test-password")
        response = self.client.get(
            reverse("remediation-create", args=("acme-rem", self.assessment.id))
            + f"?control={self.control.id}"
        )
        self.assertContains(response, "Access reviews were not performed.")
        self.assertContains(response, "Remediate AC.1")

    def test_assigned_owner_can_update_but_cannot_self_validate(self):
        plan = self._plan(root_cause="Missing schedule", corrective_action="Perform reviews")
        self.client.login(username="rem-owner", password="test-password")
        response = self.client.post(
            reverse("remediation-edit", args=("acme-rem", self.assessment.id, plan.id)),
            self._post_values(plan, corrective_action="Perform quarterly reviews",
                              validation_status=RemediationPlan.ValidationStatus.VALIDATED),
        )
        self.assertRedirects(response, reverse(
            "remediation-detail", args=("acme-rem", self.assessment.id, plan.id)
        ))
        plan.refresh_from_db()
        self.assertEqual(plan.corrective_action, "Perform quarterly reviews")
        self.assertEqual(plan.validation_status, RemediationPlan.ValidationStatus.PENDING)

    def test_validated_plan_with_evidence_can_be_closed(self):
        plan = self._plan(
            root_cause="Missing schedule", corrective_action="Perform reviews",
            closure_criteria="Review approved and evidenced",
            status=RemediationPlan.Status.READY_VALIDATION,
        )
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Completed review", external_reference="https://example.com/review",
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED, uploaded_by=self.owner,
        )
        self.client.login(username="rem-assessor", password="test-password")
        response = self.client.post(
            reverse("remediation-edit", args=("acme-rem", self.assessment.id, plan.id)),
            self._post_values(
                plan, status=RemediationPlan.Status.CLOSED,
                actual_completion=date.today().isoformat(),
                closure_evidence=[artifact.id],
                validation_status=RemediationPlan.ValidationStatus.VALIDATED,
                validation_notes="Closure verified.",
            ),
        )
        self.assertRedirects(response, reverse(
            "remediation-detail", args=("acme-rem", self.assessment.id, plan.id)
        ))
        plan.refresh_from_db()
        self.assertEqual(plan.status, RemediationPlan.Status.CLOSED)
        self.assertEqual(plan.validated_by, self.assessor)

    def test_cross_tenant_remediation_access_is_denied(self):
        plan = self._plan()
        self.client.login(username="rem-outsider", password="test-password")
        response = self.client.get(reverse(
            "remediation-detail", args=("acme-rem", self.assessment.id, plan.id)
        ))
        self.assertEqual(response.status_code, 404)

    def test_export_uses_workbook_poam_columns(self):
        self._plan(root_cause="Cause", corrective_action="Action")
        self.client.login(username="rem-assessor", password="test-password")
        response = self.client.get(reverse(
            "remediation-export", args=("acme-rem", self.assessment.id)
        ))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Disposition"],
            f'attachment; filename="Omni-{self.assessment.id}-Remediation-Action-Plan.xlsx"',
        )
        content = b"".join(response.streaming_content)
        sheet = load_workbook(BytesIO(content), read_only=True).active
        self.assertEqual(sheet.cell(5, 1).value, "Remediation ID (POA&M ID)")
        self.assertEqual(sheet.cell(5, 24).value, "Assessor Notes")
        self.assertEqual(sheet.cell(6, 2).value, "AC.1")


class SprintFiveReportingTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.assessor = users.objects.create_user("report-assessor", password="test-password")
        self.outsider = users.objects.create_user("report-outsider", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic Defense", slug="synthetic")
        self.other_org = Organization.objects.create(name="Other Synthetic", slug="other-synthetic")
        Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        Membership.objects.create(
            user=self.outsider, organization=self.other_org, role=Membership.Role.ASSESSOR
        )
        self.system = System.objects.create(
            organization=self.organization, name="Synthetic Enclave", scope="CUI enclave",
            cage_code="1AB23", system_owner_name="Synthetic Owner",
        )
        self.framework = Framework.objects.create(code="CMMC-REPORT", name="CMMC Level 2", version="2.13")
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="AC.L2-3.1.1", domain="AC",
            title="Authorized Access Control", statement="Limit access.", full_deduction=5,
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Synthetic Assessment",
            created_by=self.assessor, status=Assessment.Status.COMPLETE,
        )
        self.result = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement,
            status=ControlAssessment.Status.MET,
            implementation_state=ControlAssessment.Implementation.FULL,
            assessor_notes_findings="Access is restricted to authorized users.",
            ssp_reference="SSP 3.1.1",
        )

    def test_workbook_binds_web_assessment_data(self):
        from .reporting import build_assessment_workbook
        content = build_assessment_workbook(self.assessment)
        workbook = load_workbook(BytesIO(content), data_only=False, read_only=True)
        try:
            cover = workbook["Cover"]
            self.assertEqual(cover["C6"].value, "Synthetic Defense")
            assessment = workbook["Assessment"]
            self.assertEqual(assessment["I6"].value, "MET")
            self.assertEqual(assessment["R6"].value, "Access is restricted to authorized users.")
            crosswalk = workbook["SSP Crosswalk"]
            self.assertEqual(crosswalk["C6"].value, "SSP 3.1.1")
        finally:
            workbook.close()

    def test_package_embeds_upload_and_external_reference_text(self):
        from .reporting import build_package
        with tempfile.TemporaryDirectory() as directory:
            template = f"{directory}/private-template.docx"
            Path(template).touch()
            with override_settings(MEDIA_ROOT=directory, OMNI_SSP_TEMPLATE=template):
                artifact = EvidenceArtifact.objects.create(
                    organization=self.organization, assessment=self.assessment,
                    title="Access Review Evidence",
                    file=SimpleUploadedFile("access-review.txt", b"synthetic evidence"),
                    external_reference="https://example.com/access-review",
                    source="Synthetic repository",
                    review_status=EvidenceArtifact.ReviewStatus.ACCEPTED,
                    uploaded_by=self.assessor,
                )
                artifact.controls.add(self.result)
                with patch("webapp.reporting.build_word_ssp", return_value=b"synthetic docx"):
                    content, readiness = build_package(self.assessment, self.assessor)
            self.assertTrue(readiness["ready"])
            with zipfile.ZipFile(BytesIO(content)) as package:
                names = package.namelist()
                self.assertIn("Deliverables/Omni-Assessment-Workbook.xlsx", names)
                self.assertIn("Evidence/Evidence-Index.csv", names)
                upload = next(name for name in names if name.endswith("access-review.txt"))
                reference = next(name for name in names if name.endswith("External_Reference.txt"))
                self.assertEqual(package.read(upload), b"synthetic evidence")
                reference_text = package.read(reference).decode()
                self.assertIn("https://example.com/access-review", reference_text)
                self.assertIn("AC.L2-3.1.1", reference_text)

    def test_risk_register_is_packaged_only_when_both_toggles_are_enabled(self):
        from .reporting import build_package
        RiskRegisterEntry.objects.create(
            organization=self.organization, system=self.system, assessment=self.assessment,
            risk_id="RISK-0001", title="Synthetic risk", description="Synthetic exposure.",
            category="Access Control", created_by=self.assessor,
        )
        with tempfile.TemporaryDirectory() as directory:
            template = f"{directory}/private-template.docx"
            Path(template).touch()
            with override_settings(OMNI_SSP_TEMPLATE=template), patch(
                "webapp.reporting.build_word_ssp", return_value=b"synthetic docx"
            ):
                without_risk, _ = build_package(self.assessment, self.assessor)
                self.assessment.risk_management_enabled = True
                self.assessment.include_risk_in_reports = True
                self.assessment.save(update_fields=("risk_management_enabled", "include_risk_in_reports"))
                with_risk, _ = build_package(self.assessment, self.assessor)
        with zipfile.ZipFile(BytesIO(without_risk)) as package:
            self.assertNotIn("Deliverables/Omni-Risk-Register.csv", package.namelist())
        with zipfile.ZipFile(BytesIO(with_risk)) as package:
            self.assertIn("Deliverables/Omni-Risk-Register.csv", package.namelist())
            self.assertIn(b"RISK-0001", package.read("Deliverables/Omni-Risk-Register.csv"))

    def test_download_creates_metadata_history_without_persisting_client_file(self):
        self.client.login(username="report-assessor", password="test-password")
        with patch("webapp.reporting.build_assessment_workbook", return_value=b"workbook"):
            response = self.client.get(reverse(
                "report-download", args=("synthetic", self.assessment.id, "WORKBOOK")
            ))
        self.assertEqual(response.content, b"workbook")
        record = GeneratedDocument.objects.get(assessment=self.assessment)
        self.assertEqual(record.kind, GeneratedDocument.Kind.WORKBOOK)
        self.assertEqual(record.size_bytes, 8)
        self.assertFalse(hasattr(record, "file"))

    def test_cross_tenant_report_center_is_denied(self):
        self.client.login(username="report-outsider", password="test-password")
        response = self.client.get(reverse(
            "report-center", args=("synthetic", self.assessment.id)
        ))
        self.assertEqual(response.status_code, 404)


class SprintFivePointFiveFrameworkTests(TestCase):
    def setUp(self):
        user = get_user_model()
        self.assessor = user.objects.create_user("multi-assessor", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic Multi", slug="multi")
        Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.system = System.objects.create(organization=self.organization, name="Shared System")
        self.cmmc = Framework.objects.create(
            code="CMMC-MULTI", name="CMMC Level 2", version="2.13",
            authority="DoD", scoring_method=Framework.ScoringMethod.SPRS,
            maximum_score=110,
        )
        self.iso = Framework.objects.create(
            code="ISO-MULTI", name="ISO Synthetic", version="2026",
            authority="ISO", scoring_method=Framework.ScoringMethod.NONE,
        )
        self.cmmc_req = Requirement.objects.create(
            framework=self.cmmc, requirement_id="AC.1", domain="AC",
            title="CMMC Access", statement="Restrict access.", full_deduction=5,
        )
        self.iso_req = Requirement.objects.create(
            framework=self.iso, requirement_id="AC.1", domain="Access",
            title="ISO Access", statement="Manage access.", full_deduction=1,
        )
        self.client.login(username="multi-assessor", password="test-password")

    def test_create_selects_multiple_frameworks_and_loads_native_requirements(self):
        response = self.client.post(
            reverse("assessment-create", args=("multi", self.system.id)),
            {"name": "Integrated Assessment", "frameworks": [self.cmmc.id, self.iso.id],
             "primary_framework": self.cmmc.id},
        )
        assessment = Assessment.objects.get(name="Integrated Assessment")
        self.assertRedirects(response, reverse(
            "assessment-dashboard", args=("multi", assessment.id)
        ))
        self.assertEqual(assessment.framework, self.cmmc)
        self.assertEqual(set(assessment.frameworks.all()), {self.cmmc, self.iso})
        self.assertEqual(assessment.control_results.count(), 2)
        self.assertEqual(
            assessment.control_results.filter(requirement__requirement_id="AC.1").count(), 2
        )

    def test_framework_scores_are_independent_and_non_scored_framework_has_no_score(self):
        assessment = Assessment.objects.create(
            system=self.system, framework=self.cmmc, name="Scored Multi",
            created_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.cmmc, is_primary=True,
            added_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.iso, added_by=self.assessor,
        )
        ControlAssessment.objects.create(
            assessment=assessment, requirement=self.cmmc_req,
            status=ControlAssessment.Status.NOT_MET,
        )
        ControlAssessment.objects.create(
            assessment=assessment, requirement=self.iso_req,
            status=ControlAssessment.Status.MET,
        )
        response = self.client.get(reverse(
            "assessment-dashboard", args=("multi", assessment.id)
        ))
        self.assertContains(response, "105")
        self.assertContains(response, "ISO-MULTI")
        self.assertContains(response, "1/1 assessed")

    def test_framework_with_recorded_work_cannot_be_removed(self):
        assessment = Assessment.objects.create(
            system=self.system, framework=self.cmmc, name="Protected Multi",
            created_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.cmmc, is_primary=True,
            added_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.iso, added_by=self.assessor,
        )
        ControlAssessment.objects.create(assessment=assessment, requirement=self.cmmc_req)
        ControlAssessment.objects.create(
            assessment=assessment, requirement=self.iso_req,
            status=ControlAssessment.Status.MET,
            assessor_notes_findings="Control is conforming.",
        )
        response = self.client.post(
            reverse("assessment-frameworks", args=("multi", assessment.id)),
            {"name": assessment.name, "frameworks": [self.cmmc.id],
             "primary_framework": self.cmmc.id},
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "cannot be removed")
        self.assertTrue(assessment.frameworks.filter(pk=self.iso.id).exists())

    def test_framework_catalog_displays_metadata(self):
        response = self.client.get(reverse("framework-catalog"))
        self.assertContains(response, "ISO Synthetic")
        self.assertContains(response, "No numeric score")

    def test_crosswalk_preserves_traceability_and_artifact_reuse(self):
        assessment = Assessment.objects.create(
            system=self.system, framework=self.cmmc, name="Mapped Multi",
            created_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.cmmc, is_primary=True,
            added_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=assessment, framework=self.iso, added_by=self.assessor,
        )
        cmmc_result = ControlAssessment.objects.create(
            assessment=assessment, requirement=self.cmmc_req
        )
        iso_result = ControlAssessment.objects.create(
            assessment=assessment, requirement=self.iso_req
        )
        RequirementMapping.objects.create(
            source=self.cmmc_req, target=self.iso_req,
            relationship=RequirementMapping.Relationship.EQUIVALENT,
            mapping_reference="Synthetic crosswalk",
        )
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=assessment,
            title="Shared access evidence", external_reference="https://example.com/shared",
            uploaded_by=self.assessor,
        )
        artifact.controls.add(cmmc_result, iso_result)
        self.assertEqual(artifact.controls.count(), 2)
        response = self.client.get(reverse(
            "control-edit", args=("multi", assessment.id, cmmc_result.id)
        ))
        self.assertContains(response, "ISO-MULTI · AC.1")
        self.assertContains(response, "Equivalent")


class SprintSixExecutionTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.assessor = users.objects.create_user("s6-assessor", password="test-password")
        self.admin = users.objects.create_user("s6-admin", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic S6", slug="s6")
        self.assessor_member = Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.admin_member = Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.system = System.objects.create(organization=self.organization, name="S6 System")
        self.framework = Framework.objects.create(
            code="S6-FW", name="Sprint Six", version="1",
            scoring_method=Framework.ScoringMethod.DEDUCTION, maximum_score=100,
        )
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="S6.1", domain="S6",
            title="Objective control", statement="Meet every objective.", full_deduction=5,
        )
        self.objective_a = AssessmentObjective.objects.create(
            requirement=self.requirement, objective_id="a", text="First objective"
        )
        self.objective_b = AssessmentObjective.objects.create(
            requirement=self.requirement, objective_id="b", text="Second objective"
        )
        AssessmentProcedure.objects.create(
            requirement=self.requirement, method=AssessmentProcedure.Method.EXAMINE,
            sequence=1, assessment_object="Policy",
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework, name="S6 Assessment",
            created_by=self.assessor,
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=self.framework, is_primary=True,
            added_by=self.assessor,
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement
        )
        self.result_a = ObjectiveAssessment.objects.create(
            control_result=self.control, objective=self.objective_a
        )
        self.result_b = ObjectiveAssessment.objects.create(
            control_result=self.control, objective=self.objective_b
        )

    def test_planning_records_scope_team_and_sampling_methodology(self):
        self.client.login(username="s6-assessor", password="test-password")
        url = reverse("assessment-plan", args=("s6", self.assessment.id))
        response = self.client.post(url, {
            "action": "plan", "engagement_start": "2026-08-01",
            "engagement_end": "2026-08-31", "scope_boundaries": "Synthetic boundary",
            "assessment_locations": "Remote", "sampling_methodology": "Random sample",
        })
        self.assertRedirects(response, url)
        self.assessment.refresh_from_db()
        self.assertEqual(self.assessment.sampling_methodology, "Random sample")
        response = self.client.post(url, {
            "action": "team", "team-membership": self.assessor_member.id,
            "team-role": AssessmentTeamMember.Role.LEAD,
        })
        self.assertRedirects(response, url)
        self.assertTrue(self.assessment.team_members.filter(role="LEAD").exists())

    def test_objectives_derive_control_status(self):
        self.client.login(username="s6-assessor", password="test-password")
        url_a = reverse("objective-edit", args=("s6", self.assessment.id, self.result_a.id))
        self.client.post(url_a, {
            "status": ObjectiveAssessment.Status.MET,
            "assessor_notes": "First objective met.",
        })
        url_b = reverse("objective-edit", args=("s6", self.assessment.id, self.result_b.id))
        self.client.post(url_b, {
            "status": ObjectiveAssessment.Status.NOT_MET,
            "assessor_notes": "Second objective failed.",
        })
        self.control.refresh_from_db()
        self.assertEqual(self.control.status, ControlAssessment.Status.NOT_MET)
        self.assertEqual(self.control.calculated_deduction, 5)

    def test_sample_cannot_exceed_population(self):
        self.client.login(username="s6-assessor", password="test-password")
        response = self.client.post(reverse(
            "sample-create", args=("s6", self.assessment.id)
        ), {
            "name": "Invalid sample", "population_description": "Accounts",
            "population_size": 5, "sample_size": 6,
            "selection_method": "Random",
        })
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Cannot exceed the population size")
        self.assertFalse(AssessmentSample.objects.exists())

    def test_signoff_requires_review_then_locks_and_admin_reopens_with_reason(self):
        ObjectiveAssessment.objects.filter(control_result=self.control).update(
            status=ObjectiveAssessment.Status.MET, assessor_notes="Met"
        )
        self.control.derive_from_objectives()
        self.assessment.quality_review_status = "APPROVED"
        self.assessment.save(update_fields=("quality_review_status",))
        self.client.login(username="s6-assessor", password="test-password")
        response = self.client.post(reverse(
            "assessment-signoff", args=("s6", self.assessment.id)
        ))
        self.assertRedirects(response, reverse(
            "quality-review", args=("s6", self.assessment.id)
        ))
        self.assessment.refresh_from_db()
        self.assertTrue(self.assessment.locked)
        edit = self.client.get(reverse(
            "objective-edit", args=("s6", self.assessment.id, self.result_a.id)
        ))
        self.assertEqual(edit.status_code, 404)
        self.client.login(username="s6-admin", password="test-password")
        reopen = self.client.post(reverse(
            "assessment-reopen", args=("s6", self.assessment.id)
        ), {"reason": "New evidence requires reassessment."})
        self.assertRedirects(reopen, reverse(
            "assessment-dashboard", args=("s6", self.assessment.id)
        ))
        self.assessment.refresh_from_db()
        self.assertFalse(self.assessment.locked)
        self.assertEqual(self.assessment.reopen_reason, "New evidence requires reassessment.")


class SprintSevenDashboardTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user("executive", password="test-password")
        self.outsider = user_model.objects.create_user("outside-exec", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic Analytics", slug="analytics")
        other = Organization.objects.create(name="Other Analytics", slug="other-analytics")
        self.member = Membership.objects.create(
            user=self.user, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        Membership.objects.create(user=self.outsider, organization=other, role=Membership.Role.VIEWER)
        system = System.objects.create(organization=self.organization, name="Analytics System")
        self.framework = Framework.objects.create(
            code="AN-FW", name="Analytics Framework", version="1",
            scoring_method=Framework.ScoringMethod.DEDUCTION, maximum_score=100,
        )
        met_requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="AC-1", domain="AC",
            title="Access", statement="Control access.", full_deduction=5,
        )
        gap_requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="IR-1", domain="IR",
            title="Respond", statement="Respond to incidents.", full_deduction=10,
        )
        self.assessment = Assessment.objects.create(
            system=system, framework=self.framework, name="Executive Assessment",
            created_by=self.user, status=Assessment.Status.IN_PROGRESS,
            quality_review_status="IN_REVIEW",
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=self.framework, is_primary=True,
            added_by=self.user,
        )
        met = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=met_requirement,
            status=ControlAssessment.Status.MET,
            implementation_state=ControlAssessment.Implementation.FULL,
            primary_owner=self.member,
        )
        gap = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=gap_requirement,
            status=ControlAssessment.Status.NOT_MET,
            implementation_state=ControlAssessment.Implementation.NONE,
            primary_owner=self.member,
        )
        objective = AssessmentObjective.objects.create(
            requirement=met_requirement, objective_id="a", text="Verify access"
        )
        ObjectiveAssessment.objects.create(
            control_result=met, objective=objective, status=ObjectiveAssessment.Status.MET,
            assessed_by=self.user,
        )
        EvidenceRequest.objects.create(
            assessment=self.assessment, title="Overdue evidence",
            due_date=date.today() - timedelta(days=1), created_by=self.user,
        )
        RemediationPlan.objects.create(
            assessment=self.assessment, remediation_id="RAP-001", title="Close incident gap",
            weakness_description="Synthetic finding", created_by=self.user,
            date_identified=date.today(), planned_completion=date.today() - timedelta(days=1),
            priority=RemediationPlan.Priority.HIGH,
        ).controls.add(gap)

    def test_dashboard_shows_executive_analytics_and_domain_filter(self):
        self.client.login(username="executive", password="test-password")
        response = self.client.get(reverse(
            "assessment-dashboard", args=("analytics", self.assessment.id)
        ))
        self.assertContains(response, "Executive assessment dashboard")
        self.assertContains(response, "Objective completion")
        self.assertContains(response, "1 overdue")
        self.assertContains(response, "Control-owner workload")
        filtered = self.client.get(reverse(
            "assessment-dashboard", args=("analytics", self.assessment.id)
        ), {"domain": "AC"})
        self.assertContains(filtered, "Access")
        self.assertNotContains(filtered, "Respond")

    def test_dashboard_renders_sprs_weighted_risk_heatmap(self):
        self.framework.scoring_method = Framework.ScoringMethod.SPRS
        self.framework.save(update_fields=("scoring_method",))
        self.assessment.risk_management_enabled = True
        self.assessment.save(update_fields=("risk_management_enabled",))
        self.client.login(username="executive", password="test-password")

        response = self.client.get(reverse(
            "assessment-dashboard", args=("analytics", self.assessment.id)
        ))

        self.assertContains(response, "Weighted risk heatmap")
        self.assertContains(response, "SPRS weights")
        self.assertContains(response, "Control exposure by domain")

    def test_csv_snapshot_is_tenant_scoped_and_audited(self):
        url = reverse("dashboard-export", args=("analytics", self.assessment.id))
        self.client.login(username="outside-exec", password="test-password")
        self.assertEqual(self.client.get(url).status_code, 404)
        self.client.login(username="executive", password="test-password")
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "text/csv; charset=utf-8")
        content = response.content.decode("utf-8")
        self.assertIn("Omni Executive Assessment Snapshot", content)
        self.assertIn("AN-FW,AC-1,AC,MET", content)
        self.assertTrue(AuditEvent.objects.filter(
            organization=self.organization, action="dashboard.exported"
        ).exists())


class SprintEightWorkflowTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.assessor = user_model.objects.create_user(
            "workflow-assessor", email="assessor@example.test", password="test-password"
        )
        self.owner = user_model.objects.create_user(
            "workflow-owner", email="owner@example.test", password="test-password"
        )
        self.organization = Organization.objects.create(name="Workflow Test", slug="workflow")
        Membership.objects.create(
            user=self.assessor, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        self.owner_member = Membership.objects.create(
            user=self.owner, organization=self.organization, role=Membership.Role.CLIENT
        )
        system = System.objects.create(organization=self.organization, name="Workflow System")
        framework = Framework.objects.create(code="WF", name="Workflow", version="1")
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="WF-1", domain="WF",
            title="Workflow", statement="Operate workflow."
        )
        self.assessment = Assessment.objects.create(
            system=system, framework=framework, name="Workflow Assessment",
            created_by=self.assessor,
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=requirement
        )

    @override_settings(
        OMNI_EMAIL_ENABLED=True,
        EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
        DEFAULT_FROM_EMAIL="Omni by R!SC <omni@example.test>",
    )
    def test_assignment_creates_in_app_email_audit_and_workflow_history(self):
        NotificationPreference.objects.create(
            user=self.owner, delivery=NotificationPreference.Delivery.EMAIL
        )
        self.client.login(username="workflow-assessor", password="test-password")
        response = self.client.post(reverse(
            "evidence-request-create", args=("workflow", self.assessment.id)
        ), {
            "title": "Access policy", "description": "Provide the policy",
            "status": EvidenceRequest.Status.REQUESTED, "owner": self.owner_member.id,
            "due_date": date.today() + timedelta(days=7), "notify_owner": "on",
            "controls": [self.control.id],
        })
        self.assertRedirects(response, reverse(
            "evidence-list", args=("workflow", self.assessment.id)
        ))
        notice = Notification.objects.get(recipient=self.owner)
        self.assertEqual(notice.email_status, Notification.EmailStatus.SENT)
        self.assertEqual(len(mail.outbox), 1)
        self.assertNotIn("Provide the policy", mail.outbox[0].body)
        self.assertTrue(WorkflowHistory.objects.filter(event="evidence_request.assigned").exists())

    def test_notification_center_marks_only_own_notification_read(self):
        item = Notification.objects.create(
            recipient=self.owner, organization=self.organization, assessment=self.assessment,
            category=Notification.Category.SYSTEM, title="Assigned", message="Open Omni"
        )
        self.client.login(username="workflow-assessor", password="test-password")
        response = self.client.post(reverse("notification-read", args=(item.id,)))
        self.assertEqual(response.status_code, 404)
        item.refresh_from_db()
        self.assertIsNone(item.read_at)
        self.client.login(username="workflow-owner", password="test-password")
        self.client.post(reverse("notification-read", args=(item.id,)))
        item.refresh_from_db()
        self.assertIsNotNone(item.read_at)

    def test_rejected_evidence_requires_review_comment(self):
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment, title="Synthetic link",
            external_reference="https://example.test/evidence", uploaded_by=self.owner,
        )
        self.client.login(username="workflow-assessor", password="test-password")
        response = self.client.post(reverse(
            "evidence-artifact-edit", args=("workflow", self.assessment.id, artifact.id)
        ), {
            "title": artifact.title, "external_reference": artifact.external_reference,
            "review_status": EvidenceArtifact.ReviewStatus.REJECTED,
        })
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Explain why the evidence was rejected")

    def test_reminder_command_is_idempotent_per_day(self):
        EvidenceRequest.objects.create(
            assessment=self.assessment, title="Due request", owner=self.owner_member,
            due_date=date.today(), created_by=self.assessor,
        )
        call_command("send_workflow_reminders")
        call_command("send_workflow_reminders")
        self.assertEqual(Notification.objects.filter(
            recipient=self.owner, category=Notification.Category.DEADLINE
        ).count(), 1)


class SprintEightOneGovernanceTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_user(
            "policy-admin", email="admin@example.test", password="test-password"
        )
        self.owner = user_model.objects.create_user(
            "policy-owner", email="owner@example.test", password="test-password"
        )
        self.lead = user_model.objects.create_user(
            "policy-lead", email="lead@example.test", password="test-password"
        )
        self.organization = Organization.objects.create(name="Policy Test", slug="policy")
        Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.owner_member = Membership.objects.create(
            user=self.owner, organization=self.organization, role=Membership.Role.CLIENT
        )
        lead_member = Membership.objects.create(
            user=self.lead, organization=self.organization, role=Membership.Role.ASSESSOR
        )
        system = System.objects.create(
            organization=self.organization, name="Policy System",
            system_owner_name="Client Owner", system_owner_email="client-owner@example.test",
        )
        framework = Framework.objects.create(code="POL", name="Policy", version="1")
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="POL-1", domain="POL",
            title="Policy", statement="Apply policy."
        )
        self.assessment = Assessment.objects.create(
            system=system, framework=framework, name="Policy Assessment", created_by=self.admin
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=requirement
        )
        AssessmentTeamMember.objects.create(
            assessment=self.assessment, membership=lead_member,
            role=AssessmentTeamMember.Role.LEAD,
        )

    def test_per_action_toggle_suppresses_assignment_notification(self):
        self.client.login(username="policy-admin", password="test-password")
        self.client.post(reverse(
            "evidence-request-create", args=("policy", self.assessment.id)
        ), {
            "title": "Silent request", "status": EvidenceRequest.Status.REQUESTED,
            "owner": self.owner_member.id, "due_date": date.today() + timedelta(days=7),
            "controls": [self.control.id],
        })
        self.assertFalse(Notification.objects.filter(recipient=self.owner).exists())

    @override_settings(
        OMNI_EMAIL_ENABLED=True,
        EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
        DEFAULT_FROM_EMAIL="Omni <omni@example.test>",
    )
    def test_daily_digest_queues_then_sends(self):
        NotificationPreference.objects.create(
            user=self.owner, delivery=NotificationPreference.Delivery.DAILY
        )
        from .notifications import notify
        notify(
            recipients=[self.owner], organization=self.organization,
            assessment=self.assessment, category=Notification.Category.ASSIGNMENT,
            title="Digest assignment", message="Open Omni", actor=self.admin,
        )
        notice = Notification.objects.get(recipient=self.owner)
        self.assertEqual(notice.email_status, Notification.EmailStatus.QUEUED)
        call_command("send_notification_digests")
        notice.refresh_from_db()
        self.assertEqual(notice.email_status, Notification.EmailStatus.SENT)
        self.assertEqual(len(mail.outbox), 1)

    @override_settings(
        OMNI_EMAIL_ENABLED=True,
        EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
        DEFAULT_FROM_EMAIL="Omni <omni@example.test>",
    )
    def test_overdue_escalates_to_owner_lead_and_system_owner(self):
        NotificationPolicy.objects.create(
            organization=self.organization,
            escalation_recipients=NotificationPolicy.Escalation.CLIENT,
        )
        EvidenceRequest.objects.create(
            assessment=self.assessment, title="Escalated request", owner=self.owner_member,
            due_date=date.today() - timedelta(days=1), created_by=self.admin,
        )
        call_command("send_workflow_reminders")
        recipients = set(Notification.objects.values_list("recipient__username", flat=True))
        self.assertEqual(recipients, {"policy-owner", "policy-lead"})
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, ["client-owner@example.test"])

    def test_org_admin_can_update_notification_policy(self):
        self.client.login(username="policy-admin", password="test-password")
        response = self.client.post(reverse("notification-policy", args=("policy",)), {
            "first_reminder_days": 10, "second_reminder_days": 5,
            "notify_on_due_date": "on", "overdue_escalation_days": 2,
            "repeat_overdue_days": 14,
            "escalation_recipients": NotificationPolicy.Escalation.CLIENT,
        })
        self.assertRedirects(response, reverse("notification-policy", args=("policy",)))
        policy = self.organization.notification_policy
        self.assertEqual(policy.repeat_overdue_days, 14)
        self.assertFalse(policy.notifications_enabled)


class SprintNineAccessGovernanceTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_user(
            "access-admin", email="admin@example.test", password="test-password"
        )
        self.member_user = user_model.objects.create_user(
            "access-member", email="member@example.test", password="test-password"
        )
        self.other_user = user_model.objects.create_user(
            "access-other", email="other@example.test", password="test-password"
        )
        self.organization = Organization.objects.create(name="Access Test", slug="access")
        self.admin_member = Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.member = Membership.objects.create(
            user=self.member_user, organization=self.organization, role=Membership.Role.CLIENT
        )
        self.other_member = Membership.objects.create(
            user=self.other_user, organization=self.organization, role=Membership.Role.CLIENT
        )
        self.system = System.objects.create(organization=self.organization, name="Access System")
        framework = Framework.objects.create(code="ACCESS", name="Access", version="1")
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="ACCESS-1", domain="AC",
            title="Access", statement="Restrict access."
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=framework, name="Restricted Assessment",
            created_by=self.admin,
        )
        self.control = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=requirement, primary_owner=self.member
        )

    @override_settings(
        OMNI_EMAIL_ENABLED=True,
        EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
        DEFAULT_FROM_EMAIL="Omni <omni@example.test>",
    )
    @patch("webapp.views.secrets.token_urlsafe", return_value="fixed-secure-token")
    def test_invitation_email_acceptance_creates_account_and_membership(self, token_mock):
        self.client.login(username="access-admin", password="test-password")
        response = self.client.post(reverse("membership-list", args=("access",)), {
            "action": "invite", "email": "new-user@example.test",
            "role": Membership.Role.VIEWER,
        })
        self.assertRedirects(response, reverse("membership-list", args=("access",)))
        invitation = OrganizationInvitation.objects.get(email="new-user@example.test")
        self.assertNotIn("fixed-secure-token", invitation.token_digest)
        self.assertEqual(len(mail.outbox), 1)
        self.client.logout()
        accepted = self.client.post(reverse("invitation-accept", args=("fixed-secure-token",)), {
            "first_name": "New", "last_name": "User",
            "password1": "A-Strong-Test-Password-482!",
            "password2": "A-Strong-Test-Password-482!",
        })
        self.assertRedirects(accepted, reverse("organization-list"))
        invitation.refresh_from_db()
        self.assertEqual(invitation.status, OrganizationInvitation.Status.ACCEPTED)
        self.assertTrue(Membership.objects.filter(
            user__email="new-user@example.test", organization=self.organization,
            role=Membership.Role.VIEWER, active=True,
        ).exists())

    def test_expired_invitation_is_rejected(self):
        invitation = OrganizationInvitation.objects.create(
            organization=self.organization, email="expired@example.test",
            role=Membership.Role.CLIENT, token_digest=hashlib.sha256(b"expired").hexdigest(),
            invited_by=self.admin, expires_at=timezone.now() - timedelta(minutes=1),
        )
        response = self.client.get(reverse("invitation-accept", args=("expired",)))
        self.assertEqual(response.status_code, 410)
        invitation.refresh_from_db()
        self.assertEqual(invitation.status, OrganizationInvitation.Status.PENDING)

    def test_explicit_assessment_access_restricts_other_members(self):
        AssessmentAccess.objects.create(
            assessment=self.assessment, membership=self.member,
            access=AssessmentAccess.Access.VIEW, granted_by=self.admin,
        )
        url = reverse("assessment-dashboard", args=("access", self.assessment.id))
        self.client.login(username="access-member", password="test-password")
        self.assertEqual(self.client.get(url).status_code, 200)
        self.client.login(username="access-other", password="test-password")
        self.assertEqual(self.client.get(url).status_code, 404)
        self.client.login(username="access-admin", password="test-password")
        self.assertEqual(self.client.get(url).status_code, 200)

    def test_last_admin_cannot_be_deactivated(self):
        self.client.login(username="access-admin", password="test-password")
        response = self.client.post(reverse(
            "membership-toggle", args=("access", self.admin_member.id)
        ), {"confirm": "yes"})
        self.assertRedirects(response, reverse("membership-list", args=("access",)))
        self.admin_member.refresh_from_db()
        self.assertTrue(self.admin_member.active)

    def test_assignment_warning_precedes_deactivation(self):
        self.client.login(username="access-admin", password="test-password")
        url = reverse("membership-toggle", args=("access", self.member.id))
        self.client.post(url)
        self.member.refresh_from_db()
        self.assertTrue(self.member.active)
        self.client.post(url, {"confirm": "yes"})
        self.member.refresh_from_db()
        self.assertFalse(self.member.active)

    def test_profile_and_access_review_export(self):
        self.client.login(username="access-admin", password="test-password")
        profile = self.client.post(reverse("user-profile"), {
            "first_name": "Access", "last_name": "Admin", "email": "admin@example.test",
            "job_title": "GRC Administrator", "phone": "555-0100",
            "time_zone": "America/New_York",
        })
        self.assertRedirects(profile, reverse("user-profile"))
        self.assertEqual(UserProfile.objects.get(user=self.admin).job_title, "GRC Administrator")
        export = self.client.get(reverse("access-review-export", args=("access",)))
        self.assertEqual(export.status_code, 200)
        self.assertIn("access-member", export.content.decode())
        self.assertTrue(AuditEvent.objects.filter(action="access_review.exported").exists())


class SprintTenLocalSecurityTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_user(
            "security-admin", email="security@example.test", password="test-password"
        )
        self.organization = Organization.objects.create(name="Security Test", slug="security")
        Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )

    @override_settings(OMNI_LOGIN_FAILURE_LIMIT=3, OMNI_LOGIN_LOCKOUT_MINUTES=15)
    def test_login_throttle_blocks_repeated_failures(self):
        url = reverse("login")
        for _ in range(3):
            self.client.post(url, {"username": "security-admin", "password": "wrong"})
        blocked = self.client.post(url, {
            "username": "security-admin", "password": "test-password"
        })
        self.assertEqual(blocked.status_code, 200)
        self.assertContains(blocked, "Too many unsuccessful attempts")
        attempt = LoginAttempt.objects.get(identifier="security-admin")
        self.assertIsNotNone(attempt.blocked_until)

    def test_security_headers_are_present(self):
        response = self.client.get(reverse("login"))
        self.assertIn("default-src 'self'", response["Content-Security-Policy"])
        self.assertEqual(response["Permissions-Policy"], "camera=(), microphone=(), geolocation=()")
        self.assertEqual(response["X-Content-Type-Options"], "nosniff")

    @override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
    def test_password_reset_uses_generic_response_and_sends_secure_link(self):
        known = self.client.post(reverse("password_reset"), {"email": "security@example.test"})
        unknown = self.client.post(reverse("password_reset"), {"email": "unknown@example.test"})
        self.assertEqual(known.status_code, unknown.status_code)
        self.assertEqual(known.url, unknown.url)
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("/accounts/reset/", mail.outbox[0].body)
        self.assertNotIn("test-password", mail.outbox[0].body)

    def test_backup_and_verification_commands(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            database = root / "source.sqlite3"
            connection = sqlite3.connect(database)
            connection.execute("CREATE TABLE sample (value TEXT)")
            connection.execute("INSERT INTO sample VALUES ('synthetic')")
            connection.commit()
            connection.close()
            media = root / "private"
            media.mkdir()
            (media / "evidence.txt").write_text("synthetic evidence", encoding="utf-8")
            backup = root / "backups"
            with patch.dict(settings.DATABASES["default"], {"NAME": str(database)}), patch.object(
                settings, "MEDIA_ROOT", media
            ), patch.object(settings, "OMNI_BACKUP_DIR", backup):
                call_command("backup_omni")
                archive = next(backup.glob("*.zip"))
                call_command("verify_omni_backup", str(archive))
                with zipfile.ZipFile(archive) as source:
                    self.assertIn("private_uploads/evidence.txt", source.namelist())

    def test_expired_invitations_and_admin_health(self):
        OrganizationInvitation.objects.create(
            organization=self.organization, email="expired@example.test",
            role=Membership.Role.VIEWER, token_digest="e" * 64, invited_by=self.admin,
            expires_at=timezone.now() - timedelta(days=1),
        )
        call_command("expire_invitations")
        self.assertEqual(
            OrganizationInvitation.objects.get().status,
            OrganizationInvitation.Status.EXPIRED,
        )
        self.client.login(username="security-admin", password="test-password")
        response = self.client.get(reverse("system-health", args=("security",)))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Future deployment gate")


class SprintElevenPilotAcceptanceTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_superuser(
            "pilot-admin", email="pilot-admin@example.test", password="test-password"
        )
        framework = Framework.objects.create(
            code="PILOT", name="Pilot Framework", version="RC1",
            scoring_method=Framework.ScoringMethod.DEDUCTION, maximum_score=100,
        )
        requirement = Requirement.objects.create(
            framework=framework, requirement_id="PILOT-1", domain="AC",
            title="Synthetic pilot control", statement="Perform the synthetic pilot control.",
            full_deduction=5,
        )
        AssessmentObjective.objects.create(
            requirement=requirement, objective_id="a", text="Verify the synthetic objective."
        )

    def test_synthetic_pilot_is_idempotent_and_complete_enough_for_uat(self):
        call_command("seed_local_pilot", "--confirm-synthetic", "--created-by", "pilot-admin")
        call_command("seed_local_pilot", "--confirm-synthetic", "--created-by", "pilot-admin")
        organization = Organization.objects.get(slug="omni-synthetic-pilot")
        self.assertIn("NOT CLIENT DATA", organization.name)
        self.assertEqual(organization.systems.count(), 1)
        assessment = organization.systems.get().assessments.get()
        self.assertEqual(assessment.control_results.count(), 1)
        self.assertEqual(ObjectiveAssessment.objects.filter(
            control_result__assessment=assessment
        ).count(), 1)
        self.assertEqual(assessment.evidence_requests.count(), 1)
        self.assertEqual(assessment.remediation_plans.count(), 1)

    def test_pilot_dashboard_and_workbook_open_without_client_data(self):
        call_command("seed_local_pilot", "--confirm-synthetic", "--created-by", "pilot-admin")
        assessment = Assessment.objects.get(system__organization__slug="omni-synthetic-pilot")
        self.client.login(username="pilot-admin", password="test-password")
        dashboard = self.client.get(reverse(
            "assessment-dashboard", args=("omni-synthetic-pilot", assessment.id)
        ))
        self.assertEqual(dashboard.status_code, 200)
        self.assertContains(dashboard, "Executive assessment dashboard")
        from .reporting import build_assessment_workbook
        workbook = load_workbook(BytesIO(build_assessment_workbook(assessment)))
        self.assertIn("Dashboard", workbook.sheetnames)
        self.assertIn("Objective Results", workbook.sheetnames)


class SprintTwelveFrameworkImportTests(TestCase):
    def setUp(self):
        self.temp_media = tempfile.TemporaryDirectory()
        self.override = override_settings(MEDIA_ROOT=self.temp_media.name)
        self.override.enable()
        user_model = get_user_model()
        self.admin = user_model.objects.create_superuser("catalog-admin", "admin@example.test", "test-password")
        self.user = user_model.objects.create_user("regular-user", password="test-password")

    def tearDown(self):
        self.override.disable()
        self.temp_media.cleanup()

    def test_csv_dry_run_requires_explicit_approval_and_preserves_provenance(self):
        self.client.login(username="catalog-admin", password="test-password")
        source = b"Control ID,Domain,Title,Statement\nCCF-1,Governance,Policy,Maintain a security policy.\n"
        response = self.client.post(reverse("framework-import-upload"), {
            "code": "CCF-2026", "name": "Comprehensive Controls Framework",
            "version": "2026.1", "authority": "R!SC",
            "source_file": SimpleUploadedFile("ccf.csv", source, content_type="text/csv"),
        })
        self.assertEqual(response.status_code, 302)
        job = FrameworkImport.objects.get()
        self.assertEqual(job.status, FrameworkImport.Status.PREVIEW)
        self.assertFalse(Framework.objects.filter(code="CCF-2026").exists())
        response = self.client.post(reverse("framework-import-preview", args=(job.id,)))
        self.assertRedirects(response, reverse("framework-catalog"))
        requirement = Requirement.objects.get(framework__code="CCF-2026")
        self.assertEqual(requirement.source_row, 2)
        self.assertEqual(requirement.framework.source_sha256, hashlib.sha256(source).hexdigest())

    def test_existing_code_is_never_overwritten(self):
        Framework.objects.create(code="CCF", name="Existing", version="1")
        self.client.login(username="catalog-admin", password="test-password")
        source = b"id,title,statement\n1,One,Statement\n"
        self.client.post(reverse("framework-import-upload"), {
            "code": "CCF", "name": "Replacement", "version": "2",
            "source_file": SimpleUploadedFile("ccf.csv", source),
        })
        job = FrameworkImport.objects.get()
        response = self.client.post(reverse("framework-import-preview", args=(job.id,)), follow=True)
        self.assertContains(response, "never overwrites")
        self.assertEqual(Framework.objects.get(code="CCF").name, "Existing")

    def test_framework_ingestion_is_superuser_only(self):
        self.client.login(username="regular-user", password="test-password")
        self.assertEqual(self.client.get(reverse("framework-import-list")).status_code, 404)

    def test_excel_adapter_normalizes_a_standard_control_sheet(self):
        from openpyxl import Workbook
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Requirement ID", "Family", "Requirement Title", "Requirement"])
        sheet.append(["AC-1", "AC", "Access policy", "Establish access control policy."])
        stream = BytesIO()
        workbook.save(stream)
        upload = SimpleUploadedFile("framework.xlsx", stream.getvalue())
        normalized, report, source_format, _ = parse_upload(upload, {
            "code": "TEST", "name": "Test", "version": "1"
        })
        self.assertTrue(report["valid"])
        self.assertEqual(source_format, FrameworkImport.SourceFormat.XLSX)
        self.assertEqual(normalized["requirements"][0]["requirement_id"], "AC-1")

    def test_ccf_matrix_adapter_retains_cross_framework_references(self):
        from openpyxl import Workbook
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["CCF Domain", "CCF Control", "CCF #",
                      "Comprehensive Controls Framework (CCF)\nControl Description",
                      "CCF Control Question", "NIST SP 800-171", "ISO 27001"])
        sheet.append(["Governance", "Security policy", "GOV-01", "Maintain policy.",
                      "Is policy maintained?", "3.12.4", "5.1"])
        stream = BytesIO()
        workbook.save(stream)
        normalized, report, _, _ = parse_upload(
            SimpleUploadedFile("ccf.xlsx", stream.getvalue()),
            {"code": "CCF", "name": "Comprehensive Controls Framework", "version": "2026.2"},
        )
        self.assertTrue(report["valid"])
        self.assertEqual(report["mapping_reference_count"], 2)
        self.assertEqual(report["authority_count"], 2)
        self.assertEqual(normalized["requirements"][0]["mapping_refs"][0]["source_column"], 6)
        self.assertEqual(normalized["requirements"][0]["mapping_refs"][0]["target_requirement"], "3.12.4")

    def test_scanned_pdf_is_flagged_for_ocr(self):
        from pypdf import PdfWriter
        writer, stream = PdfWriter(), BytesIO()
        writer.add_blank_page(width=72, height=72)
        writer.write(stream)
        _, report, source_format, _ = parse_upload(
            SimpleUploadedFile("scan.pdf", stream.getvalue()),
            {"code": "SCAN", "name": "Scan", "version": "1"},
        )
        self.assertEqual(source_format, FrameworkImport.SourceFormat.PDF)
        self.assertFalse(report["valid"])
        self.assertIn("OCR_REQUIRED", {issue["code"] for issue in report["issues"]})


class SprintThirteenHarmonizationTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_user("harmonizer", password="test-password")
        self.outsider = user_model.objects.create_user("harmonizer-outsider", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic Harmonization", slug="synthetic-harmonization")
        self.other = Organization.objects.create(name="Other Synthetic", slug="other-synthetic")
        Membership.objects.create(user=self.admin, organization=self.organization, role=Membership.Role.ADMIN)
        Membership.objects.create(user=self.outsider, organization=self.other, role=Membership.Role.ADMIN)
        self.system = System.objects.create(organization=self.organization, name="Synthetic system")
        self.hub = Framework.objects.create(
            code="OMNI-CF", name="Omni Control Framework", version="2026.2",
            is_omni_control_framework=True,
        )
        self.framework_a = Framework.objects.create(code="FRAME-A", name="Framework A", version="1")
        self.framework_b = Framework.objects.create(code="FRAME-B", name="Framework B", version="1")
        self.hub_requirement = Requirement.objects.create(
            framework=self.hub, requirement_id="OMNI-1", domain="Governance",
            title="Policy governance", statement="Maintain policy governance.",
        )
        self.requirement_a = Requirement.objects.create(
            framework=self.framework_a, requirement_id="A-1", domain="Governance",
            title="Policy A", statement="Maintain policy A.",
        )
        self.requirement_b = Requirement.objects.create(
            framework=self.framework_b, requirement_id="B-1", domain="Governance",
            title="Policy B", statement="Maintain policy B.",
        )
        RequirementMapping.objects.create(
            source=self.hub_requirement, target=self.requirement_a,
            relationship=RequirementMapping.Relationship.EQUIVALENT,
        )
        RequirementMapping.objects.create(
            source=self.hub_requirement, target=self.requirement_b,
            relationship=RequirementMapping.Relationship.PARTIAL,
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework_a, name="Synthetic multi-framework",
            created_by=self.admin,
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=self.framework_a, is_primary=True,
            added_by=self.admin,
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=self.framework_b, added_by=self.admin,
        )
        self.result_a = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement_a,
            status=ControlAssessment.Status.MET, updated_by=self.admin,
        )
        self.result_b = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement_b,
        )

    def test_omni_hub_derives_reviewable_relationship_without_propagating_result(self):
        summary = refresh_harmonization(self.assessment)
        self.assertEqual(summary["created"], 1)
        decision = AssessmentReuseDecision.objects.get()
        self.assertEqual(decision.basis, AssessmentReuseDecision.Basis.OMNI_DERIVED)
        self.assertEqual(decision.relationship, RequirementMapping.Relationship.PARTIAL)
        self.assertEqual(decision.mapping_path, ["A-1", "OMNI-1", "B-1"])
        self.assertEqual(decision.status, AssessmentReuseDecision.Status.SUGGESTED)
        self.result_b.refresh_from_db()
        self.assertEqual(self.result_b.status, ControlAssessment.Status.NOT_ASSESSED)

    def test_approval_reuses_only_accepted_evidence_and_never_control_outcome(self):
        accepted = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment, title="Accepted policy",
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED, uploaded_by=self.admin,
        )
        rejected = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment, title="Rejected draft",
            review_status=EvidenceArtifact.ReviewStatus.REJECTED, uploaded_by=self.admin,
        )
        accepted.controls.add(self.result_a)
        rejected.controls.add(self.result_a)
        refresh_harmonization(self.assessment)
        decision = AssessmentReuseDecision.objects.get()
        linked = review_reuse(decision, self.admin, True)
        self.assertEqual(linked, 1)
        self.assertTrue(accepted.controls.filter(pk=self.result_b.pk).exists())
        self.assertFalse(rejected.controls.filter(pk=self.result_b.pk).exists())
        self.result_b.refresh_from_db()
        self.assertEqual(self.result_b.status, ControlAssessment.Status.NOT_ASSESSED)

    def test_harmonization_page_is_tenant_scoped_and_records_review_audit(self):
        refresh_harmonization(self.assessment)
        decision = AssessmentReuseDecision.objects.get()
        self.client.login(username="harmonizer-outsider", password="test-password")
        url = reverse("assessment-harmonization", args=(self.organization.slug, self.assessment.pk))
        self.assertEqual(self.client.get(url).status_code, 404)
        self.client.login(username="harmonizer", password="test-password")
        response = self.client.post(url, {
            "action": "approve", "decision_id": decision.pk,
            "reuse_evidence": "on", "reuse_testing": "on", "rationale": "Validated overlap.",
        })
        self.assertRedirects(response, url)
        decision.refresh_from_db()
        self.assertEqual(decision.status, AssessmentReuseDecision.Status.APPROVED)
        self.assertTrue(AuditEvent.objects.filter(action="harmonization.approved").exists())

    def test_retained_omni_mapping_resolves_after_external_framework_exists(self):
        FrameworkImport.objects.create(
            source_file=SimpleUploadedFile("omni-map.csv", b"synthetic"),
            source_filename="omni-map.csv", source_format=FrameworkImport.SourceFormat.CSV,
            source_sha256="0" * 64, status=FrameworkImport.Status.IMPORTED,
            normalized_data={"framework": {"code": self.hub.code}, "requirements": [{
                "requirement_id": self.hub_requirement.requirement_id,
                "source_reference": "synthetic matrix row 2",
                "mapping_refs": [{
                    "target_framework": self.framework_b.name,
                    "target_requirement": self.requirement_b.requirement_id,
                }],
            }]},
            validation_report={"valid": True}, created_by=self.admin,
            approved_by=self.admin, approved_at=timezone.now(), imported_framework=self.hub,
        )
        RequirementMapping.objects.filter(
            source=self.hub_requirement, target=self.requirement_b
        ).delete()
        result = resolve_catalog_mappings(self.admin)
        self.assertEqual(result["created"], 1)
        mapping = RequirementMapping.objects.get(
            source=self.hub_requirement, target=self.requirement_b
        )
        self.assertEqual(mapping.source_reference, "synthetic matrix row 2")


class SprintFourteenMappingCurationTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.admin = user_model.objects.create_superuser("mapping-admin", "mapping@example.test", "test-password")
        self.user = user_model.objects.create_user("mapping-user", password="test-password")
        self.framework = Framework.objects.create(
            code="OMNI-CF-TEST", name="Omni Control Framework", version="test",
            is_omni_control_framework=True,
        )
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="GOV-01", domain="Governance",
            title="Governance", statement="Maintain governance.",
        )
        self.job = FrameworkImport.objects.create(
            source_file=SimpleUploadedFile("omni.xlsx", b"synthetic"), source_filename="omni.xlsx",
            source_format=FrameworkImport.SourceFormat.XLSX, source_sha256="1" * 64,
            status=FrameworkImport.Status.IMPORTED, normalized_data={"framework": {}, "requirements": [{
                "requirement_id": "GOV-01", "source_row": 2,
                "mapping_refs": [{"target_framework": "Example Standard", "target_requirement": "A.1\nA.2", "source_column": 6}],
            }]}, validation_report={"valid": True}, created_by=self.admin,
            approved_by=self.admin, approved_at=timezone.now(), imported_framework=self.framework,
        )

    def test_materialization_preserves_raw_cell_and_provenance(self):
        result = materialize_mapping_references(self.job)
        self.assertEqual(result, {"authorities": 1, "references": 1})
        reference = MappingReference.objects.get()
        self.assertEqual(reference.raw_reference, "A.1\nA.2")
        self.assertEqual(reference.source_row, 2)
        self.assertEqual(reference.source_column, 6)
        self.assertEqual(reference.authority.aliases, ["Example Standard"])

    def test_mapping_quality_is_superuser_only_and_bulk_reviewed(self):
        materialize_mapping_references(self.job)
        reference = MappingReference.objects.get()
        url = reverse("mapping-quality")
        self.client.login(username="mapping-user", password="test-password")
        self.assertEqual(self.client.get(url).status_code, 404)
        self.client.login(username="mapping-admin", password="test-password")
        response = self.client.post(url, {"action": "approve", "selected": [reference.pk]})
        self.assertRedirects(response, url)
        reference.refresh_from_db()
        self.assertEqual(reference.review_status, MappingReference.ReviewStatus.APPROVED)
        self.assertEqual(reference.reviewed_by, self.admin)


class SprintFifteenSharedWorkTests(SprintThirteenHarmonizationTests):
    def test_shared_workspace_records_applicability_without_changing_conclusion(self):
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment, title="Shared policy",
            external_reference="https://example.test/policy",
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED, uploaded_by=self.admin,
        )
        artifact.controls.add(self.result_a)
        refresh_harmonization(self.assessment)
        decision = AssessmentReuseDecision.objects.get()
        review_reuse(decision, self.admin, True)
        self.client.login(username="harmonizer", password="test-password")
        url = reverse("shared-work-workspace", args=(self.organization.slug, self.assessment.pk))
        response = self.client.post(url, {"action": "applicability", "artifact_id": artifact.pk,
            "control_id": self.result_b.pk, "applicability": "PARTIAL",
            "rationale": "Same policy, supplemental test needed.", "scope_limitations": "Framework-specific clause."})
        self.assertRedirects(response, url)
        self.assertEqual(EvidenceApplicability.objects.get().applicability, "PARTIAL")
        self.result_b.refresh_from_db()
        self.assertEqual(self.result_b.status, ControlAssessment.Status.NOT_ASSESSED)

    def test_request_consolidation_preserves_source_requests(self):
        first = EvidenceRequest.objects.create(assessment=self.assessment, title="Policy copy", created_by=self.admin)
        second = EvidenceRequest.objects.create(assessment=self.assessment, title="Policy approval", created_by=self.admin)
        first.controls.add(self.result_a); second.controls.add(self.result_b)
        self.client.login(username="harmonizer", password="test-password")
        url = reverse("shared-work-workspace", args=(self.organization.slug, self.assessment.pk))
        self.client.post(url, {"action": "consolidate", "requests": [first.pk, second.pk]})
        first.refresh_from_db(); second.refresh_from_db()
        consolidated = [item for item in (first, second) if item.consolidated_into_id]
        primary = [item for item in (first, second) if not item.consolidated_into_id]
        self.assertEqual(len(consolidated), 1)
        self.assertEqual(len(primary), 1)
        self.assertEqual(primary[0].controls.count(), 2)

    def test_multiframework_reports_and_traceability_are_generated(self):
        from .reporting import build_multi_framework_report, build_traceability_csv
        report = build_multi_framework_report(self.assessment)
        framework_report = build_multi_framework_report(self.assessment, self.framework_a)
        traceability = build_traceability_csv(self.assessment)
        self.assertTrue(report.startswith(b"PK"))
        self.assertTrue(framework_report.startswith(b"PK"))
        self.assertIn(b"FRAME-A", traceability)
        self.assertIn(b"A-1", traceability)


class SprintSeventeenMappingGovernanceTests(TestCase):
    def setUp(self):
        users = get_user_model(); self.author = users.objects.create_user("map-author"); self.reviewer = users.objects.create_user("map-reviewer")
        first = Framework.objects.create(code="MAP-A", name="Map A", version="1")
        second = Framework.objects.create(code="MAP-B", name="Map B", version="1")
        source = Requirement.objects.create(framework=first, requirement_id="A", domain="G", title="A", statement="A")
        target = Requirement.objects.create(framework=second, requirement_id="B", domain="G", title="B", statement="B")
        self.mapping = RequirementMapping.objects.create(source=source, target=target, relationship="RELATED")
        self.change = MappingChangeRequest.objects.create(mapping=self.mapping, proposed_relationship="PARTIAL", proposed_rationale="Overlap only", reason="Source changed", requested_by=self.author)

    def test_independent_approval_versions_mapping_and_preserves_history(self):
        with self.assertRaises(ValueError): review_change(self.change, self.author, True)
        review_change(self.change, self.reviewer, True)
        self.mapping.refresh_from_db(); self.change.refresh_from_db()
        self.assertEqual(self.mapping.relationship, "PARTIAL")
        self.assertEqual(self.mapping.revision, 2)
        self.assertEqual(self.change.status, "APPROVED")
        self.assertEqual(MappingHistory.objects.filter(mapping=self.mapping).count(), 2)


class SprintSeventeenPointFiveEvidenceCatalogTests(TestCase):
    def test_private_catalog_import_preserves_source_and_normalizes_cmmc_alias(self):
        from openpyxl import Workbook
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "evidence.xlsx"; workbook = Workbook(); sheet = workbook.active
            sheet.append(["#", "ERL #", "Area of Focus", "Documentation Artifact", "Artifact Description", "SCF Control Mappings", "Relevant CMMC"])
            sheet.append([1, "E-1", "Governance", "Security Plan", "Plan evidence", "GOV-01\nGOV-02", "AC.L1-3.1.1"])
            sheet.append([2, "E-2", "Governance", "Novel Evidence", "Novel evidence", "GOV-03", ""]); workbook.save(path)
            report = import_catalog(path, apply=True)
        self.assertEqual(report["requests"], 2); self.assertEqual(report["exact"], 1)
        exact = OmniEvidenceSourceRequest.objects.get(source_identifier="E-1")
        self.assertEqual(exact.canonical_evidence_code, "EV-0001")
        self.assertEqual(exact.source_cmmc_ids, ["AC.L1-3.1.1"])
        self.assertEqual(exact.normalized_cmmc_ids, ["AC.L2-3.1.1"])
        self.assertEqual(OmniEvidenceSourceRequest.objects.get(source_identifier="E-2").resolution, "REVIEW")

    def test_catalog_curation_is_superuser_only(self):
        user = get_user_model().objects.create_user("catalog-reader", password="test-password")
        self.client.login(username="catalog-reader", password="test-password")
        self.assertEqual(self.client.get(reverse("omni-evidence-catalog")).status_code, 404)


class SprintNineteenFoundationTests(TestCase):
    def setUp(self):
        self.admin = get_user_model().objects.create_user(
            "s19-admin", password="test-password"
        )
        self.organization = Organization.objects.create(name="Synthetic S19", slug="s19")
        Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.system = System.objects.create(
            organization=self.organization, name="S19 Synthetic System"
        )
        self.framework = Framework.objects.create(
            code="S19-FW", name="Sprint Nineteen Framework", version="1"
        )
        self.requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="S19-1", domain="Governance",
            title="Synthetic requirement", statement="Maintain a synthetic control."
        )
        self.objective = AssessmentObjective.objects.create(
            requirement=self.requirement, objective_id="a", text="Evaluate the synthetic control."
        )
        self.assessment = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Prior S19 Assessment",
            created_by=self.admin, scope_boundaries="Synthetic boundary",
        )
        AssessmentFramework.objects.create(
            assessment=self.assessment, framework=self.framework,
            is_primary=True, added_by=self.admin,
        )
        result = ControlAssessment.objects.create(
            assessment=self.assessment, requirement=self.requirement,
            status=ControlAssessment.Status.MET, assessor_notes_findings="Prior conclusion",
        )
        ObjectiveAssessment.objects.create(
            control_result=result, objective=self.objective,
            status=ObjectiveAssessment.Status.MET, assessor_notes="Prior objective conclusion",
        )
        request = EvidenceRequest.objects.create(
            assessment=self.assessment, title="Synthetic policy", created_by=self.admin
        )
        request.controls.add(result)
        self.client.login(username="s19-admin", password="test-password")

    def test_template_creates_fresh_assessment_without_conclusions(self):
        create_url = reverse("assessment-template-create", args=(self.organization.slug,))
        response = self.client.post(create_url, {
            "source_assessment": self.assessment.id,
            "name": "Annual synthetic review", "description": "Reusable configuration",
            "primary_framework": self.framework.id, "frameworks": [self.framework.id],
            "scope_boundaries": "Synthetic boundary", "assessment_locations": "Remote",
            "sampling_methodology": "Representative sample",
            "notifications_enabled": "on", "email_notifications_enabled": "on",
            "recurrence": AssessmentTemplate.Recurrence.ANNUAL,
            "next_start_date": "2026-09-01", "default_duration_days": 30, "active": "on",
        })
        self.assertRedirects(response, reverse(
            "assessment-template-list", args=(self.organization.slug,)
        ))
        template = AssessmentTemplate.objects.get(name="Annual synthetic review")
        self.assertEqual(len(template.evidence_request_blueprints), 1)
        instantiate_url = reverse(
            "assessment-from-template",
            args=(self.organization.slug, self.system.id, template.id),
        )
        response = self.client.post(instantiate_url, {
            "name": "2027 Synthetic Review", "engagement_start": "2026-09-01",
            "engagement_end": "2026-09-30", "prior_assessment": self.assessment.id,
        })
        created = Assessment.objects.get(name="2027 Synthetic Review")
        self.assertRedirects(response, reverse(
            "assessment-dashboard", args=(self.organization.slug, created.id)
        ))
        self.assertEqual(created.prior_assessment, self.assessment)
        self.assertEqual(created.source_template, template)
        self.assertFalse(created.control_results.exclude(
            status=ControlAssessment.Status.NOT_ASSESSED
        ).exists())
        self.assertFalse(ObjectiveAssessment.objects.filter(
            control_result__assessment=created
        ).exclude(status=ObjectiveAssessment.Status.NOT_ASSESSED).exists())
        self.assertEqual(created.evidence_requests.count(), 1)
        self.assertEqual(created.evidence_artifacts.count(), 0)
        template.refresh_from_db()
        self.assertEqual(template.next_start_date, date(2027, 9, 1))

    def test_integration_policy_is_admin_scoped_and_connector_stays_inactive(self):
        url = reverse("integration-settings", args=(self.organization.slug,))
        response = self.client.post(url, {
            "delivery": IntegrationPolicy.Delivery.EXTERNAL,
            "provider": IntegrationPolicy.Provider.JIRA,
            "external_ticketing_enabled": "on", "create_for_remediation": "on",
        }, follow=True)
        self.assertEqual(response.status_code, 200)
        policy = IntegrationPolicy.objects.get(organization=self.organization)
        self.assertEqual(policy.provider, IntegrationPolicy.Provider.JIRA)
        self.assertTrue(policy.external_ticketing_enabled)
        self.assertEqual(OutboundWorkItem.objects.count(), 0)
        self.assertContains(response, "Foundation only")

    def test_policy_driven_freshness_uses_effective_date(self):
        evidence_request = self.assessment.evidence_requests.get()
        evidence_request.freshness_days = 365
        evidence_request.renewal_lead_days = 30
        evidence_request.save()
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Annual policy", external_reference="https://example.invalid/policy",
            effective_on=timezone.localdate() - timedelta(days=350),
            uploaded_by=self.admin,
        )
        artifact.requests.add(evidence_request)
        self.assertEqual(
            artifact.freshness_deadline,
            artifact.effective_on + timedelta(days=365),
        )
        self.assertEqual(artifact.freshness, "AGING")

    def test_manual_renewal_is_idempotent_and_preserves_request_context(self):
        evidence_request = self.assessment.evidence_requests.get()
        evidence_request.freshness_days = 90
        evidence_request.renewal_lead_days = 20
        evidence_request.save()
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Quarterly access review",
            external_reference="https://example.invalid/access-review",
            effective_on=timezone.localdate() - timedelta(days=80),
            uploaded_by=self.admin,
        )
        artifact.requests.add(evidence_request)
        url = reverse("evidence-artifact-renew", args=(
            self.organization.slug, self.assessment.id, artifact.id,
        ))
        self.client.post(url)
        self.client.post(url)
        renewal = EvidenceRequest.objects.get(renewal_of=evidence_request)
        self.assertEqual(renewal.title, "Renew: Synthetic policy")
        self.assertEqual(
            list(renewal.controls.values_list("id", flat=True)),
            list(evidence_request.controls.values_list("id", flat=True)),
        )
        self.assertEqual(evidence_request.renewal_requests.count(), 1)
        self.assertTrue(AuditEvent.objects.filter(
            action="evidence_artifact.renewal_requested", object_id=str(artifact.id)
        ).exists())

    def test_workflow_command_opens_automatic_renewal_once(self):
        evidence_request = self.assessment.evidence_requests.get()
        evidence_request.freshness_days = 30
        evidence_request.renewal_lead_days = 10
        evidence_request.auto_renew = True
        evidence_request.save()
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Monthly evidence", external_reference="https://example.invalid/monthly",
            effective_on=timezone.localdate() - timedelta(days=25),
            uploaded_by=self.admin,
        )
        artifact.requests.add(evidence_request)
        call_command("send_workflow_reminders")
        call_command("send_workflow_reminders")
        self.assertEqual(evidence_request.renewal_requests.count(), 1)

    def test_monitoring_event_creates_reassessment_without_changing_conclusion(self):
        control = self.assessment.control_results.get()
        url = reverse("control-monitoring-event-create", args=(
            self.organization.slug, self.assessment.id,
        ))
        response = self.client.post(url, {
            "title": "Material authentication change",
            "event_type": ControlMonitoringEvent.EventType.CHANGE,
            "severity": ControlMonitoringEvent.Severity.HIGH,
            "occurred_on": timezone.localdate(),
            "description": "The identity provider configuration changed.",
            "source_reference": "CHG-SYNTHETIC-1",
            "controls": [control.id],
        })
        self.assertRedirects(response, reverse(
            "control-monitoring", args=(self.organization.slug, self.assessment.id)
        ))
        task = ControlReassessmentTask.objects.get()
        self.assertEqual(task.prior_conclusion["status"], ControlAssessment.Status.MET)
        control.refresh_from_db()
        self.assertEqual(control.status, ControlAssessment.Status.MET)
        task_url = reverse("control-reassessment-task-edit", args=(
            self.organization.slug, self.assessment.id, task.id,
        ))
        invalid = self.client.post(task_url, {
            "status": ControlReassessmentTask.Status.COMPLETED,
            "due_date": task.due_date,
        })
        self.assertEqual(invalid.status_code, 200)
        self.client.post(task_url, {
            "status": ControlReassessmentTask.Status.COMPLETED,
            "due_date": task.due_date,
            "resolution": "Retested; the existing conclusion remains supported.",
        })
        task.refresh_from_db(); task.event.refresh_from_db()
        self.assertEqual(task.status, ControlReassessmentTask.Status.COMPLETED)
        self.assertEqual(task.event.status, ControlMonitoringEvent.Status.REVIEWED)

    def test_scheduler_creates_duplicate_safe_periodic_reassessment(self):
        control = self.assessment.control_results.get()
        ControlMonitoringProfile.objects.create(
            control_result=control, enabled=True, review_frequency_days=30,
            next_review_date=timezone.localdate(), updated_by=self.admin,
        )
        call_command("send_workflow_reminders")
        call_command("send_workflow_reminders")
        self.assertEqual(ControlMonitoringEvent.objects.filter(
            event_type=ControlMonitoringEvent.EventType.SCHEDULED
        ).count(), 1)
        self.assertEqual(ControlReassessmentTask.objects.count(), 1)
        profile = control.monitoring_profile
        profile.refresh_from_db()
        self.assertEqual(profile.next_review_date, timezone.localdate() + timedelta(days=30))

    def test_expired_evidence_triggers_affected_controls_only(self):
        control = self.assessment.control_results.get()
        artifact = EvidenceArtifact.objects.create(
            organization=self.organization, assessment=self.assessment,
            title="Expired configuration", external_reference="https://example.invalid/config",
            expires_on=timezone.localdate() - timedelta(days=1), uploaded_by=self.admin,
        )
        artifact.controls.add(control)
        call_command("send_workflow_reminders")
        event = ControlMonitoringEvent.objects.get(
            event_type=ControlMonitoringEvent.EventType.EVIDENCE
        )
        self.assertEqual(list(event.controls.all()), [control])
        self.assertEqual(event.reassessment_tasks.count(), 1)

    def test_unsigned_assessment_cannot_be_captured_as_baseline(self):
        response = self.client.post(reverse("assessment-baseline-create", args=(
            self.organization.slug, self.assessment.id,
        )), {"name": "Invalid draft", "description": "Not signed off"})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Only a signed-off assessment")
        self.assertFalse(AssessmentBaseline.objects.exists())

    def test_approved_baseline_is_immutable_comparison_point(self):
        self.assessment.locked = True
        self.assessment.status = Assessment.Status.COMPLETE
        self.assessment.quality_review_status = "APPROVED"
        self.assessment.signed_off_by = self.admin
        self.assessment.signed_off_at = timezone.now()
        self.assessment.save()
        create_url = reverse("assessment-baseline-create", args=(
            self.organization.slug, self.assessment.id,
        ))
        self.client.post(create_url, {
            "name": "Approved synthetic baseline", "description": "Period one",
        })
        baseline = AssessmentBaseline.objects.get()
        self.assertEqual(baseline.status, AssessmentBaseline.Status.DRAFT)
        self.client.post(reverse("assessment-baseline-approve", args=(
            self.organization.slug, self.assessment.id, baseline.id,
        )))
        baseline.refresh_from_db()
        self.assertEqual(baseline.status, AssessmentBaseline.Status.APPROVED)
        control = self.assessment.control_results.get()
        control.status = ControlAssessment.Status.NOT_MET
        control.implementation_state = ControlAssessment.Implementation.NONE
        control.assessor_notes_findings = "A new synthetic finding."
        control.save()
        response = self.client.get(reverse("assessment-baselines", args=(
            self.organization.slug, self.assessment.id,
        )), {"baseline": baseline.id})
        self.assertContains(response, "REGRESSED")
        self.assertEqual(
            next(iter(baseline.snapshot["controls"].values()))["status"],
            ControlAssessment.Status.MET,
        )
        self.assertTrue(AuditEvent.objects.filter(
            action="assessment_baseline.approved", object_id=str(baseline.id)
        ).exists())
        export = self.client.get(reverse("assessment-baseline-comparison-export", args=(
            self.organization.slug, self.assessment.id, baseline.id,
        )))
        self.assertEqual(export.status_code, 200)
        self.assertIn(b"REGRESSED", export.content)
        self.assertIn(baseline.checksum.encode(), export.content)
        preserved_snapshot = baseline.snapshot
        self.client.post(reverse("assessment-baseline-retire", args=(
            self.organization.slug, self.assessment.id, baseline.id,
        )))
        baseline.refresh_from_db()
        self.assertEqual(baseline.status, AssessmentBaseline.Status.RETIRED)
        self.assertEqual(baseline.snapshot, preserved_snapshot)
        baseline.snapshot = {"tampered": True}
        with self.assertRaisesRegex(ValueError, "immutable"):
            baseline.save()

    def test_baseline_integrity_failure_prevents_approval(self):
        from .assessment_baselines import assessment_snapshot
        self.assessment.locked = True
        self.assessment.quality_review_status = "APPROVED"
        self.assessment.save()
        snapshot = assessment_snapshot(self.assessment)
        baseline = AssessmentBaseline.objects.create(
            assessment=self.assessment, name="Tampered baseline", snapshot=snapshot,
            checksum="0" * 64, created_by=self.admin,
        )
        response = self.client.post(reverse("assessment-baseline-approve", args=(
            self.organization.slug, self.assessment.id, baseline.id,
        )), follow=True)
        baseline.refresh_from_db()
        self.assertEqual(baseline.status, AssessmentBaseline.Status.DRAFT)
        self.assertContains(response, "failed its integrity check")


class SprintNineteenPointFivePortfolioTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.admin = users.objects.create_user("portfolio-admin", password="test-password")
        self.viewer = users.objects.create_user("portfolio-viewer", password="test-password")
        self.other = users.objects.create_user("portfolio-other", password="test-password")
        self.organization = Organization.objects.create(name="Synthetic Portfolio", slug="portfolio")
        Membership.objects.create(
            user=self.admin, organization=self.organization, role=Membership.Role.ADMIN
        )
        self.viewer_membership = Membership.objects.create(
            user=self.viewer, organization=self.organization, role=Membership.Role.VIEWER
        )
        self.other_membership = Membership.objects.create(
            user=self.other, organization=self.organization, role=Membership.Role.VIEWER
        )
        self.system = System.objects.create(
            organization=self.organization, name="Visible Portfolio System"
        )
        self.framework = Framework.objects.create(
            code="PORT-FW", name="Portfolio Framework", version="1"
        )
        requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="PORT-1", domain="Governance",
            title="Portfolio requirement", statement="Synthetic statement",
        )
        self.visible = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Visible Assessment",
            status=Assessment.Status.IN_PROGRESS, created_by=self.admin,
        )
        self.hidden = Assessment.objects.create(
            system=self.system, framework=self.framework, name="Restricted Assessment",
            status=Assessment.Status.COMPLETE, created_by=self.admin,
        )
        for assessment in (self.visible, self.hidden):
            AssessmentFramework.objects.create(
                assessment=assessment, framework=self.framework,
                is_primary=True, added_by=self.admin,
            )
        AssessmentAccess.objects.create(
            assessment=self.hidden, membership=self.other_membership,
            access=AssessmentAccess.Access.VIEW, granted_by=self.admin,
        )
        ControlAssessment.objects.create(
            assessment=self.visible, requirement=requirement,
            status=ControlAssessment.Status.MET,
            implementation_state=ControlAssessment.Implementation.FULL,
        )
        ControlAssessment.objects.create(
            assessment=self.hidden, requirement=requirement,
            status=ControlAssessment.Status.NOT_MET,
            implementation_state=ControlAssessment.Implementation.NONE,
        )
        self.outside = Organization.objects.create(name="Outside Tenant", slug="outside-portfolio")
        outside_system = System.objects.create(
            organization=self.outside, name="Outside Confidential System"
        )
        Assessment.objects.create(
            system=outside_system, framework=self.framework,
            name="Outside Confidential Assessment", created_by=self.admin,
        )

    def test_viewer_portfolio_respects_assessment_grants(self):
        self.client.login(username="portfolio-viewer", password="test-password")
        response = self.client.get(reverse("portfolio-dashboard", args=(self.organization.slug,)))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Visible Assessment")
        self.assertNotContains(response, "Restricted Assessment")
        self.assertNotContains(response, "Outside Confidential")
        self.assertEqual(response.context["analytics"]["metrics"]["assessments"], 1)
        self.assertEqual(response.context["analytics"]["metrics"]["findings"], 0)

    def test_admin_filters_and_export_are_tenant_scoped_and_audited(self):
        self.client.login(username="portfolio-admin", password="test-password")
        url = reverse("portfolio-dashboard", args=(self.organization.slug,))
        response = self.client.get(url, {"status": Assessment.Status.COMPLETE})
        self.assertEqual(response.context["analytics"]["metrics"]["assessments"], 1)
        self.assertEqual(response.context["analytics"]["metrics"]["findings"], 1)
        export = self.client.get(
            reverse("portfolio-export", args=(self.organization.slug,)),
            {"status": Assessment.Status.COMPLETE},
        )
        self.assertEqual(export.status_code, 200)
        self.assertIn(b"Restricted Assessment", export.content)
        self.assertNotIn(b"Outside Confidential", export.content)
        self.assertTrue(AuditEvent.objects.filter(action="portfolio.exported").exists())

    def test_unassigned_tenant_portfolio_is_not_found(self):
        self.client.login(username="portfolio-viewer", password="test-password")
        response = self.client.get(reverse("portfolio-dashboard", args=(self.outside.slug,)))
        self.assertEqual(response.status_code, 404)


class SprintNineteenPointSixAutomationTests(TestCase):
    def setUp(self):
        users = get_user_model()
        self.admin = users.objects.create_user("calendar-admin", password="test-password")
        self.viewer = users.objects.create_user("calendar-viewer", password="test-password")
        self.restricted_user = users.objects.create_user(
            "calendar-restricted", password="test-password"
        )
        self.org_a = Organization.objects.create(name="Calendar Alpha", slug="calendar-alpha")
        self.org_b = Organization.objects.create(name="Calendar Beta", slug="calendar-beta")
        self.admin_a = Membership.objects.create(
            user=self.admin, organization=self.org_a, role=Membership.Role.ADMIN
        )
        self.admin_b = Membership.objects.create(
            user=self.admin, organization=self.org_b, role=Membership.Role.ADMIN
        )
        self.viewer_a = Membership.objects.create(
            user=self.viewer, organization=self.org_a, role=Membership.Role.VIEWER
        )
        self.restricted_a = Membership.objects.create(
            user=self.restricted_user, organization=self.org_a, role=Membership.Role.VIEWER
        )
        self.framework = Framework.objects.create(
            code="CAL-FW", name="Calendar Framework", version="1"
        )
        requirement = Requirement.objects.create(
            framework=self.framework, requirement_id="CAL-1", domain="Governance",
            title="Calendar requirement", statement="Synthetic calendar control",
        )
        self.assessments = {}
        for organization, membership, suffix in (
            (self.org_a, self.admin_a, "Alpha"), (self.org_b, self.admin_b, "Beta")
        ):
            system = System.objects.create(
                organization=organization, name=f"{suffix} Calendar System"
            )
            assessment = Assessment.objects.create(
                system=system, framework=self.framework,
                name=f"{suffix} Calendar Assessment", created_by=self.admin,
                engagement_start=timezone.localdate(),
                engagement_end=timezone.localdate() + timedelta(days=30),
            )
            AssessmentFramework.objects.create(
                assessment=assessment, framework=self.framework,
                is_primary=True, added_by=self.admin,
            )
            result = ControlAssessment.objects.create(
                assessment=assessment, requirement=requirement
            )
            evidence = EvidenceRequest.objects.create(
                assessment=assessment, title=f"{suffix} evidence due",
                owner=membership, due_date=timezone.localdate(), created_by=self.admin,
            )
            evidence.controls.add(result)
            self.assessments[organization.slug] = assessment
        hidden = Assessment.objects.create(
            system=self.assessments[self.org_a.slug].system, framework=self.framework,
            name="Restricted Calendar Assessment", created_by=self.admin,
            engagement_end=timezone.localdate() + timedelta(days=2),
        )
        AssessmentFramework.objects.create(
            assessment=hidden, framework=self.framework, is_primary=True, added_by=self.admin
        )
        AssessmentAccess.objects.create(
            assessment=hidden, membership=self.restricted_a,
            access=AssessmentAccess.Access.VIEW, granted_by=self.admin,
        )
        EvidenceRequest.objects.create(
            assessment=hidden, title="Restricted calendar evidence",
            due_date=timezone.localdate(), created_by=self.admin,
        )
        self.hidden = hidden

    def test_calendar_and_export_are_tenant_and_grant_scoped(self):
        self.client.login(username="calendar-viewer", password="test-password")
        url = reverse("compliance-calendar", args=(self.org_a.slug,))
        response = self.client.get(url, {"category": "EVIDENCE"})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Alpha evidence due")
        self.assertNotContains(response, "Beta evidence due")
        self.assertNotContains(response, "Restricted calendar evidence")
        export = self.client.get(
            reverse("compliance-calendar-export", args=(self.org_a.slug,)),
            {"category": "EVIDENCE"},
        )
        self.assertIn(b"Alpha evidence due", export.content)
        self.assertNotIn(b"Beta evidence due", export.content)
        self.assertNotIn(b"Restricted calendar evidence", export.content)
        self.assertTrue(AuditEvent.objects.filter(
            organization=self.org_a, action="compliance_calendar.exported"
        ).exists())

    def test_automation_runner_is_organization_scoped_and_audited(self):
        for organization in (self.org_a, self.org_b):
            ComplianceAutomationPolicy.objects.create(
                organization=organization, enabled=True,
                next_run_on=timezone.localdate(), updated_by=self.admin,
            )
        call_command(
            "run_compliance_automation", organization_slug=self.org_a.slug, force=True
        )
        self.assertEqual(ComplianceAutomationRun.objects.filter(
            policy__organization=self.org_a,
            status=ComplianceAutomationRun.Status.SUCCESS,
        ).count(), 1)
        self.assertEqual(ComplianceAutomationRun.objects.filter(
            policy__organization=self.org_b
        ).count(), 0)
        self.assertTrue(Notification.objects.filter(
            organization=self.org_a, title__contains="due today"
        ).exists())
        self.assertFalse(Notification.objects.filter(organization=self.org_b).exists())
        self.assertTrue(AuditEvent.objects.filter(
            organization=self.org_a, action="compliance_automation.completed"
        ).exists())

    def test_automation_settings_are_admin_only_and_disabled_by_default(self):
        self.client.login(username="calendar-viewer", password="test-password")
        self.assertEqual(self.client.get(reverse(
            "compliance-automation-settings", args=(self.org_a.slug,)
        )).status_code, 404)
        self.client.login(username="calendar-admin", password="test-password")
        response = self.client.get(reverse(
            "compliance-automation-settings", args=(self.org_a.slug,)
        ))
        self.assertEqual(response.status_code, 200)
        policy = ComplianceAutomationPolicy.objects.get(organization=self.org_a)
        self.assertFalse(policy.enabled)
