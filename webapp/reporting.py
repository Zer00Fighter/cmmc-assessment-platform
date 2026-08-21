from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import tempfile
import zipfile
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from openpyxl import load_workbook
from docx import Document

from src.ssp_export import SSPExportMetadata, export_ssp
from src.workbook import WorkbookBuilder

from .models import (
    AuthoritativeDocument, ControlAssessment, EvidenceArtifact, ObjectiveAssessment,
    RemediationPlan, Soc2AssessmentProfile,
)
from .soc2_activity_import import TSC_FRAMEWORK_CODE
from .remediation_export import build_remediation_workbook


class ReportNotReady(ValueError):
    def __init__(self, issues):
        self.issues = issues
        super().__init__("Assessment deliverable is not ready.")


def assessment_readiness(assessment, *, require_template=False) -> dict:
    blockers, warnings = [], []
    organization = assessment.system.organization
    required = {
        "Organization name": organization.name,
        "System name": assessment.system.name,
        "Assessment scope": assessment.system.scope,
        "CAGE code": assessment.system.cage_code,
    }
    for field, value in required.items():
        if not str(value or "").strip():
            blockers.append(f"{field} is missing.")
    results = assessment.control_results.select_related("requirement")
    for result in results:
        control = result.requirement.requirement_id
        if result.status == ControlAssessment.Status.NOT_ASSESSED:
            blockers.append(f"{control} has not been assessed.")
        elif not result.assessor_notes_findings.strip():
            blockers.append(f"{control} requires an Assessor Notes/Findings statement.")
        if not result.ssp_reference.strip():
            warnings.append(f"{control} has no Security Plan reference.")
        if not result.evidence_artifacts.filter(
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED
        ).exists():
            warnings.append(f"{control} has no accepted supporting artifact.")
        unassessed_objectives = result.objective_results.filter(
            status="NOT_ASSESSED"
        ).count()
        if unassessed_objectives:
            blockers.append(f"{control} has {unassessed_objectives} unassessed objective(s).")
    template = Path(settings.OMNI_SSP_TEMPLATE) if settings.OMNI_SSP_TEMPLATE else None
    if require_template and (not template or not template.is_file()):
        blockers.append(
            "Word SSP template is not configured. Set OMNI_SSP_TEMPLATE to the private template path."
        )
    total = results.count()
    assessed = results.exclude(status=ControlAssessment.Status.NOT_ASSESSED).count()
    return {
        "ready": not blockers,
        "blockers": blockers,
        "warnings": warnings,
        "total_controls": total,
        "assessed_controls": assessed,
        "completion_percent": round(assessed / total * 100, 2) if total else 0,
    }


def multi_framework_readiness(assessment) -> dict:
    report = assessment_readiness(assessment)
    for artifact in assessment.evidence_artifacts.prefetch_related("requests"):
        if artifact.freshness in {"AGING", "EXPIRED", "SUPERSEDED"}:
            report["warnings"].append(f"{artifact.title} is {artifact.freshness.lower()}.")
    for decision in assessment.reuse_decisions.filter(status="APPROVED", reuse_evidence=True):
        if not decision.target_result.evidence_applicability.exists():
            report["blockers"].append(
                f"Reused evidence for {decision.target_result.requirement.requirement_id} needs applicability review."
            )
    report["ready"] = not report["blockers"]
    return report


def build_traceability_csv(assessment) -> bytes:
    output = io.StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(["Framework", "Requirement", "Status", "Mapping basis", "Mapping path",
                     "Evidence", "Evidence applicability", "Testing references", "Finding"])
    decisions = {item.target_result_id: item for item in assessment.reuse_decisions.select_related("target_result")}
    for result in assessment.control_results.select_related("requirement__framework").prefetch_related(
        "evidence_artifacts__applicability_reviews", "objective_results__reused_tests"
    ):
        decision = decisions.get(result.pk)
        evidence = "; ".join(result.evidence_artifacts.values_list("title", flat=True))
        applicability = "; ".join(result.evidence_applicability.values_list("applicability", flat=True))
        tests = sum(item.reused_tests.count() for item in result.objective_results.all())
        writer.writerow([result.requirement.framework.code, result.requirement.requirement_id,
                         result.status, decision.get_basis_display() if decision else "Original",
                         " > ".join(decision.mapping_path) if decision else "", evidence,
                         applicability, tests, result.assessor_notes_findings])
    return output.getvalue().encode("utf-8-sig")


def build_multi_framework_report(assessment, framework=None) -> bytes:
    document = Document()
    title = f"{framework.name} Assessment Report" if framework else "Consolidated Multi-Framework Assessment Report"
    document.add_heading(title, 0)
    document.add_paragraph(f"Organization: {assessment.system.organization.name}")
    document.add_paragraph(f"System: {assessment.system.name}")
    document.add_paragraph(f"Assessment: {assessment.name}")
    selected = [framework] if framework else list(assessment.frameworks.all())
    document.add_heading("Assessment scope and frameworks", level=1)
    document.add_paragraph(assessment.system.scope or "Scope not recorded.")
    document.add_paragraph(", ".join(f"{item.name} {item.version}" for item in selected))
    for item in selected:
        results = assessment.control_results.filter(requirement__framework=item).select_related("requirement")
        document.add_heading(f"{item.name} {item.version}", level=1)
        table = document.add_table(rows=1, cols=5); table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ("Requirement", "Title", "Status", "Evidence", "Conclusion/Finding")):
            cell.text = label
        for result in results:
            cells = table.add_row().cells
            values = (result.requirement.requirement_id, result.requirement.title, result.status,
                      str(result.evidence_artifacts.count()), result.assessor_notes_findings)
            for cell, value in zip(cells, values): cell.text = str(value)
        sources = AuthoritativeDocument.objects.filter(
            authority__canonical_name__iexact=item.name, active=True
        )
        if sources:
            document.add_heading(f"{item.code} authoritative sources", level=2)
            for source in sources:
                document.add_paragraph(f"{source.formal_name} — {source.official_url or 'Source URL pending'}")
    stream = io.BytesIO(); document.save(stream); return stream.getvalue()


def soc2_report_readiness(assessment) -> dict:
    """Validate SOC 2-specific scope and criterion execution before reporting."""
    blockers, warnings = [], []
    try:
        profile = assessment.soc2_profile
    except Soc2AssessmentProfile.DoesNotExist:
        profile = None
    if not profile:
        blockers.append("SOC 2 examination type, scope, and measurement period are not configured.")
        return {"ready": False, "blockers": blockers, "warnings": warnings,
                "total_criteria": 0, "assessed_criteria": 0, "completion_percent": 0}
    if not profile.service_commitments.strip():
        blockers.append("SOC 2 service commitments are not documented.")
    if not (assessment.system.description.strip() or assessment.system.scope.strip()):
        blockers.append("A system description or assessment scope is not documented.")
    results = assessment.control_results.filter(
        in_scope=True, requirement__framework__code=TSC_FRAMEWORK_CODE
    ).select_related("requirement").prefetch_related(
        "evidence_artifacts", "objective_results__reused_tests"
    )
    if not results.exists():
        blockers.append("No in-scope SOC 2 criteria are loaded.")
    for result in results:
        criterion = result.requirement.requirement_id
        if result.status == ControlAssessment.Status.NOT_ASSESSED:
            blockers.append(f"{criterion} has not been assessed.")
        objectives = result.objective_results.all()
        if not objectives:
            blockers.append(f"{criterion} has no assessment objective result.")
        for objective in objectives:
            if objective.status == ObjectiveAssessment.Status.NOT_ASSESSED:
                blockers.append(f"{criterion} objective has not been assessed.")
            if objective.design_conclusion == ObjectiveAssessment.Conclusion.NOT_ASSESSED:
                blockers.append(f"{criterion} requires a design conclusion.")
            if objective.implementation_conclusion == ObjectiveAssessment.Conclusion.NOT_ASSESSED:
                blockers.append(f"{criterion} requires an implementation conclusion.")
            if (profile.examination_type == "TYPE_II" and
                    objective.operating_effectiveness_conclusion ==
                    ObjectiveAssessment.OperatingConclusion.NOT_TESTED):
                blockers.append(f"{criterion} requires an operating-effectiveness conclusion.")
            if (profile.examination_type == "TYPE_I" and
                    objective.operating_effectiveness_conclusion !=
                    ObjectiveAssessment.OperatingConclusion.NOT_APPLICABLE):
                blockers.append(f"{criterion} Type I operating effectiveness must be not applicable.")
            if objective.status != ObjectiveAssessment.Status.NOT_ASSESSED and not objective.assessor_notes.strip():
                blockers.append(f"{criterion} requires an objective conclusion or finding narrative.")
        accepted = result.evidence_artifacts.filter(
            review_status=EvidenceArtifact.ReviewStatus.ACCEPTED
        )
        reused_tests = sum(item.reused_tests.count() for item in objectives)
        if not accepted.exists() and not reused_tests:
            warnings.append(f"{criterion} has no accepted evidence or approved testing reference.")
        if profile.examination_type == "TYPE_II":
            for artifact in accepted:
                if not artifact.period_start or not artifact.period_end:
                    warnings.append(f"{criterion}: {artifact.title} has no evidence period.")
                elif artifact.period_start > profile.period_start or artifact.period_end < profile.period_end:
                    warnings.append(f"{criterion}: {artifact.title} does not cover the full examination period.")
    for decision in assessment.reuse_decisions.filter(
        status="APPROVED", reuse_evidence=True,
        target_result__in=results,
    ).select_related("target_result__requirement"):
        if not decision.target_result.evidence_applicability.exists():
            blockers.append(
                f"{decision.target_result.requirement.requirement_id} reused evidence needs an applicability review."
            )
    total = results.count()
    assessed = results.exclude(status=ControlAssessment.Status.NOT_ASSESSED).count()
    return {"ready": not blockers, "blockers": blockers, "warnings": warnings,
            "total_criteria": total, "assessed_criteria": assessed,
            "completion_percent": round(assessed / total * 100, 2) if total else 0}


def build_soc2_report(assessment) -> bytes:
    readiness = soc2_report_readiness(assessment)
    if not readiness["ready"]:
        raise ReportNotReady(readiness["blockers"])
    profile = assessment.soc2_profile
    results = assessment.control_results.filter(
        in_scope=True, requirement__framework__code=TSC_FRAMEWORK_CODE
    ).select_related("requirement").prefetch_related(
        "evidence_artifacts", "evidence_applicability",
        "objective_results__evidence", "objective_results__reused_tests__source_test",
        "remediation_plans",
    )
    document = Document()
    document.add_heading("SOC 2 Readiness Assessment Work Program", 0)
    document.add_paragraph(f"Organization: {assessment.system.organization.name}")
    document.add_paragraph(f"System: {assessment.system.name}")
    document.add_paragraph(f"Assessment: {assessment.name}")
    document.add_heading("Important limitation", level=1)
    document.add_paragraph(
        "This Omni-generated work program documents readiness-assessment procedures and results. "
        "It is not an AICPA SOC 2 report, attestation opinion, or substitute for an examination "
        "performed and issued by an independent licensed CPA firm."
    )
    document.add_heading("Executive summary", level=1)
    counts = {status: results.filter(status=status).count() for status in (
        ControlAssessment.Status.MET, ControlAssessment.Status.NOT_MET,
        ControlAssessment.Status.NOT_APPLICABLE, ControlAssessment.Status.NOT_ASSESSED,
    )}
    document.add_paragraph(
        f"{profile.get_examination_type_display()} · {readiness['assessed_criteria']} of "
        f"{readiness['total_criteria']} in-scope criteria assessed · "
        f"MET {counts[ControlAssessment.Status.MET]} · "
        f"NOT MET {counts[ControlAssessment.Status.NOT_MET]} · "
        f"N/A {counts[ControlAssessment.Status.NOT_APPLICABLE]}."
    )
    document.add_heading("System description and readiness-assessment scope", level=1)
    document.add_paragraph(assessment.system.description or "System description not recorded.")
    document.add_paragraph(assessment.system.scope or "System scope not recorded.")
    document.add_paragraph("Trust Services categories: " + ", ".join(profile.included_category_labels))
    if profile.examination_type == "TYPE_I":
        document.add_paragraph(f"Measurement date: {profile.as_of_date:%B %d, %Y}")
    else:
        document.add_paragraph(
            f"Examination period: {profile.period_start:%B %d, %Y} through "
            f"{profile.period_end:%B %d, %Y}"
        )
    document.add_paragraph(f"Service commitments: {profile.service_commitments or 'Not recorded.'}")
    document.add_heading("Document request list", level=1)
    if assessment.evidence_requests.exists():
        for request in assessment.evidence_requests.prefetch_related("controls__requirement"):
            controls = ", ".join(item.requirement.requirement_id for item in request.controls.all())
            document.add_paragraph(
                f"{request.title} · {request.get_status_display()} · {controls or 'General'}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No document requests have been recorded.")
    document.add_heading("Criterion results and testing traceability", level=1)
    table = document.add_table(rows=1, cols=8); table.style = "Table Grid"
    labels = ("Criterion", "Category", "Result", "Design", "Implementation",
              "Operating effectiveness", "Evidence / reused tests", "Conclusion or finding")
    for cell, label in zip(table.rows[0].cells, labels): cell.text = label
    for result in results:
        objective = result.objective_results.first()
        evidence = [f"EA-{item.id:04d} {item.title}" for item in result.evidence_artifacts.all()]
        reused = [] if not objective else [
            f"Test {item.source_test_id} ({item.source_test.get_outcome_display()})"
            for item in objective.reused_tests.all()
        ]
        values = (
            result.requirement.requirement_id, result.requirement.domain,
            result.get_status_display(),
            objective.get_design_conclusion_display() if objective else "Not assessed",
            objective.get_implementation_conclusion_display() if objective else "Not assessed",
            objective.get_operating_effectiveness_conclusion_display() if objective else "Not tested",
            "; ".join(evidence + reused) or "None recorded",
            (objective.assessor_notes if objective else result.assessor_notes_findings),
        )
        for cell, value in zip(table.add_row().cells, values): cell.text = str(value or "")
    exceptions = [item for item in results if item.status == ControlAssessment.Status.NOT_MET]
    document.add_heading("Exceptions, findings, and corrective actions", level=1)
    if not exceptions:
        document.add_paragraph("No NOT MET criteria were recorded.")
    for result in exceptions:
        document.add_heading(result.requirement.requirement_id, level=2)
        document.add_paragraph(result.assessor_notes_findings or result.objective_results.first().assessor_notes)
        for plan in result.remediation_plans.all():
            document.add_paragraph(f"{plan.remediation_id}: {plan.title} ({plan.get_status_display()})")
    if readiness["warnings"]:
        document.add_heading("Reporting limitations and warnings", level=1)
        for warning in readiness["warnings"]: document.add_paragraph(warning, style="List Bullet")
    stream = io.BytesIO(); document.save(stream); return stream.getvalue()


def _display_member(membership) -> str:
    if not membership:
        return ""
    return membership.user.get_full_name() or membership.user.username


def build_assessment_workbook(assessment) -> bytes:
    with tempfile.TemporaryDirectory(prefix="omni-report-") as directory:
        path = Path(directory) / "assessment.xlsx"
        WorkbookBuilder(project_root=Path(settings.BASE_DIR), output_path=path).build()
        workbook = load_workbook(path)
        cover = workbook["Cover"]
        cover_values = {
            "Organization Name": assessment.system.organization.name,
            "Assessment Name": assessment.name,
            "Assessment Scope": assessment.system.scope,
            "CAGE Code": assessment.system.cage_code,
            "Assessment Start Date": assessment.created_at.date(),
            "Assessment End Date": timezone.localdate() if assessment.status == "COMPLETE" else "",
            "Lead Assessor": assessment.created_by.get_full_name() or assessment.created_by.username,
        }
        for row in range(1, cover.max_row + 1):
            label = cover.cell(row, 2).value
            if label in cover_values:
                cover.cell(row, 3, cover_values[label])

        results = {
            item.requirement.requirement_id: item
            for item in assessment.control_results.filter(
                requirement__framework=assessment.framework
            ).select_related(
                "requirement", "primary_owner__user"
            ).prefetch_related("evidence_artifacts", "remediation_plans")
        }
        multi_sheet = workbook.create_sheet("Multi-Framework Results", 3)
        multi_headers = (
            "Framework Code", "Framework Name", "Version", "Primary",
            "Requirement ID", "Domain", "Title", "Statement", "Status",
            "Implementation State", "Calculated Deduction", "Control Owner",
            "SSP Reference", "Assessor Notes / Findings", "Evidence References",
        )
        for column, header in enumerate(multi_headers, 1):
            multi_sheet.cell(1, column, header)
        for offset, result in enumerate(
            assessment.control_results.select_related(
                "requirement__framework", "primary_owner__user"
            ).prefetch_related("evidence_artifacts"), 2
        ):
            framework = result.requirement.framework
            artifacts = list(result.evidence_artifacts.all())
            values = (
                framework.code, framework.name, framework.version,
                "Yes" if framework.pk == assessment.framework_id else "No",
                result.requirement.requirement_id, result.requirement.domain,
                result.requirement.title, result.requirement.statement, result.status,
                result.implementation_state,
                result.calculated_deduction if framework.scoring_method != "NONE" else "",
                _display_member(result.primary_owner) or result.control_owner,
                result.ssp_reference, result.assessor_notes_findings,
                "; ".join(f"EA-{item.id:04d} {item.title}" for item in artifacts),
            )
            for column, value in enumerate(values, 1):
                multi_sheet.cell(offset, column, value)
        multi_sheet.freeze_panes = "A2"
        multi_sheet.auto_filter.ref = multi_sheet.dimensions
        for column in range(1, len(multi_headers) + 1):
            multi_sheet.column_dimensions[chr(64 + column)].width = 20
        objective_sheet = workbook.create_sheet("Objective Results", 4)
        objective_headers = (
            "Framework", "Requirement ID", "Objective ID", "Objective Text",
            "Status", "Assessor Notes", "Assessed By", "Assessed At",
            "Evidence References", "Examine Objects", "Interview Objects", "Test Objects",
        )
        for column, header in enumerate(objective_headers, 1):
            objective_sheet.cell(1, column, header)
        objective_results = ObjectiveAssessment.objects.filter(
            control_result__assessment=assessment
        ).select_related(
            "objective__requirement__framework", "assessed_by"
        ).prefetch_related("evidence", "objective__requirement__assessment_procedures")
        for row, result in enumerate(objective_results, 2):
            procedures = list(result.objective.requirement.assessment_procedures.all())
            method_text = lambda method: "; ".join(
                item.assessment_object for item in procedures if item.method == method
            )
            values = (
                result.objective.requirement.framework.code,
                result.objective.requirement.requirement_id, result.objective.objective_id,
                result.objective.text, result.get_status_display(), result.assessor_notes,
                (result.assessed_by.get_full_name() or result.assessed_by.username)
                if result.assessed_by else "", result.assessed_at,
                "; ".join(f"EA-{item.id:04d} {item.title}" for item in result.evidence.all()),
                method_text("EXAMINE"), method_text("INTERVIEW"), method_text("TEST"),
            )
            for column, value in enumerate(values, 1):
                objective_sheet.cell(row, column, value)
        objective_sheet.freeze_panes = "A2"
        objective_sheet.auto_filter.ref = objective_sheet.dimensions
        sheet = workbook["Assessment"]
        crosswalk = workbook["SSP Crosswalk"]
        for row in range(6, sheet.max_row + 1):
            requirement_id = str(sheet.cell(row, 2).value or "").strip()
            result = results.get(requirement_id)
            if not result:
                continue
            artifacts = list(result.evidence_artifacts.all())
            owner = _display_member(result.primary_owner) or result.control_owner
            sheet.cell(row, 9, result.status)
            sheet.cell(row, 10, result.implementation_state)
            sheet.cell(row, 13, "Yes" if result.calculated_deduction == 3 else "No")
            sheet.cell(row, 14, result.calculated_deduction)
            sheet.cell(row, 15, "Complete" if any(
                item.review_status == EvidenceArtifact.ReviewStatus.ACCEPTED for item in artifacts
            ) else ("In Progress" if artifacts else "Not Started"))
            sheet.cell(row, 16, owner)
            sheet.cell(row, 17, result.ssp_reference)
            sheet.cell(row, 18, result.assessor_notes_findings)
            sheet.cell(row, 19, "Yes" if result.remediation_plans.exclude(
                status__in=(RemediationPlan.Status.CLOSED, RemediationPlan.Status.RISK_ACCEPTED)
            ).exists() else "No")
            evidence_refs = "; ".join(
                f"EA-{item.id:04d} {item.title}" for item in artifacts
            )
            crosswalk.cell(row, 3, result.ssp_reference)
            crosswalk.cell(row, 5, evidence_refs)
            crosswalk.cell(row, 6, owner)
            crosswalk.cell(row, 7, "Mapped" if result.ssp_reference and evidence_refs else (
                "Partially Mapped" if result.ssp_reference or evidence_refs else "Not Mapped"
            ))

        evidence_sheet = workbook["Evidence"]
        artifacts = assessment.evidence_artifacts.select_related("uploaded_by").prefetch_related(
            "controls__requirement"
        )
        for offset, artifact in enumerate(artifacts):
            row = 6 + offset
            controls = ", ".join(
                item.requirement.requirement_id for item in artifact.controls.all()
            )
            location = artifact.external_reference or (artifact.file.name if artifact.file else "")
            values = (
                f"EA-{artifact.id:04d}", artifact.title, "Artifact", artifact.assessor_notes,
                location, artifact.source, "Complete", artifact.get_review_status_display(),
                "", artifact.updated_at.date(), artifact.period_end, controls, "", "Confidential",
                "", artifact.assessor_notes,
            )
            for column, value in enumerate(values, 1):
                evidence_sheet.cell(row, column, value)

        remediation = workbook["POA&M"]
        for row in range(6, remediation.max_row + 1):
            for column in range(1, 25):
                remediation.cell(row, column, "")
        plans = assessment.remediation_plans.prefetch_related(
            "controls__requirement", "milestones__owner__user", "closure_evidence"
        ).select_related("owner__user")
        for offset, plan in enumerate(plans):
            row = 6 + offset
            controls = list(plan.controls.all())
            milestone = plan.milestones.exclude(status="COMPLETE").first() or plan.milestones.first()
            dates_end = plan.actual_completion or timezone.localdate()
            days = max((dates_end - plan.date_identified).days, 0)
            values = (
                plan.remediation_id,
                ", ".join(item.requirement.requirement_id for item in controls),
                "; ".join(item.requirement.title for item in controls),
                ", ".join(sorted({item.requirement.domain for item in controls})),
                plan.weakness_description, plan.root_cause, plan.corrective_action,
                milestone.title if milestone else "", str(milestone.owner) if milestone and milestone.owner else "",
                plan.get_status_display(), plan.get_priority_display(), plan.get_severity_display(),
                plan.get_likelihood_display(), plan.risk_score, plan.date_identified,
                plan.planned_completion, plan.actual_completion, days, "", plan.get_residual_risk_display(),
                plan.get_validation_status_display(),
                ", ".join(f"EA-{item.id:04d}" for item in plan.closure_evidence.all()),
                ", ".join(sorted({item.ssp_reference for item in controls if item.ssp_reference})),
                plan.validation_notes,
            )
            for column, value in enumerate(values, 1):
                remediation.cell(row, column, value)
        workbook.calculation.fullCalcOnLoad = True
        workbook.calculation.forceFullCalc = True
        workbook.save(path)
        return path.read_bytes()


def build_word_ssp(assessment, workbook_bytes: bytes, generated_by) -> bytes:
    readiness = assessment_readiness(assessment, require_template=True)
    if not readiness["ready"]:
        raise ReportNotReady(readiness["blockers"])
    with tempfile.TemporaryDirectory(prefix="omni-ssp-") as directory:
        workbook_path = Path(directory) / "assessment.xlsx"
        output_path = Path(directory) / "ssp.docx"
        workbook_path.write_bytes(workbook_bytes)
        export_ssp(
            settings.OMNI_SSP_TEMPLATE, workbook_path, output_path,
            SSPExportMetadata(
                organization_name=assessment.system.organization.name,
                system_name=assessment.system.name,
                system_owner=assessment.system.system_owner_name,
                prepared_by=generated_by.get_full_name() or generated_by.username,
                version="1.0", export_date=timezone.localdate().isoformat(),
            ),
        )
        return output_path.read_bytes()


def _safe_name(value: str, fallback: str) -> str:
    result = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return result[:120] or fallback


def _artifact_reference_text(artifact) -> str:
    controls = ", ".join(item.requirement.requirement_id for item in artifact.controls.all())
    requests = ", ".join(item.title for item in artifact.requests.all())
    remediation = ", ".join(item.remediation_id for item in artifact.remediation_plans.all())
    registered_by = artifact.uploaded_by.get_full_name() or artifact.uploaded_by.username
    lines = {
        "Artifact ID": f"EA-{artifact.id:04d}", "Title": artifact.title,
        "External URL": artifact.external_reference, "Source": artifact.source,
        "Period Start": artifact.period_start or "", "Period End": artifact.period_end or "",
        "Review Status": artifact.get_review_status_display(),
        "Assessor Notes": artifact.assessor_notes, "Linked Requests": requests,
        "Linked Controls": controls, "Linked Remediation Plans": remediation,
        "Registered At": artifact.created_at.isoformat(), "Registered By": registered_by,
    }
    return "\n".join(f"{key}: {value}" for key, value in lines.items()) + "\n"


def build_risk_register_csv(assessment) -> str:
    output = io.StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(["Risk ID", "Title", "Category", "Status", "Likelihood", "Impact",
                     "Inherent Score", "Treatment", "Residual Likelihood", "Residual Impact",
                     "Residual Score", "Trend", "Target Date", "Next Review", "Controls"])
    for risk in assessment.risks.prefetch_related("controls__requirement__framework"):
        controls = "; ".join(
            f"{item.requirement.framework.code} {item.requirement.requirement_id}"
            for item in risk.controls.all()
        )
        writer.writerow([risk.risk_id, risk.title, risk.category, risk.get_status_display(),
                         risk.likelihood, risk.impact, risk.inherent_score,
                         risk.get_treatment_display(), risk.residual_likelihood or "",
                         risk.residual_impact or "", risk.residual_score or "", risk.get_trend_display(),
                         risk.target_date or "", risk.next_review_date or "", controls])
    return output.getvalue()


def build_package(assessment, generated_by) -> tuple[bytes, dict]:
    readiness = assessment_readiness(assessment, require_template=True)
    if not readiness["ready"]:
        raise ReportNotReady(readiness["blockers"])
    workbook = build_assessment_workbook(assessment)
    ssp = build_word_ssp(assessment, workbook, generated_by)
    remediation = build_remediation_workbook(assessment).getvalue()
    output = io.BytesIO()
    manifest_rows = []
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as package:
        package.writestr("Deliverables/Omni-Assessment-Workbook.xlsx", workbook)
        package.writestr("Deliverables/Omni-System-Security-Plan.docx", ssp)
        package.writestr("Deliverables/Omni-Remediation-Action-Plan.xlsx", remediation)
        if assessment.risk_management_enabled and assessment.include_risk_in_reports:
            package.writestr("Deliverables/Omni-Risk-Register.csv", build_risk_register_csv(assessment))
        artifacts = assessment.evidence_artifacts.select_related("uploaded_by").prefetch_related(
            "controls__requirement", "requests", "remediation_plans"
        )
        for artifact in artifacts:
            artifact_id = f"EA-{artifact.id:04d}"
            base = _safe_name(artifact.title, artifact_id)
            controls = ", ".join(item.requirement.requirement_id for item in artifact.controls.all())
            requests = ", ".join(item.title for item in artifact.requests.all())
            remediations = ", ".join(item.remediation_id for item in artifact.remediation_plans.all())
            if artifact.file:
                filename = _safe_name(Path(artifact.file.name).name, f"{artifact_id}.bin")
                with artifact.file.open("rb") as evidence_file:
                    package.writestr(f"Evidence/{artifact_id}_{base}/{filename}", evidence_file.read())
            if artifact.external_reference:
                package.writestr(
                    f"Evidence/{artifact_id}_{base}/{artifact_id}_External_Reference.txt",
                    _artifact_reference_text(artifact),
                )
            manifest_rows.append({
                "artifact_id": artifact_id, "title": artifact.title,
                "review_status": artifact.get_review_status_display(),
                "freshness": artifact.freshness,
                "freshness_deadline": artifact.freshness_deadline or "",
                "controls": controls, "requests": requests,
                "remediation_plans": remediations,
                "uploaded_file": Path(artifact.file.name).name if artifact.file else "",
                "external_reference": artifact.external_reference,
            })
        csv_buffer = io.StringIO()
        writer = csv.DictWriter(csv_buffer, fieldnames=(
            "artifact_id", "title", "review_status", "freshness", "freshness_deadline",
            "controls", "requests",
            "remediation_plans", "uploaded_file", "external_reference",
        ))
        writer.writeheader()
        writer.writerows(manifest_rows)
        package.writestr("Evidence/Evidence-Index.csv", csv_buffer.getvalue())
        package.writestr("Package-Manifest.json", json.dumps({
            "assessment": assessment.name,
            "organization": assessment.system.organization.name,
            "system": assessment.system.name,
            "generated_at": timezone.now().isoformat(),
            "generated_by": generated_by.get_full_name() or generated_by.username,
            "readiness": readiness, "evidence_count": len(manifest_rows),
            "risk_management_enabled": assessment.risk_management_enabled,
            "risk_reporting_included": (
                assessment.risk_management_enabled and assessment.include_risk_in_reports
            ),
        }, indent=2))
    return output.getvalue(), readiness


def digest(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()
